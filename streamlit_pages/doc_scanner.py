# streamlit_pages/doc_scanner.py
"""
AI-Сканер документов
Поддерживаемые форматы: PDF (текстовый и скан), DOCX, DOC, XLSX, JPG, PNG, TXT
OCR: EasyOCR (основной) → Tesseract (fallback)
Пересказ: LM Studio
Поиск: текстовый с синонимами через QueryExpander
"""
import os
import io
import re
import json
import time
import hashlib
import threading
from datetime import datetime
from typing import List, Dict, Optional, Tuple

import streamlit as st

# =============================================================================
# Утилиты — загрузка/сохранение базы сканов
# =============================================================================
_DB_PATH      = os.path.join("data", "doc_scanner", "scans_db.json")
_ORIGINALS_DIR = os.path.join("data", "doc_scanner", "originals")
_DB_LOCK = threading.Lock()


def _ensure_dir():
    os.makedirs(os.path.dirname(_DB_PATH), exist_ok=True)
    os.makedirs(_ORIGINALS_DIR, exist_ok=True)


def _fname(doc: Dict) -> str:
    """Совместимость: старые документы используют 'file_name', новые — 'filename'."""
    return doc.get("filename") or doc.get("file_name") or "неизвестный файл"




def _summary_path(original_path: str) -> str:
    """Путь к txt-файлу пересказа рядом с оригиналом."""
    if not original_path:
        return ""
    base = os.path.splitext(original_path)[0]
    return base + "_пересказ.txt"

def save_summary_file(original_path: str, summary: str):
    """Сохраняет пересказ в txt рядом с оригиналом."""
    path = _summary_path(original_path)
    if path:
        with open(path, "w", encoding="utf-8") as f:
            f.write(summary)

def load_summary_file(original_path: str) -> str:
    """Читает пересказ из txt, возвращает пустую строку если нет."""
    path = _summary_path(original_path)
    if path and os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    return ""

def load_db() -> Dict:
    _ensure_dir()
    if os.path.exists(_DB_PATH):
        try:
            with open(_DB_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {"documents": [], "stats": {"total": 0, "last_scan": None}}


def save_db(db: Dict):
    _ensure_dir()
    with _DB_LOCK:
        with open(_DB_PATH, "w", encoding="utf-8") as f:
            json.dump(db, f, ensure_ascii=False, indent=2)


# =============================================================================
# OCR — EasyOCR с fallback на Tesseract
# =============================================================================
_easyocr_reader      = None
_easyocr_lock        = threading.Lock()
_EASYOCR_AVAILABLE   = False
_TESSERACT_AVAILABLE = False


def _init_ocr():
    global _easyocr_reader, _EASYOCR_AVAILABLE, _TESSERACT_AVAILABLE
    # EasyOCR
    try:
        import easyocr
        with _easyocr_lock:
            if _easyocr_reader is None:
                _easyocr_reader = easyocr.Reader(
                    ["ru", "en"],
                    gpu=True,
                    verbose=False,
                )
        _EASYOCR_AVAILABLE = True
        print("[OCR] EasyOCR инициализирован")
    except ImportError:
        print("[OCR] EasyOCR не установлен: pip install easyocr")
    except Exception as e:
        print(f"[OCR] EasyOCR ошибка инициализации: {e}")
        try:
            import streamlit as _st
            _st.session_state["ocr_init_error"] = str(e)
        except Exception:
            pass

    # Tesseract fallback
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        _TESSERACT_AVAILABLE = True
        print("[OCR] Tesseract доступен как fallback")
    except Exception:
        pass


def _ocr_image(image) -> str:
    """
    Распознаёт текст из PIL Image.
    Стратегия: EasyOCR → Tesseract → пустая строка.
    """
    # EasyOCR
    if _EASYOCR_AVAILABLE and _easyocr_reader is not None:
        try:
            import numpy as np
            img_array = np.array(image.convert("RGB"))
            result = _easyocr_reader.readtext(img_array, detail=0, paragraph=True)
            return "\n".join(result)
        except Exception as e:
            print(f"[OCR] EasyOCR ошибка: {e}")

    # Tesseract fallback
    if _TESSERACT_AVAILABLE:
        try:
            import pytesseract
            return pytesseract.image_to_string(image, lang="rus+eng")
        except Exception as e:
            print(f"[OCR] Tesseract ошибка: {e}")

    return ""


# =============================================================================
# Извлечение текста по типу файла
# =============================================================================
def _extract_pdf(file_bytes: bytes, filename: str,
                 max_pages: int = 0) -> List[Dict]:
    """
    PDF → список страниц [{page, text, method}].
    max_pages=0 — читать все страницы (полный режим).
    max_pages=N — читать только первые N страниц (режим заголовка).
    """
    pages = []
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for i, page in enumerate(doc, 1):
            if max_pages and i > max_pages:
                break  # останавливаемся — остальные страницы не читаем

            # Попытка прямого извлечения текста
            text = page.get_text("text").strip()
            method = "direct"

            # Если текста мало — скан → OCR
            if len(text) < 50:
                try:
                    from PIL import Image
                    pix = page.get_pixmap(dpi=200)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    text = _ocr_image(img)
                    method = "ocr"
                except Exception as e:
                    text = f"[Ошибка OCR стр.{i}: {e}]"
                    method = "error"

            pages.append({
                "page": i,
                "text": text,
                "method": method,
                "word_count": len(text.split()),
            })
        doc.close()
    except ImportError:
        pages.append({"page": 1, "text": "[PyMuPDF не установлен: pip install pymupdf]",
                       "method": "error", "word_count": 0})
    except Exception as e:
        pages.append({"page": 1, "text": f"[Ошибка PDF: {e}]",
                       "method": "error", "word_count": 0})
    return pages


def _extract_docx(file_bytes: bytes, max_pages: int = 0) -> List[Dict]:
    """DOCX → список страниц (по 100 строк каждая)."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        # Таблицы
        for table in doc.tables:
            for row in table.rows:
                full_text += "\n" + " | ".join(c.text.strip() for c in row.cells)
        return _split_into_pages(full_text, method="direct")
    except ImportError:
        return [{"page": 1, "text": "[python-docx не установлен: pip install python-docx]",
                 "method": "error", "word_count": 0}]
    except Exception as e:
        return [{"page": 1, "text": f"[Ошибка DOCX: {e}]",
                 "method": "error", "word_count": 0}]


def _extract_doc(file_bytes: bytes) -> List[Dict]:
    """DOC (старый формат) → текст через mammoth."""
    try:
        import mammoth
        result = mammoth.extract_raw_text(io.BytesIO(file_bytes))
        return _split_into_pages(result.value, method="direct")
    except ImportError:
        return [{"page": 1, "text": "[mammoth не установлен: pip install mammoth]",
                 "method": "error", "word_count": 0}]
    except Exception as e:
        return [{"page": 1, "text": f"[Ошибка DOC: {e}]",
                 "method": "error", "word_count": 0}]


def _extract_xlsx(file_bytes: bytes) -> List[Dict]:
    """XLSX → текст из всех листов."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        lines = []
        for sheet in wb.worksheets:
            lines.append(f"=== Лист: {sheet.title} ===")
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(c) for c in row if c is not None)
                if row_text.strip():
                    lines.append(row_text)
        return _split_into_pages("\n".join(lines), method="direct")
    except ImportError:
        return [{"page": 1, "text": "[openpyxl не установлен: pip install openpyxl]",
                 "method": "error", "word_count": 0}]
    except Exception as e:
        return [{"page": 1, "text": f"[Ошибка XLSX: {e}]",
                 "method": "error", "word_count": 0}]


def _extract_image(file_bytes: bytes) -> List[Dict]:
    """JPG/PNG → OCR."""
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(file_bytes))
        text = _ocr_image(img)
        return [{"page": 1, "text": text, "method": "ocr",
                 "word_count": len(text.split())}]
    except Exception as e:
        return [{"page": 1, "text": f"[Ошибка изображения: {e}]",
                 "method": "error", "word_count": 0}]


def _extract_txt(file_bytes: bytes) -> List[Dict]:
    """TXT → список страниц (по 80 строк каждая)."""
    try:
        # Пробуем UTF-8, затем cp1251 (Windows-кириллица)
        for enc in ("utf-8", "cp1251", "latin-1"):
            try:
                text = file_bytes.decode(enc)
                break
            except UnicodeDecodeError:
                continue
        else:
            text = file_bytes.decode("utf-8", errors="replace")
        return _split_into_pages(text, method="direct")
    except Exception as e:
        return [{"page": 1, "text": f"[Ошибка TXT: {e}]",
                 "method": "error", "word_count": 0}]


def _split_into_pages(text: str, method: str = "direct", lines_per_page: int = 80) -> List[Dict]:
    """Разбивает длинный текст на условные страницы."""
    lines = text.splitlines()
    pages = []
    for i in range(0, max(1, len(lines)), lines_per_page):
        chunk = "\n".join(lines[i:i + lines_per_page]).strip()
        pages.append({
            "page": len(pages) + 1,
            "text": chunk,
            "method": method,
            "word_count": len(chunk.split()),
        })
    return pages if pages else [{"page": 1, "text": text, "method": method,
                                  "word_count": len(text.split())}]


def extract_text(file_bytes: bytes, filename: str,
                 max_pages: int = 0) -> List[Dict]:
    """
    Универсальный экстрактор — выбирает метод по расширению файла.
    max_pages=0 — читать всё (полный режим).
    max_pages=N — читать только первые N страниц (быстрый режим для заголовков).
    """
    ext = os.path.splitext(filename.lower())[1]
    if ext == ".pdf":
        return _extract_pdf(file_bytes, filename, max_pages=max_pages)
    elif ext == ".docx":
        return _extract_docx(file_bytes, max_pages=max_pages)
    elif ext in (".doc",):
        return _extract_doc(file_bytes)
    elif ext in (".xlsx", ".xls"):
        return _extract_xlsx(file_bytes)
    elif ext in (".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp"):
        return _extract_image(file_bytes)
    elif ext == ".txt":
        return _extract_txt(file_bytes)
    else:
        return [{"page": 1, "text": f"[Формат {ext} не поддерживается]",
                 "method": "error", "word_count": 0}]


# =============================================================================
# База документов
# =============================================================================
def add_to_db(db: Dict, filename: str, pages: List[Dict],
              original_path: str = "") -> str:
    doc_id = hashlib.md5(
        f"{filename}_{datetime.now().isoformat()}".encode()
    ).hexdigest()[:12]

    full_text = "\n".join(p["text"] for p in pages)
    ocr_pages = sum(1 for p in pages if p.get("method") == "ocr")

    doc = {
        "id": doc_id,
        "filename": filename,
        "pages": pages,
        "full_text": full_text,
        "word_count": sum(p["word_count"] for p in pages),
        "page_count": len(pages),
        "ocr_pages": ocr_pages,
        "processed_at": datetime.now().isoformat(),
        "original_path": original_path,
    }
    db["documents"].append(doc)
    db["stats"]["total"] = len(db["documents"])
    db["stats"]["last_scan"] = doc["processed_at"]
    return doc_id


# =============================================================================
# Поиск с синонимами
# =============================================================================
def search_documents(db: Dict, query: str,
                     include_summary: bool = False) -> List[Dict]:
    """Поиск по документам. include_summary — искать и в тексте пересказа."""
    if not query.strip():
        return []

    search_terms = [query.lower()]
    try:
        from core.query_expander import QueryExpander
        variants = QueryExpander().expand(query)
        search_terms += [
            v.lower() for v in variants
            if v != query and not v.endswith("тарифное регулирование")
        ]
    except Exception:
        pass

    results = []
    for doc in db.get("documents", []):
        matches = []

        # Поиск по страницам
        for page in doc.get("pages", []):
            page_text = page.get("text", "").lower()
            matched_terms = [t for t in search_terms if t in page_text]
            if matched_terms:
                term = matched_terms[0]
                idx = page_text.find(term)
                start = max(0, idx - 100)
                end = min(len(page_text), idx + 200)
                context = "..." + page_text[start:end].strip() + "..."
                for t in matched_terms:
                    context = re.sub(
                        re.escape(t), f"**{t.upper()}**", context, flags=re.IGNORECASE
                    )
                matches.append({
                    "page": page["page"],
                    "source": "текст",
                    "context": context,
                    "matched_terms": matched_terms,
                })

        # Поиск по пересказу (опционально)
        if include_summary:
            _sum_text = load_summary_file(doc.get("original_path", ""))
            if _sum_text:
                _sum_lower = _sum_text.lower()
                matched_terms = [t for t in search_terms if t in _sum_lower]
                if matched_terms:
                    term = matched_terms[0]
                    idx = _sum_lower.find(term)
                    start = max(0, idx - 100)
                    end = min(len(_sum_lower), idx + 200)
                    context = "..." + _sum_text[start:end].strip() + "..."
                    for t in matched_terms:
                        context = re.sub(
                            re.escape(t), f"**{t.upper()}**", context,
                            flags=re.IGNORECASE
                        )
                    matches.append({
                        "page": None,
                        "source": "пересказ",
                        "context": context,
                        "matched_terms": matched_terms,
                    })

        if matches:
            results.append({
                "doc": doc,
                "matches": matches,
                "total_matches": len(matches),
            })

    results.sort(key=lambda x: x["total_matches"], reverse=True)
    return results


# =============================================================================
# Пересказ через LM Studio — Map-Reduce для больших документов
# =============================================================================

# Порог (символов): документы длиннее этого значения обрабатываются чанками
_LARGE_DOC_THRESHOLD = 12_000
# Размер одного чанка (символов) — ~1500 токенов, комфортно для 8B модели
_CHUNK_SIZE          = 6_000
# Перекрытие между чанками — модель не теряет контекст на стыке частей
_CHUNK_OVERLAP       = 300


def _safe_print(msg: str):
    """print(), безопасный для Windows-консоли (cp1252 и др.)."""    
    try:
        print(msg)
    except UnicodeEncodeError:
        print(msg.encode("ascii", errors="replace").decode("ascii"))


def _strip_thinking(text: str) -> str:
    """
    Убирает <think>...</think> блоки Qwen3/DeepSeek.
    Модель генерирует их перед ответом — они не нужны пользователю.
    """
    import re
    cleaned = re.sub(r'<think>.*?</think>', '', text, flags=re.DOTALL)
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
    return cleaned.strip()


def _load_lm_config() -> tuple:
    """Возвращает (lm_url, model) из advisor_config.json."""
    config_path = os.path.join("config", "advisor_config.json")
    config = {}
    if os.path.exists(config_path):
        try:
            with open(config_path, "r", encoding="utf-8") as f:
                config = json.load(f)
        except Exception:
            pass
    lm_url = config.get("lm_studio_url", "http://127.0.0.1:1234/v1")
    model  = config.get("default_model", "qwen/qwen3.5-9b")
    return lm_url, model


def _split_text_chunks(text: str, chunk_size: int = _CHUNK_SIZE,
                       overlap: int = _CHUNK_OVERLAP) -> List[str]:
    """
    Разбивает текст на чанки по границам абзацев/предложений.
    Перекрытие `overlap` сохраняет контекст между соседними частями.

    Ключевое правило: разделитель принимается только если он находится
    в ПОСЛЕДНЕЙ ЧЕТВЕРТИ окна [start..end] — иначе чанк получался бы
    крошечным (когда \n\n стоит близко к началу диапазона rfind).
    """
    if len(text) <= chunk_size:
        return [text]

    chunks, start = [], 0
    while start < len(text):
        end = start + chunk_size
        if end >= len(text):
            tail = text[start:].strip()
            if tail:
                chunks.append(tail)
            break

        # Ищем точку разрыва в последних 25% окна, чтобы чанк
        # был не короче 75% chunk_size
        min_pos = start + chunk_size * 3 // 4
        split_end = end
        for sep in ('\n\n', '. ', ' '):
            pos = text.rfind(sep, min_pos, end)
            if pos != -1:
                split_end = pos + len(sep)
                break

        chunk = text[start:split_end].strip()
        if chunk:
            chunks.append(chunk)

        # Следующий чанк начинается с перекрытием, но не раньше split_end//2
        next_start = split_end - overlap
        start = max(start + chunk_size // 2, next_start)

    return chunks


def _lm_call(client, model: str, system: str, user: str, max_tokens: int) -> str:
    """Один вызов LM Studio; возвращает текст ответа или строку с ошибкой."""
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system},
                {"role": "user",   "content": user},
            ],
            max_tokens=max_tokens,
            temperature=0.3,
        )
        raw = (resp.choices[0].message.content or "").strip()
        return _strip_thinking(raw)
    except Exception as e:
        return f"[Ошибка LM: {e}]"


def summarize_document(text: str, length: str = "средний", model: str = None, _progress_cb=None) -> str:
    """
    Пересказывает текст через LM Studio.

    Для коротких документов (< _LARGE_DOC_THRESHOLD символов) — один запрос.
    Для длинных — Map-Reduce:
      MAP    : каждый чанк (~6 000 симв.) суммаризируется отдельно
      REDUCE : мини-резюме объединяются в итоговый пересказ

    """
    # length может быть строкой-ключом ИЛИ числом слов (int/str с цифрой)
    # Строковые ключи → фиксированные пресеты
    _PRESETS = {
        "1 страница":  (500,  "примерно на 1 страницу (~250 слов). Только самое главное"),
        "2 страницы":  (1000, "примерно на 2 страницы (~500 слов). Ключевые факты и детали"),
        "5 страниц":   (2500, "примерно на 5 страниц (~1250 слов). Подробно, все важные пункты"),
        "10 страниц":  (5000, "примерно на 10 страниц (~2500 слов). Максимально подробно"),
        # Обратная совместимость со старыми ключами
        "краткий":     (400,  "кратко, в 3-5 предложениях, только главное"),
        "средний":     (1000, "в 1-2 абзацах с ключевыми деталями (~500 слов)"),
        "подробный":   (2500, "подробно, с перечислением всех важных пунктов (~1000 слов)"),
    }
    # Числовой режим: length == "NNN слов" или int
    _word_target = None
    try:
        _word_target = int(str(length).replace("слов", "").replace("слова", "").strip())
    except (ValueError, TypeError):
        pass

    if _word_target:
        # ~1.5 токена на слово для русского текста
        max_tokens  = int(_word_target * 1.7)
        instruction = f"ровно примерно {_word_target} слов. Следи за объёмом"
    else:
        max_tokens, instruction = _PRESETS.get(length, _PRESETS["2 страницы"])


    try:
        from openai import OpenAI
        lm_url, default_model = _load_lm_config()
        model  = model or default_model
        client = OpenAI(base_url=lm_url, api_key="lm-studio", timeout=180.0)

        system_msg = "Ты помощник для анализа документов. Отвечаешь кратко и по делу на русском языке."

        # ── Короткий документ: один запрос ───────────────────────────────
        if len(text) <= _LARGE_DOC_THRESHOLD:
            prompt = (
                f"Перескажи содержание документа {instruction}. "
                f"Отвечай на русском языке, без лишних вступлений.\n\n"
                f"ДОКУМЕНТ:\n{text}"
            )
            result = _lm_call(client, model, system_msg, prompt, max_tokens)
            return result

        # ── Большой документ: Map-Reduce ─────────────────────────────────
        chunks   = _split_text_chunks(text)
        total    = len(chunks)
        t_start  = time.perf_counter()
        _safe_print(f"[MAP-REDUCE] {len(text)} chars -> {total} chunks, model: {model}")

        # MAP: резюме каждого чанка
        mini_summaries = []
        chunk_times    = []   # храним время каждого чанка для ETA

        for i, chunk in enumerate(chunks, 1):
            t_chunk = time.perf_counter()

            # Статус перед запросом
            elapsed  = t_chunk - t_start
            elapsed_fmt = f"{int(elapsed // 60)}м {int(elapsed % 60)}с"
            if chunk_times:
                avg_sec  = sum(chunk_times) / len(chunk_times)
                remaining = avg_sec * (total - i + 1)
                eta_fmt  = f"{int(remaining // 60)}м {int(remaining % 60)}с"
                eta_str  = f"ещё ~{eta_fmt}"
            else:
                eta_str  = "считаю время..."

            if _progress_cb:
                _pct = (i - 1) / (total + 1)
                _progress_cb(_pct, f"MAP — часть {i} / {total} | прошло: {elapsed_fmt} | {eta_str}")

            map_prompt = (
                f"Это часть {i} из {total} одного документа.\n"
                f"Сделай краткое резюме этой части (3-5 пунктов), "
                f"сохрани все факты, цифры, даты, названия.\n"
                f"Отвечай на русском, без вступлений.\n\n"
                f"ЧАСТЬ ДОКУМЕНТА:\n{chunk}"
            )
            mini = _lm_call(client, model, system_msg, map_prompt, 400)

            chunk_done = time.perf_counter()
            chunk_times.append(chunk_done - t_chunk)

            if mini.startswith("[Ошибка"):
                _safe_print(f"[MAP] chunk {i} error: {mini}")
                mini = f"[Часть {i}: данные недоступны]"

            mini_summaries.append(f"=== Часть {i}/{total} ===\n{mini}")
            _safe_print(f"[MAP] chunk {i}/{total} done ({len(mini)} chars, {chunk_times[-1]:.1f}s)")

        # REDUCE: финальный синтез
        elapsed_map = time.perf_counter() - t_start
        elapsed_fmt = f"{int(elapsed_map // 60)}м {int(elapsed_map % 60)}с"
        _safe_print(f"[MAP-REDUCE] REDUCE start, map took {elapsed_fmt}")
        if _progress_cb:
            _progress_cb(total / (total + 1), f"REDUCE — финальный синтез {total} частей | MAP: {elapsed_fmt}")

        combined = "\n\n".join(mini_summaries)
        reduce_prompt = (
            f"Ниже — резюме отдельных частей одного документа. "
            f"Создай единый итоговый пересказ {instruction}.\n"
            f"Объедини все части, устрани дублирование, сохрани все ключевые факты, "
            f"цифры, даты, названия документов.\n"
            f"Отвечай на русском, без вступлений и мета-комментариев.\n\n"
            f"РЕЗЮМЕ ЧАСТЕЙ:\n{combined}"
        )
        result = _lm_call(client, model, system_msg, reduce_prompt, max_tokens)

        total_fmt = f"{int((time.perf_counter()-t_start)//60)}м {int((time.perf_counter()-t_start)%60)}с"
        _safe_print(f"[MAP-REDUCE] done in {total_fmt}. result: {len(result)} chars")
        return result

    except Exception as e:
        return f"[Ошибка пересказа: {e}]"


# =============================================================================
# Экспорт результатов
# =============================================================================
def export_txt(doc: Dict) -> bytes:
    """Экспорт в TXT."""
    _pc = doc.get("page_count") or len(doc.get("pages", []))
    lines = [
        f"Документ: {_fname(doc)}",
        f"Обработан: {(doc.get('processed_at') or '')[:16]}",
        f"Страниц: {_pc} | Слов: {doc.get('word_count', 0)}",
        "=" * 60,
        "",
    ]
    for page in doc.get("pages", []):
        lines.append(f"--- Страница {page.get('page',1)} ({page.get('method','')}) ---")
        lines.append(page.get("text", ""))
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def export_docx(doc: Dict, summary: str = None) -> io.BytesIO:
    """Экспорт в DOCX."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        d = Document()
        _pc = doc.get("page_count") or len(doc.get("pages", []))

        h = d.add_heading(_fname(doc), level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        d.add_paragraph(f"Обработан: {(doc.get('processed_at') or '')[:16]}")
        d.add_paragraph(f"Страниц: {_pc} | Слов: {doc.get('word_count', 0)}")
        d.add_paragraph(f"OCR-страниц: {doc.get('ocr_pages', 0)}")

        if summary:
            d.add_heading("Краткое содержание", level=2)
            d.add_paragraph(summary)

        d.add_heading("Текст документа", level=2)
        for page in doc.get("pages", []):
            d.add_heading(f"Страница {page.get('page',1)}", level=3)
            d.add_paragraph(page.get("text", ""))

        buf = io.BytesIO()
        d.save(buf)
        buf.seek(0)
        return buf
    except ImportError:
        raise RuntimeError("python-docx не установлен: pip install python-docx")


# =============================================================================
# UI — главная страница сканера
# =============================================================================
def show_doc_scanner():
    st.header("Сканер документов")
    st.caption("PDF · DOCX · DOC · XLSX · JPG · PNG — распознавание, пересказ, поиск")

    # Инициализация OCR один раз
    if "ocr_initialized" not in st.session_state:
        with st.spinner("Инициализация OCR..."):
            _init_ocr()
        st.session_state["ocr_initialized"] = True

    # Статус OCR
    if _EASYOCR_AVAILABLE:
        st.success("Распознавание текста активно")
    elif _TESSERACT_AVAILABLE:
        st.warning("Основной модуль OCR недоступен, используется резервный (хуже качество).")
    else:
        _ocr_err = st.session_state.get("ocr_init_error")
        if _ocr_err:
            st.error(f"Ошибка инициализации OCR:\n```\n{_ocr_err}\n```")
        else:
            st.error("OCR не найден. Установите: `pip install paddlepaddle paddleocr pillow`")

    # Загрузка базы
    if "scanner_db" not in st.session_state:
        st.session_state["scanner_db"] = load_db()

    db = st.session_state["scanner_db"]

    # ── Вкладки ──────────────────────────────────────────────────────────
    tab_scan, tab_search, tab_docs = st.tabs([
        "Загрузка и распознавание",
        "Поиск по содержимому",
        "База документов",
    ])

    # =========================================================================
    # Вкладка 1 — Загрузка и распознавание
    # =========================================================================
    with tab_scan:
        st.subheader("Загрузка файлов")

        SUPPORTED = ["pdf", "docx", "doc", "xlsx", "xls", "jpg", "jpeg", "png", "bmp", "tiff", "txt"]
        uploaded = st.file_uploader(
            "Перетащите файлы или выберите из папки",
            type=SUPPORTED,
            accept_multiple_files=True,
            help="Поддерживаются: PDF (текстовые и сканы), DOCX, DOC, XLSX, JPG, PNG, TXT",
        )

        if uploaded:
            st.info(f"📎 Загружено файлов: **{len(uploaded)}**")

            # Кнопка запуска
            if st.button("▶ Запустить распознавание", type="primary", key="scan_btn"):
                st.warning("⚠️ Идёт распознавание — не переключайте раздел и не закрывайте вкладку")
                progress = st.progress(0)
                status   = st.empty()
                new_ids  = []

                for i, f in enumerate(uploaded):
                    status.text(f"⏳ Обрабатываю {f.name} ({i+1}/{len(uploaded)})...")
                    progress.progress((i) / len(uploaded))

                    file_bytes = f.read()
                    t0 = time.perf_counter()
                    pages = extract_text(file_bytes, f.name)
                    elapsed = time.perf_counter() - t0

                    # Сохраняем оригинал на диск
                    _orig_path = os.path.join(_ORIGINALS_DIR, f.name)
                    try:
                        with open(_orig_path, "wb") as _of:
                            _of.write(file_bytes)
                    except Exception:
                        _orig_path = ""

                    doc_id = add_to_db(db, f.name, pages, original_path=_orig_path)
                    new_ids.append(doc_id)

                    ocr_count = sum(1 for p in pages if p.get("method") == "ocr")
                    words     = sum(p["word_count"] for p in pages)
                    st.success(
                        f"✅ **{f.name}** — {len(pages)} стр., {words} слов, "
                        f"{ocr_count} OCR-стр., {elapsed:.1f} сек"
                    )

                progress.progress(1.0)
                status.text("✅ Все файлы обработаны!")
                save_db(db)
                st.session_state["scanner_db"] = db
                st.session_state["last_scanned_ids"] = new_ids

        # ── Результаты последней обработки ──────────────────────────────
        _raw_ids = st.session_state.get("last_scanned_ids", [])
        last_ids = list(dict.fromkeys(_raw_ids))
        if last_ids:
            _seen = set()
            last_docs = [
                d for d in db["documents"]
                if d["id"] in last_ids and d["id"] not in _seen
                and not _seen.add(d["id"])
            ]
            # Восстанавливаем пересказы из файлов в session_state
            for _ld in last_docs:
                _sk = f"summary_{_ld['id']}"
                if _sk not in st.session_state:
                    _loaded = load_summary_file(_ld.get("original_path", ""))
                    if _loaded:
                        st.session_state[_sk] = _loaded


            # ── Список документов ─────────────────────────────────────────
            st.divider()
            st.subheader("📄 Результаты распознавания")

            # ── Выбор документа — selectbox со встроенным поиском ───────
            def _short(name, n=45):
                return name if len(name) <= n else name[:n] + "…"

            _doc_ids   = [d["id"] for d in last_docs]
            _doc_labels = {
                d["id"]: _short(_fname(d)) + f"  ·  {d.get('page_count') or len(d.get('pages',[]))} стр."
                for d in last_docs
            }

            # Сохраняем выбранный документ в session_state
            if "scan_sel_id" not in st.session_state or                st.session_state.scan_sel_id not in _doc_ids:
                st.session_state.scan_sel_id = _doc_ids[0]

            _n_docs = len(_doc_ids)
            sel_id = st.selectbox(
                f"Выберите документ  ({_n_docs} {'документ' if _n_docs == 1 else 'документа' if 2 <= _n_docs <= 4 else 'документов'})",
                options=_doc_ids,
                format_func=lambda x: _doc_labels[x],
                key="scan_sel_id",
            )

            doc = next((d for d in last_docs if d["id"] == sel_id), last_docs[0])

            # ── Шапка выбранного документа ────────────────────────────────
            _pc2 = doc.get("page_count") or len(doc.get("pages", []))
            _wc2 = doc.get("word_count", 0)
            _has_summary = bool(st.session_state.get(f"summary_{doc['id']}"))
            st.markdown(
                f"<div style='background:#ffffff;border:1px solid #dce3ec;"
                f"border-radius:6px;padding:12px 16px;margin-bottom:12px;"
                f"display:flex;align-items:center;justify-content:space-between;'>"
                f"<div>"
                f"<div style='font-size:1rem;font-weight:700;color:#1a2a3a;"
                f"word-break:break-all;'>{_fname(doc)}</div>"
                f"<div style='font-size:0.78rem;color:#5a6a7a;margin-top:3px;'>"
                f"{_pc2} стр. · {_wc2:,} слов"
                f"{'  ·  ✅ пересказ готов' if _has_summary else ''}</div>"
                f"</div></div>",
                unsafe_allow_html=True,
            )

            # ── Навигация по страницам ────────────────────────────────────
            pages_list = doc.get("pages", [])
            page_key   = f"page_idx_{doc['id']}"
            if page_key not in st.session_state:
                st.session_state[page_key] = 0

            cur_idx  = st.session_state[page_key]
            cur_idx  = max(0, min(cur_idx, len(pages_list) - 1))
            cur_page = pages_list[cur_idx] if pages_list else {}

            nav1, nav2, nav3 = st.columns([1, 4, 1])
            with nav1:
                if st.button("◀ Пред.", key=f"prev_btn_{doc['id']}",
                             disabled=(cur_idx == 0)):
                    st.session_state[page_key] = cur_idx - 1
                    st.rerun()
            with nav2:
                st.markdown(
                    f"<div style='text-align:center;padding-top:6px;'>"
                    f"Страница <b>{cur_idx+1}</b> из <b>{len(pages_list)}</b> "
                    f"<span style='color:grey;font-size:0.85em;'>"
                    f"({cur_page.get('method','')} · "
                    f"{cur_page.get('word_count') or len(cur_page.get('text','').split())} слов)"
                    f"</span></div>",
                    unsafe_allow_html=True,
                )
            with nav3:
                if st.button("След. ▶", key=f"next_btn_{doc['id']}",
                             disabled=(cur_idx >= len(pages_list) - 1)):
                    st.session_state[page_key] = cur_idx + 1
                    st.rerun()

            st.markdown(
                f"<div style='font-size:0.87em;line-height:1.55;"
                f"background:#f8f9fa;border:1px solid #e0e0e0;"
                f"border-radius:6px;padding:12px 14px;"
                f"max-height:320px;overflow-y:auto;"
                f"white-space:pre-wrap;word-break:break-word;'>"
                f"{cur_page.get('text','').replace('<','&lt;').replace('>','&gt;')}"
                f"</div>",
                unsafe_allow_html=True,
            )

            st.divider()

            # ── Пересказ одного документа ─────────────────────────────────
            summary_key  = f"summary_{doc['id']}"
            sum_mode_key = f"sum_mode_{doc['id']}"
            sum_words_key= f"sum_words_{doc['id']}"

            _doc_short_name = _fname(doc)
            if len(_doc_short_name) > 50:
                _doc_short_name = _doc_short_name[:50] + "…"

            st.markdown(f"**Пересказ: {_doc_short_name}**")
            sum_mode = st.selectbox(
                "Объём пересказа",
                ["1 страница", "2 страницы", "5 страниц", "10 страниц",
                 "Точное кол-во слов"],
                index=1,
                key=sum_mode_key,
            )
            if sum_mode == "Точное кол-во слов":
                sum_words = st.number_input(
                    "Количество слов",
                    min_value=50, max_value=5000, value=300, step=50,
                    key=sum_words_key,
                )
                summary_length = str(int(sum_words))
            else:
                summary_length = sum_mode

            _already_has = bool(st.session_state.get(summary_key))

            _bL, _bC, _bR = st.columns([2, 3, 2])
            with _bC:
                _do_summary = st.button(
                    f"▶ Пересказать «{_doc_short_name}»",
                    key=f"sum_btn_{doc['id']}", type="primary",
                    use_container_width=True,
                )

            if _already_has:
                st.caption(
                    "Пересказ уже создан. Повторное нажатие запустит процесс заново."
                )

            if _do_summary:
                _gft  = doc.get("full_text", "")
                _gnc  = max(1, len(_gft) // _CHUNK_SIZE + 1)
                _gisl = len(_gft) > _LARGE_DOC_THRESHOLD

                if _gisl:
                    st.caption(f"📄 Map-Reduce: ~{_gnc} частей · {len(_gft):,} символов")

                _gprog = st.progress(0.0)
                _gcap  = st.empty()
                _gcap.caption("⏳ Подготовка...")

                def _gcb(pct, msg, _p=_gprog, _c=_gcap):
                    _pct = min(float(pct), 1.0)
                    _p.progress(_pct)
                    _c.caption(f"⏳ {msg}  ·  {int(_pct * 100)}%")

                st.warning("⚠️ Идёт генерация пересказа — не переключайте раздел и не закрывайте вкладку")
                with st.spinner("🤖 Генерирую пересказ..."):
                    _gresult = summarize_document(_gft, summary_length, _progress_cb=_gcb)

                _gprog.empty()
                _gcap.empty()
                if not _gresult or not _gresult.strip():
                    _gresult = "⚠️ Модель вернула пустой ответ. Попробуйте ещё раз."
                st.session_state[summary_key] = _gresult

                _orig_p = doc.get("original_path", "")
                if _orig_p:
                    save_summary_file(_orig_p, _gresult)
                # Нет st.rerun() — результат показывается сразу ниже

            # ── Результат пересказа ───────────────────────────────────────
            _summary_text = st.session_state.get(summary_key)
            if _summary_text:
                import streamlit.components.v1 as _stc
                _wc = len(_summary_text.split())

                # Заголовок с кол-вом слов
                st.markdown(
                    "<div style='display:flex;align-items:center;"
                    "justify-content:space-between;margin-bottom:8px;margin-top:4px;'>"
                    "<span style='font-weight:600;font-size:0.95em;'>🤖 Пересказ</span>"
                    f"<span style='color:#888;font-size:0.82em;'>{_wc} слов</span>"
                    "</div>",
                    unsafe_allow_html=True,
                )

                # Рендерим как markdown — LLM возвращает структуру с заголовками и списками
                st.markdown(_summary_text)

                # Кнопка копирования через iframe (текст в скрытом span)
                _safe = (
                    _summary_text
                    .replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                    .replace('"', "&quot;")
                )
                _stc.html(
                    "<div style='font-family:sans-serif'>"
                    f"<span id='t' style='display:none'>{_safe}</span>"
                    "<button onclick=\"navigator.clipboard.writeText("
                    "document.getElementById('t').textContent).then(()=>{"
                    "this.textContent='✅ Скопировано';"
                    "setTimeout(()=>this.textContent='📋 Скопировать пересказ',2000)})\" "
                    "style='font-size:13px;padding:5px 16px;cursor:pointer;"
                    "border:1px solid #d0d8e4;border-radius:5px;"
                    "background:#fff;color:#333;margin-top:2px'>"
                    "📋 Скопировать пересказ</button></div>",
                    height=44,
                )

            st.divider()

            # ── Экспорт ───────────────────────────────────────────────────
            st.markdown("**Сохранить результат:**")
            exp_col1, exp_col2 = st.columns(2)
            with exp_col1:
                st.download_button(
                    label="Скачать TXT",
                    data=export_txt(doc),
                    file_name=f"{os.path.splitext(_fname(doc))[0]}_распознан.txt",
                    mime="text/plain",
                    key=f"dl_txt_{doc['id']}",
                    use_container_width=True,
                )
            with exp_col2:
                try:
                    docx_buf = export_docx(doc, summary=st.session_state.get(summary_key))
                    st.download_button(
                        label="Скачать DOCX",
                        data=docx_buf,
                        file_name=f"{os.path.splitext(_fname(doc))[0]}_распознан.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_docx_{doc['id']}",
                        use_container_width=True,
                    )
                except Exception as e:
                    st.warning(f"DOCX недоступен: {e}")

            # ── Пересказ всех загруженных файлов ─────────────────────────
            if len(last_docs) > 1:
                st.divider()
                st.markdown("**Пересказать все загруженные файлы**")

                _batch_mode_key  = "batch_sum_mode"
                _batch_words_key = "batch_sum_words"
                _ba, _bb = st.columns([3, 2])
                with _ba:
                    _batch_mode = st.selectbox(
                        "Объём пересказа для каждого файла",
                        ["1 страница", "2 страницы", "5 страниц", "10 страниц",
                         "Точное кол-во слов"],
                        index=1,
                        key=_batch_mode_key,
                    )
                with _bb:
                    if _batch_mode == "Точное кол-во слов":
                        _batch_words = st.number_input(
                            "Количество слов",
                            min_value=50, max_value=5000, value=300, step=50,
                            key=_batch_words_key,
                        )
                        _batch_length = str(int(_batch_words))
                    else:
                        _batch_length = _batch_mode
                        _batch_words  = None

                # Счётчик уже готовых
                _batch_done = sum(
                    1 for _bd in last_docs
                    if st.session_state.get(f"summary_{_bd['id']}")
                )
                if _batch_done:
                    st.caption(f"Уже готово: {_batch_done} из {len(last_docs)}")

                _bpL, _bpC, _bpR = st.columns([2, 3, 2])
                with _bpC:
                    _do_batch = st.button(
                        f"▶ Пересказать все {len(last_docs)} файла(-ов)",
                        key="batch_sum_btn",
                        type="primary",
                        use_container_width=True,
                    )

                if _do_batch:
                    st.warning("⚠️ Идёт пакетный пересказ — не переключайте раздел и не закрывайте вкладку")
                    _batch_progress = st.progress(0.0)
                    _batch_status   = st.empty()
                    _batch_errors   = []

                    for _bi, _bd in enumerate(last_docs):
                        _bsk = f"summary_{_bd['id']}"
                        _bname = _fname(_bd)
                        _batch_status.caption(
                            f"⏳ Файл {_bi + 1}/{len(last_docs)}: {_bname}"
                        )
                        _batch_progress.progress(_bi / len(last_docs))

                        _bft = _bd.get("full_text", "")

                        # Вложенный прогресс для Map-Reduce
                        _inner_cap = st.empty()
                        def _bcb(pct, msg, _c=_inner_cap, _i=_bi, _t=len(last_docs)):
                            _c.caption(
                                f"  [{_i + 1}/{_t}] {msg}  ·  {int(min(pct, 1.0) * 100)}%"
                            )

                        _bres = summarize_document(
                            _bft, _batch_length, _progress_cb=_bcb
                        )
                        _inner_cap.empty()

                        if not _bres or not _bres.strip():
                            _bres = "⚠️ Модель вернула пустой ответ."
                            _batch_errors.append(_bname)

                        st.session_state[_bsk] = _bres

                        _borig = _bd.get("original_path", "")
                        if _borig:
                            save_summary_file(_borig, _bres)

                    _batch_progress.progress(1.0)
                    _batch_status.empty()

                    if _batch_errors:
                        st.warning(
                            "Не удалось получить пересказ для: "
                            + ", ".join(_batch_errors)
                        )
                    else:
                        st.success(
                            f"✅ Пересказы готовы для всех {len(last_docs)} файлов!"
                        )

    # =========================================================================
    # Вкладка 2 — Поиск
    # =========================================================================


    with tab_search:
        st.subheader("Поиск по содержимому")

        if not db.get("documents"):
            st.info("Сначала распознайте хотя бы один документ на вкладке «Загрузка».")
        else:
            _sq = st.text_input(
                "Поисковый запрос",
                placeholder="например: неподконтрольные расходы, ДМС, НДС",
                key="search_input",
            )

            so1, so2, so3 = st.columns(3)
            with so1:
                _match_type = st.radio(
                    "Тип совпадения",
                    ["Точное", "По словам", "Нечёткое"],
                    key="search_match_type",
                    help=(
                        "Точное — ищет фразу целиком. "
                        "По словам — все слова должны присутствовать. "
                        "Нечёткое — находит похожие слова (по началу)."
                    ),
                )
            with so2:
                _scope = st.radio(
                    "Где искать",
                    ["Текст документов", "Пересказы", "Везде"],
                    key="search_scope_adv",
                )
            with so3:
                _presence = st.radio(
                    "Наличие",
                    ["Содержит", "Не содержит"],
                    key="search_presence",
                    help="Найти документы где запрос присутствует или отсутствует",
                )

            _do_search = st.button("Найти", key="search_exec_btn", type="primary")
            st.divider()

            # ── Функция поиска с типами совпадения ───────────────────────────
            def _doc_matches(doc: dict, query: str, match_type: str, scope: str) -> list:
                """Возвращает список [(контекст, источник)] или [] если не найдено."""
                import re as _re
                q = query.strip()
                if not q:
                    return []

                def _search_in(text: str, src: str) -> list:
                    hits = []
                    if match_type == "Точное":
                        idx = 0
                        while True:
                            pos = text.lower().find(q.lower(), idx)
                            if pos < 0:
                                break
                            start   = max(0, pos - 80)
                            end     = min(len(text), pos + len(q) + 80)
                            snippet = (
                                ("..." if start > 0 else "")
                                + text[start:end]
                                + ("..." if end < len(text) else "")
                            )
                            hi = _re.sub(f"(?i)({_re.escape(q)})", r"**\1**", snippet)
                            hits.append((hi, src))
                            idx = pos + 1
                            if len(hits) >= 3:
                                break
                    elif match_type == "По словам":
                        words = q.lower().split()
                        tl    = text.lower()
                        if all(w in tl for w in words):
                            pos     = tl.find(words[0])
                            start   = max(0, pos - 60)
                            end     = min(len(text), pos + 120)
                            snippet = ("..." if start > 0 else "") + text[start:end] + "..."
                            pattern = "|".join(_re.escape(w) for w in words)
                            hi = _re.sub(f"(?i)({pattern})", r"**\1**", snippet)
                            hits.append((hi, src))
                    elif match_type == "Нечёткое":
                        words       = q.lower().split()
                        tl          = text.lower()
                        found_words = []
                        for w in words:
                            stem = w[:max(4, len(w) - 2)]
                            if stem in tl:
                                found_words.append(w)
                        if len(found_words) >= max(1, len(words) // 2):
                            pos     = tl.find(found_words[0][:4])
                            start   = max(0, pos - 60)
                            end     = min(len(text), pos + 150)
                            snippet = ("..." if start > 0 else "") + text[start:end] + "..."
                            pattern = "|".join(_re.escape(w[:4]) for w in found_words)
                            hi = _re.sub(f"(?i)({pattern})", r"**\1**", snippet)
                            hits.append((hi, src))
                    return hits

                all_hits = []

                # Поиск по тексту страниц
                if scope in ("Текст документов", "Везде"):
                    full_text = doc.get("full_text", "") or "\n".join(
                        p.get("text", "") for p in doc.get("pages", [])
                    )
                    all_hits += _search_in(full_text, "текст документа")

                # Поиск по пересказу
                if scope in ("Пересказы", "Везде"):
                    orig_path    = doc.get("original_path", "")
                    summary_text = load_summary_file(orig_path) if orig_path else ""
                    if summary_text:
                        all_hits += _search_in(summary_text, "пересказ")

                return all_hits

            # ── Выполнение поиска ─────────────────────────────────────────────
            if _do_search or st.session_state.get("search_last_query"):
                if _do_search:
                    st.session_state.search_last_query    = _sq
                    st.session_state.search_last_match    = _match_type
                    st.session_state.search_last_scope    = _scope
                    st.session_state.search_last_presence = _presence

                query_to_use    = st.session_state.get("search_last_query", _sq)
                match_to_use    = st.session_state.get("search_last_match", _match_type)
                scope_to_use    = st.session_state.get("search_last_scope", _scope)
                presence_to_use = st.session_state.get("search_last_presence", _presence)
                want_match      = (presence_to_use == "Содержит")

                if query_to_use.strip():
                    _results = []
                    for _d in db.get("documents", []):
                        _hits    = _doc_matches(_d, query_to_use, match_to_use, scope_to_use)
                        has_hits = len(_hits) > 0
                        if has_hits == want_match:
                            _results.append((_d, _hits))

                    _icon = "+" if want_match else "−"
                    st.caption(
                        f"{_icon} Найдено документов: **{len(_results)}** "
                        f"из {len(db.get('documents', []))} "
                        f"· запрос: «{query_to_use}» "
                        f"· {match_to_use} · {scope_to_use}"
                    )

                    if not _results:
                        st.warning(
                            "Ничего не найдено. Попробуйте изменить тип совпадения "
                            "или расширить область поиска."
                        )
                    else:
                        for _d, _hits in _results:
                            _dfn  = _fname(_d)
                            _pc   = _d.get("page_count", 0) or len(_d.get("pages", []))
                            _wc   = _d.get("word_count", 0)
                            _orig = _d.get("original_path", "")
                            _has_sum = bool(load_summary_file(_orig))
                            _exp_label = (
                                f"{'📄' if _hits else '○'} {_dfn} · "
                                f"{_pc} стр. · {_wc:,} сл. · "
                                f"{len(_hits)} совп."
                                + (" · есть пересказ" if _has_sum else "")
                            )
                            with st.expander(_exp_label, expanded=(len(_results) == 1)):
                                if _hits:
                                    for ctx, src in _hits[:5]:
                                        st.markdown(
                                            f"<div style='background:#f4f6f9;"
                                            f"border-left:3px solid #1B5C74;"
                                            f"padding:6px 10px;margin:4px 0;"
                                            f"border-radius:0 4px 4px 0;"
                                            f"font-size:0.85rem;'>{ctx}</div>",
                                            unsafe_allow_html=True,
                                        )
                                        st.caption(f"Источник: {src}")
                                elif not want_match:
                                    st.caption("Запрос не найден в этом документе")

                                # ── Кнопки действий ──────────────────────────
                                _ca, _cb = st.columns(2)
                                with _ca:
                                    st.download_button(
                                        "Скачать TXT",
                                        data=export_txt(_d),
                                        file_name=f"{os.path.splitext(_dfn)[0]}_распознан.txt",
                                        mime="text/plain",
                                        key=f"search_dl_txt_{_d['id']}",
                                        use_container_width=True,
                                    )
                                with _cb:
                                    _dsum_s = load_summary_file(_orig)
                                    try:
                                        _dbuf = export_docx(_d, summary=_dsum_s or None)
                                        st.download_button(
                                            "Скачать DOCX",
                                            data=_dbuf,
                                            file_name=f"{os.path.splitext(_dfn)[0]}_распознан.docx",
                                            mime=(
                                                "application/vnd.openxmlformats-"
                                                "officedocument.wordprocessingml.document"
                                            ),
                                            key=f"search_dl_docx_{_d['id']}",
                                            use_container_width=True,
                                        )
                                    except Exception:
                                        pass

    # =========================================================================
    # Вкладка 3 — База документов
    # =========================================================================
    with tab_docs:
        st.subheader("База распознанных документов")

        docs = db.get("documents", [])
        stats = db.get("stats", {})

        # ── Метрики ──────────────────────────────────────────────────────────
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Документов", stats.get("total", 0))
        m2.metric("Страниц", sum(d.get("page_count", 0) for d in docs))
        m3.metric("Слов", f"{sum(d.get('word_count', 0) for d in docs):,}")
        m4.metric("Последний", (stats.get("last_scan") or "—")[:10])

        if docs:
            st.divider()

            # ── Хранилище: метрики занятого места ────────────────────────────
            _sz_orig = _sz_sum = _sz_db = 0
            for _sd in docs:
                _op = _sd.get("original_path", "")
                if _op and os.path.exists(_op):
                    _sz_orig += os.path.getsize(_op)
                _sp2 = _summary_path(_op)
                if _sp2 and os.path.exists(_sp2):
                    _sz_sum += os.path.getsize(_sp2)
            try:
                _sz_db = os.path.getsize(_DB_PATH) if os.path.exists(_DB_PATH) else 0
            except Exception:
                pass
            def _fmt_sz(b):
                if b < 1024: return f"{b} Б"
                if b < 1024**2: return f"{b/1024:.1f} КБ"
                return f"{b/1024**2:.2f} МБ"
            _sz_total = _sz_orig + _sz_sum + _sz_db
            _sw1, _sw2, _sw3, _sw4 = st.columns(4)
            _sw1.metric("Итого на диске", _fmt_sz(_sz_total))
            _sw2.metric("Оригиналы", _fmt_sz(_sz_orig))
            _sw3.metric("Пересказы (txt)", _fmt_sz(_sz_sum))
            _sw4.metric("База (json)", _fmt_sz(_sz_db))

            st.divider()

            # ── Фильтры ───────────────────────────────────────────────────────
            ff1, ff2, ff3 = st.columns([2, 2, 1])
            with ff1:
                _db_search = st.text_input(
                    "Поиск по имени",
                    placeholder="Поиск по имени файла...",
                    key="db_search",
                )
            with ff2:
                _all_exts = sorted({
                    os.path.splitext(_fname(d))[1].lower().lstrip(".")
                    for d in docs
                    if os.path.splitext(_fname(d))[1]
                })
                _db_fmt = st.multiselect(
                    "Формат файла",
                    options=_all_exts,
                    key="db_fmt",
                    placeholder="Все форматы",
                )
            with ff3:
                _db_only_sum = st.toggle("Только с пересказом", key="db_only_sum")

            ff4, ff5 = st.columns([2, 2])
            with ff4:
                _days_opts = [1, 3, 7, 14, 30, 90, 0]
                _days_labels = {1:"Сегодня",3:"3 дня",7:"7 дней",
                                14:"2 недели",30:"Месяц",90:"3 месяца",0:"Все время"}
                _db_days = st.select_slider(
                    "Период загрузки",
                    options=_days_opts,
                    value=7,
                    format_func=lambda x: _days_labels[x],
                    key="db_days",
                )
            with ff5:
                _db_sort = st.selectbox(
                    "Сортировка",
                    ["По дате (новые)", "По дате (старые)", "По имени А-Я",
                     "По размеру (больше)", "Сначала с пересказом"],
                    key="db_sort",
                )

            # ── Применяем фильтры ─────────────────────────────────────────────
            from datetime import datetime as _dt, timedelta as _td
            _filtered = list(reversed(docs))

            # Фильтр по дате
            if _db_days > 0:
                _cutoff = _dt.now() - _td(days=_db_days)
                def _doc_date(d):
                    try:
                        return _dt.fromisoformat(d.get("processed_at","1970-01-01"))
                    except Exception:
                        return _dt(1970,1,1)
                _filtered = [d for d in _filtered if _doc_date(d) >= _cutoff]

            # Фильтр по формату
            if _db_fmt:
                _filtered = [
                    d for d in _filtered
                    if os.path.splitext(_fname(d))[1].lower().lstrip(".") in _db_fmt
                ]

            # Фильтр по имени
            if _db_search.strip():
                _q = _db_search.strip().lower()
                _filtered = [d for d in _filtered if _q in _fname(d).lower()]

            # Фильтр "только с пересказом"
            if _db_only_sum:
                _filtered = [
                    d for d in _filtered
                    if st.session_state.get(f"summary_{d['id']}")
                    or load_summary_file(d.get("original_path", ""))
                ]

            # Сортировка
            if _db_sort == "По дате (старые)":
                _filtered = list(reversed(_filtered))
            elif _db_sort == "По имени А-Я":
                _filtered.sort(key=lambda d: _fname(d).lower())
            elif _db_sort == "По размеру (больше)":
                _filtered.sort(key=lambda d: d.get("word_count", 0), reverse=True)
            elif _db_sort == "Сначала с пересказом":
                def _has_sum(d):
                    return bool(
                        st.session_state.get(f"summary_{d['id']}")
                        or load_summary_file(d.get("original_path", ""))
                    )
                _filtered.sort(key=_has_sum, reverse=True)

            # Подсказка активных фильтров
            _active = []
            if _db_days > 0:
                _active.append(f"за {_days_labels[_db_days].lower()}")
            if _db_fmt:
                _active.append(f"форматы: {', '.join(_db_fmt)}")
            if _db_only_sum:
                _active.append("только с пересказом")
            _flt_hint = "  ·  ".join(_active) if _active else "все документы"
            st.caption(f"Показано: {len(_filtered)} из {len(docs)}  ·  {_flt_hint}")

            # ── Список карточек ──────────────────────────────────────────────
            _open_key = "db_open_card"
            if _open_key not in st.session_state:
                st.session_state[_open_key] = None

            for _d in _filtered:
                _did  = _d["id"]
                _dfn  = _fname(_d)
                _dpc  = _d.get("page_count") or len(_d.get("pages", []))
                _dwc  = _d.get("word_count", 0)
                _docr = _d.get("ocr_pages", 0)
                _ddt  = (_d.get("processed_at") or "")[:10]
                _dorig = _d.get("original_path", "")
                _dsum  = (
                    st.session_state.get(f"summary_{_did}")
                    or load_summary_file(_dorig)
                )
                _has_s  = bool(_dsum)
                _orig_ok = bool(_dorig and os.path.exists(_dorig))
                _is_open = st.session_state[_open_key] == _did

                # Карточка — шапка
                _badge = "✅ пересказ" if _has_s else "— пересказа нет"
                _orig_badge = "" if not _dorig else (
                    "  ·  файл доступен" if _orig_ok else "  ·  ⚠ файл перемещён"
                )
                st.markdown(
                    f"<div style='background:#ffffff;border:1px solid #dce3ec;"
                    f"border-radius:6px;padding:10px 14px;margin-bottom:4px;'>"
                    f"<div style='font-weight:600;font-size:0.95rem;"
                    f"color:#1a2a3a;word-break:break-all;'>{_dfn}</div>"
                    f"<div style='font-size:0.75rem;color:#5a6a7a;margin-top:3px;'>"
                    f"{_dpc} стр. · {_dwc:,} слов · OCR: {_docr} · {_ddt}"
                    f"  ·  {_badge}{_orig_badge}</div>"
                    f"</div>",
                    unsafe_allow_html=True,
                )

                # Кнопки-действия в строке
                _ca, _cb_btn, _cc, _cd = st.columns([2, 1, 1, 1])
                with _ca:
                    _btn_label = "Скрыть детали" if _is_open else "Показать детали"
                    if st.button(_btn_label, key=f"db_toggle_{_did}",
                                 use_container_width=True):
                        st.session_state[_open_key] = None if _is_open else _did
                        st.rerun()
                with _cb_btn:
                    _file_label = "Открыть файл" if _orig_ok else "Файл недоступен"
                    if st.button(_file_label, key=f"db_open_{_did}",
                                 disabled=not _orig_ok, use_container_width=True):
                        try:
                            import subprocess as _sp, sys as _sys
                            if _sys.platform == "win32":
                                os.startfile(_dorig)
                            elif _sys.platform == "darwin":
                                _sp.Popen(["open", _dorig])
                            else:
                                _sp.Popen(["xdg-open", _dorig])
                        except Exception as _oe:
                            st.error(str(_oe))
                with _cc:
                    if _has_s:
                        _sfname = os.path.basename(_summary_path(_dorig))                                   or f"{os.path.splitext(_dfn)[0]}_пересказ.txt"
                        st.download_button(
                            "Скачать пересказ",
                            data=_dsum.encode("utf-8"),
                            file_name=_sfname,
                            mime="text/plain",
                            key=f"db_dl_sum_{_did}",
                            use_container_width=True,
                        )
                    else:
                        st.button("Пересказа нет", key=f"db_nosum_{_did}",
                                  disabled=True, use_container_width=True)
                with _cd:
                    _ce, _cf = st.columns(2)
                    with _ce:
                        st.download_button(
                            "Скачать TXT",
                            data=export_txt(_d),
                            file_name=f"{os.path.splitext(_dfn)[0]}_распознан.txt",
                            mime="text/plain",
                            key=f"db_dl_txt2_{_did}",
                            use_container_width=True,
                        )
                    with _cf:
                        try:
                            _docx_buf = export_docx(_d, summary=_dsum or None)
                            st.download_button(
                                "Скачать DOCX",
                                data=_docx_buf,
                                file_name=f"{os.path.splitext(_dfn)[0]}_распознан.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"db_dl_docx_{_did}",
                                use_container_width=True,
                            )
                        except Exception:
                            pass

                # Развёрнутая карточка
                if _is_open:
                    with st.container():
                        st.markdown(
                            "<div style='border:1px solid #dce3ec;border-top:none;"
                            "border-radius:0 0 6px 6px;padding:14px 16px;"
                            "background:#fafbfc;margin-bottom:12px;'>",
                            unsafe_allow_html=True,
                        )

                        # Навигация по страницам
                        _pages  = _d.get("pages", [])
                        _pkey   = f"db_page_{_did}"
                        if _pkey not in st.session_state:
                            st.session_state[_pkey] = 0
                        _pidx = max(0, min(st.session_state[_pkey], len(_pages)-1))
                        _pcur = _pages[_pidx] if _pages else {}

                        _n1, _n2, _n3 = st.columns([1, 4, 1])
                        with _n1:
                            if st.button("◀", key=f"db_prev_{_did}",
                                         disabled=(_pidx == 0)):
                                st.session_state[_pkey] = _pidx - 1
                                st.rerun()
                        with _n2:
                            _pw = _pcur.get("word_count") or len(_pcur.get("text","").split())
                            st.markdown(
                                f"<div style='text-align:center;padding-top:6px;'>"
                                f"Страница <b>{_pidx+1}</b> из <b>{len(_pages)}</b> "
                                f"<span style='color:#888;font-size:0.82em;'>"
                                f"({_pcur.get('method','')} · {_pw} слов)</span></div>",
                                unsafe_allow_html=True,
                            )
                        with _n3:
                            if st.button("▶", key=f"db_next_{_did}",
                                         disabled=(_pidx >= len(_pages)-1)):
                                st.session_state[_pkey] = _pidx + 1
                                st.rerun()

                        st.markdown(
                            f"<div style='font-size:0.85em;line-height:1.55;"
                            f"background:#f8f9fa;border:1px solid #e0e0e0;"
                            f"border-radius:6px;padding:10px 12px;margin-top:6px;"
                            f"max-height:280px;overflow-y:auto;"
                            f"white-space:pre-wrap;word-break:break-word;'>"
                            f"{_pcur.get('text','').replace('<','&lt;').replace('>','&gt;')}"
                            f"</div>",
                            unsafe_allow_html=True,
                        )

                        # Пересказ
                        if _dsum:
                            st.divider()
                            _swc = len(_dsum.split())
                            st.markdown(
                                f"<div style='display:flex;justify-content:space-between;"
                                f"align-items:center;margin-bottom:4px;'>"
                                f"<span style='font-weight:600;'>Пересказ</span>"
                                f"<span style='color:#5a6a7a;font-size:0.78em;'>{_swc} слов</span>"
                                f"</div>",
                                unsafe_allow_html=True,
                            )
                            _sh = max(120, min(400, _swc * 6))
                            st.markdown(
                                f"<div style='font-size:0.85em;line-height:1.6;"
                                f"background:#f0f4f8;border:1px solid #d0d8e4;"
                                f"border-radius:6px;padding:12px 14px;"
                                f"max-height:{_sh}px;overflow-y:auto;"
                                f"white-space:pre-wrap;word-break:break-word;'>"
                                f"{_dsum.replace('<','&lt;').replace('>','&gt;')}"
                                f"</div>",
                                unsafe_allow_html=True,
                            )

                        st.markdown("</div>", unsafe_allow_html=True)
                else:
                    st.markdown("<div style='margin-bottom:8px;'></div>",
                                unsafe_allow_html=True)

            st.divider()

            # Очистка базы
            _clr_key = "confirm_clear_db"
            _clr_pending = st.session_state.get(_clr_key, False)

            if not _clr_pending:
                if st.button("Очистить базу документов", type="secondary",
                             key="clear_scan_db", use_container_width=True):
                    st.session_state[_clr_key] = True
                    st.rerun()
            else:
                st.warning(
                    "Будут удалены все оригиналы файлов, результаты распознавания "
                    "и пересказы. Это действие необратимо."
                )
                _cy, _cn = st.columns(2)
                if _cy.button("Да, удалить всё", key="clear_confirm_yes",
                              type="primary", use_container_width=True):
                    # Удаляем файлы originals + пересказы
                    _del_ok, _del_err = 0, 0
                    for _ddoc in docs:
                        for _fpath in [
                            _ddoc.get("original_path", ""),
                            _summary_path(_ddoc.get("original_path", "")),
                        ]:
                            if _fpath and os.path.exists(_fpath):
                                try:
                                    os.remove(_fpath)
                                    _del_ok += 1
                                except Exception:
                                    _del_err += 1

                    # Сбрасываем БД
                    _empty_db = {"documents": [], "stats": {"total": 0, "last_scan": None}}
                    st.session_state["scanner_db"] = _empty_db
                    save_db(_empty_db)

                    # Очищаем session_state от пересказов и прочего
                    for _k in list(st.session_state.keys()):
                        if any(_k.startswith(p) for p in (
                            "summary_", "gen_trigger_", "pending_confirm_",
                            "exp_open_", "page_idx_", "sum_mode_", "sum_words_",
                        )):
                            st.session_state.pop(_k, None)
                    st.session_state.pop("last_scanned_ids", None)
                    st.session_state.pop("search_results", None)
                    st.session_state.pop("scan_sel_id", None)
                    st.session_state[_clr_key] = False

                    _msg = f"База очищена. Удалено файлов: {_del_ok}."
                    if _del_err:
                        _msg += f" Не удалось удалить: {_del_err}."
                    st.success(_msg)
                    st.rerun()

                if _cn.button("Отмена", key="clear_confirm_no",
                              use_container_width=True):
                    st.session_state[_clr_key] = False
                    st.rerun()
        else:
            st.info("📭 База пуста. Загрузите документы на вкладке «Загрузка».")