# streamlit_pages/risk_calculator.py
import streamlit as st
import json
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# =============================================================================
# 🛠 Функции
# =============================================================================

def load_json_data(uploaded_file):
    """Загружает JSON из Экспорта ФГИС"""
    try:
        uploaded_file.seek(0)
        data = json.load(uploaded_file)
        return data
    except Exception as e:
        st.error(f"❌ Ошибка чтения JSON: {e}")
        return None

def calculate_article_risks(article, sphere, total_revenue):
    """Рассчитывает риски по статье затрат (включение/исключение)"""
    
    code = article.get("code", "")
    name = article.get("name", "").lower()
    amount = article.get("amount", 0)
    percent = (amount / total_revenue * 100) if total_revenue > 0 else 0
    
    # Базовые риски
    inclusion_risk = 50  # Риск исключения статьи из заявки
    exclusion_risk = 50  # Риск если статью не включить
    
    # Шаблоны обоснований
    inclusion_reason = "Требуется дополнительное обоснование"
    exclusion_reason = "Статья не является обязательной по методике"
    recommendations_inc = []
    recommendations_exc = []
    
    # Анализ по ключевым словам
    if "ремонт" in name or "ос" in name:
        inclusion_risk = 70
        inclusion_reason = "Высокий риск: часто исключают без дефектной ведомости"
        recommendations_inc = [
            "📎 Приложите дефектную ведомость",
            "📎 Приложите акт осмотра ОС",
            "📜 Ссылка: Приказ ФАС №1746-э, п. 67"
        ]
        exclusion_risk = 30
        exclusion_reason = "Исключение безопасно: не обязательные затраты"
        recommendations_exc = ["⚠️ Приведёт к недофинансированию ремонта"]
    
    elif "зарплат" in name or "труд" in name or "ауп" in name:
        inclusion_risk = 40
        inclusion_reason = "Средний риск: сверьте с нормативом численности"
        recommendations_inc = [
            "📎 Приложите штатное расписание",
            "📎 Приложите приказ о зарплате",
            "📜 Ссылка: Приказ ФАС №1746-э, п. 45-50"
        ]
        exclusion_risk = 85
        exclusion_reason = "Высокий риск: обязательные затраты по методике"
        recommendations_exc = ["⚠️ Нарушение методики расчёта", "⚠️ Претензии регулятора"]
    
    elif "электроэнерги" in name or "тепло" in name or "топлив" in name:
        inclusion_risk = 15
        inclusion_reason = "Низкий риск: обязательные производственные затраты"
        recommendations_inc = ["📎 Приложите счета от поставщиков"]
        exclusion_risk = 95
        exclusion_reason = "Критический риск: невозможно без этих затрат"
        recommendations_exc = ["⚠️ Невозможно оказывать услугу", "⚠️ 100% отказ"]
    
    elif "программ" in name or "по" in name or "лицензи" in name:
        inclusion_risk = 55
        inclusion_reason = "Средний риск: требуется обоснование необходимости"
        recommendations_inc = [
            "📎 Приложите лицензионный договор",
            "📎 Обоснуйте необходимость для деятельности",
            "📜 Ссылка: Письмо ФАС №АЦ/30445-ПР"
        ]
        exclusion_risk = 40
        exclusion_reason = "Исключение возможно: не производственные затраты"
        recommendations_exc = ["⚠️ Ухудшение учёта и отчётности"]
    
    elif "тко" in name or "транспортирован" in name or "полигон" in name:
        inclusion_risk = 25
        inclusion_reason = "Низкий риск: профильные затраты для ТКО"
        recommendations_inc = ["📎 Приложите договоры с полигоном"]
        exclusion_risk = 90
        exclusion_reason = "Высокий риск: обязательные для ТКО затраты"
        recommendations_exc = ["⚠️ Невозможно оказывать услугу"]
    
    elif "вод" in name or "перекачк" in name or "химреагент" in name:
        inclusion_risk = 20
        inclusion_reason = "Низкий риск: профильные затраты для водоснабжения"
        recommendations_inc = ["📎 Приложите договоры с поставщиками"]
        exclusion_risk = 92
        exclusion_reason = "Высокий риск: обязательные для водоснабжения затраты"
        recommendations_exc = ["⚠️ Невозможно оказывать услугу"]
    
    # Корректировка по доле в выручке
    if percent > 20:
        inclusion_risk = min(100, inclusion_risk + 15)
        inclusion_reason += " + высокая доля в выручке"
    
    if percent < 1:
        exclusion_risk = max(0, exclusion_risk - 20)
        exclusion_reason += " + низкая доля"
    
    # Определение статуса
    if inclusion_risk >= 70:
        inclusion_status = "🔴"
    elif inclusion_risk >= 40:
        inclusion_status = "🟡"
    else:
        inclusion_status = "🟢"
    
    if exclusion_risk >= 70:
        exclusion_status = "🔴"
    elif exclusion_risk >= 40:
        exclusion_status = "🟡"
    else:
        exclusion_status = "🟢"
    
    return {
        "code": code,
        "name": article.get("name", ""),
        "amount": amount,
        "percent": round(percent, 1),
        "inclusion_risk": inclusion_risk,
        "inclusion_status": inclusion_status,
        "inclusion_reason": inclusion_reason,
        "inclusion_recommendations": recommendations_inc,
        "exclusion_risk": exclusion_risk,
        "exclusion_status": exclusion_status,
        "exclusion_reason": exclusion_reason,
        "exclusion_recommendations": recommendations_exc
    }

def calculate_application_risk(articles_risks, completeness_score=80):
    """Рассчитывает общий риск заявки"""
    
    if not articles_risks:
        return 50, "🟡", "Нет данных для расчёта"
    
    # Средневзвешенный риск по статьям
    total_amount = sum(a["amount"] for a in articles_risks)
    weighted_risk = 0
    
    for article in articles_risks:
        weight = article["amount"] / total_amount if total_amount > 0 else 0
        weighted_risk += article["inclusion_risk"] * weight
    
    # Учёт комплектности
    final_risk = weighted_risk * 0.7 + (100 - completeness_score) * 0.3
    
    # Определение статуса
    if final_risk >= 70:
        status = "🔴"
        reason = "Высокий риск отказа: требуется доработка"
    elif final_risk >= 40:
        status = "🟡"
        reason = "Средний риск: рекомендуются улучшения"
    else:
        status = "🟢"
        reason = "Низкий риск: заявка готова к подаче"
    
    return round(final_risk, 1), status, reason

def generate_risk_report(articles_risks, application_risk, application_status, org_name, inn):
    """Генерирует DOCX отчёт о рисках"""
    doc = Document()
    
    # Стили
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    # Заголовок
    p = doc.add_paragraph("ОТЧЁТ О РИСКАХ ТАРИФНОЙ ЗАЯВКИ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph(f"Организация: {org_name}")
    doc.add_paragraph(f"ИНН: {inn}")
    doc.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    
    # Общая оценка
    doc.add_paragraph()
    p = doc.add_paragraph("ОБЩАЯ ОЦЕНКА РИСКОВ")
    run = p.runs[0]
    run.bold = True
    
    doc.add_paragraph(f"Вероятность одобрения: {100 - application_risk:.0f}% {application_status}")
    
    # Таблица рисков
    doc.add_paragraph()
    p = doc.add_paragraph("РИСКИ ПО СТАТЬЯМ ЗАТРАТ")
    run = p.runs[0]
    run.bold = True
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    headers = ["Статья", "Сумма (₽)", "% от ВВ", "Риск вкл.", "Риск искл.", "Статус"]
    for i, h in enumerate(headers):
        header_cells[i].text = h
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    for article in articles_risks[:20]:  # ТОП-20
        row = table.add_row().cells
        row[0].text = article["name"][:40]
        row[1].text = f"{article['amount']:,.0f}"
        row[2].text = f"{article['percent']:.1f}%"
        row[3].text = f"{article['inclusion_risk']}%"
        row[4].text = f"{article['exclusion_risk']}%"
        row[5].text = f"{article['inclusion_status']}/{article['exclusion_status']}"
    
    # Рекомендации
    doc.add_paragraph()
    p = doc.add_paragraph("РЕКОМЕНДАЦИИ ПО СНИЖЕНИЮ РИСКОВ")
    run = p.runs[0]
    run.bold = True
    
    high_risk_articles = [a for a in articles_risks if a["inclusion_risk"] >= 70]
    if high_risk_articles:
        for article in high_risk_articles[:5]:
            doc.add_paragraph(f"• {article['name']}: {article['inclusion_reason']}", style='List Bullet')
            for rec in article["inclusion_recommendations"][:2]:
                doc.add_paragraph(f"  - {rec}", style='List Bullet')
    else:
        doc.add_paragraph("Критических рисков не выявлено")
    
    return doc

def validate_risk_calculator(fgis_data, articles_risks):
    """Проверяет данные для калькулятора"""
    errors = []
    
    if not fgis_data:
        errors.append("❌ Загрузите JSON из Экспорта ФГИС")
    
    if not articles_risks or len(articles_risks) == 0:
        errors.append("❌ Нет данных для расчёта рисков")
    
    return errors

# =============================================================================
# 🎨 Интерфейс Streamlit
# =============================================================================

def show_risk_calculator():
    """Страница калькулятора рисков"""
    
    st.header("📊 Калькулятор рисков тарифной заявки")
    st.info("📌 Загрузите JSON из Экспорта ФГИС — система оценит риски по заявке и каждой статье затрат")
    
    # ─────────────────────────────────────────────────────────────────────
    # 📁 Шаг 1: Загрузка JSON из Экспорта ФГИС
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("1️⃣ Загрузка данных из Экспорта ФГИС")
    
    uploaded_json = st.file_uploader(
        "📄 Загрузите JSON файл из модуля Экспорт ФГИС",
        type=['json'],
        help="Файл должен быть сформирован в модуле Экспорт в ФГИС Тариф"
    )
    
    fgis_data = None
    articles_risks = []
    
    if uploaded_json:
        fgis_data = load_json_data(uploaded_json)
        
        if fgis_data:
            st.success(f"✅ JSON загружен: {uploaded_json.name}")
            
            # Извлекаем статьи затрат
            articles = fgis_data.get("application", {}).get("revenue", {}).get("articles", [])
            total_revenue = fgis_data.get("application", {}).get("revenue", {}).get("total", 0)
            sphere = fgis_data.get("application", {}).get("activity_sphere", "")
            
            if articles:
                # Рассчитываем риски по каждой статье
                for article in articles:
                    risk_data = calculate_article_risks(article, sphere, total_revenue)
                    articles_risks.append(risk_data)
                
                # Сортируем по риску включения (убывание)
                articles_risks.sort(key=lambda x: x["inclusion_risk"], reverse=True)
                
                # Общая оценка заявки
                app_risk, app_status, app_reason = calculate_application_risk(articles_risks)
                
                # Сохраняем в session_state
                st.session_state.risk_articles = articles_risks
                st.session_state.risk_app = app_risk
                st.session_state.risk_status = app_status
                st.session_state.risk_reason = app_reason
                
                # Показываем сводку
                col1, col2, col3 = st.columns(3)
                col1.metric("📊 Статьей затрат", len(articles_risks))
                col2.metric("💰 Валовая выручка", f"{total_revenue:,.0f} ₽")
                col3.metric("🎯 Риск заявки", f"{app_risk}%", delta_color="inverse")
    
    st.divider()
    
    # ─────────────────────────────────────────────────────────────────────
    # 📊 Шаг 2: Общая оценка рисков
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("2️⃣ Общая оценка рисков заявки")
    
    if "risk_app" in st.session_state:
        app_risk = st.session_state.risk_app
        app_status = st.session_state.risk_status
        app_reason = st.session_state.risk_reason
        
        # Крупный индикатор
        approval_chance = 100 - app_risk
        
        col1, col2, col3 = st.columns([2, 1, 1])
        
        with col1:
            st.metric(
                "Вероятность одобрения заявки",
                f"{approval_chance:.0f}%",
                delta=f"{app_status} {app_reason}",
                delta_color="normal" if approval_chance >= 70 else "inverse"
            )
        
        with col2:
            if app_risk >= 70:
                st.error("🔴 Высокий риск")
            elif app_risk >= 40:
                st.warning("🟡 Средний риск")
            else:
                st.success("🟢 Низкий риск")
        
        with col3:
            st.info(f"Риск отказа: {app_risk:.0f}%")
        
        # Прогресс-бар
        st.progress(approval_chance / 100)
        st.caption(f"Цель: 85%+ для уверенного одобрения")
    
    else:
        st.info("💡 Загрузите JSON файл для расчёта рисков")
    
    st.divider()
    
    # ─────────────────────────────────────────────────────────────────────
    # 📋 Шаг 3: Матрица рисков по статьям затрат
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("3️⃣ Матрица рисков по статьям затрат")
    st.info("📌 Риск включения = вероятность исключения статьи регулятором | Риск исключения = последствия если не включить")
    
    if "risk_articles" in st.session_state:
        articles_risks = st.session_state.risk_articles
        
        # Фильтры
        col1, col2 = st.columns(2)
        with col1:
            risk_filter = st.selectbox(
                "Фильтр по риску",
                ["Все", "🔴 Высокий (≥70%)", "🟡 Средний (40-69%)", "🟢 Низкий (<40%)"]
            )
        with col2:
            sort_by = st.selectbox(
                "Сортировка",
                ["По риску включения", "По сумме", "По доле в выручке"]
            )
        
        # Применяем фильтры
        filtered_articles = articles_risks.copy()
        
        if risk_filter == "🔴 Высокий (≥70%)":
            filtered_articles = [a for a in articles_risks if a["inclusion_risk"] >= 70]
        elif risk_filter == "🟡 Средний (40-69%)":
            filtered_articles = [a for a in articles_risks if 40 <= a["inclusion_risk"] < 70]
        elif risk_filter == "🟢 Низкий (<40%)":
            filtered_articles = [a for a in articles_risks if a["inclusion_risk"] < 40]
        
        if sort_by == "По сумме":
            filtered_articles.sort(key=lambda x: x["amount"], reverse=True)
        elif sort_by == "По доле в выручке":
            filtered_articles.sort(key=lambda x: x["percent"], reverse=True)
        
        st.write(f"**📊 Показано {len(filtered_articles)} из {len(articles_risks)} статей**")
        
        # Таблица рисков
        df_risks = pd.DataFrame(filtered_articles)
        
        # Форматируем для отображения
        df_display = df_risks[[
            "code", "name", "amount", "percent",
            "inclusion_status", "inclusion_risk",
            "exclusion_status", "exclusion_risk"
        ]].copy()
        
        df_display.columns = [
            "Код", "Статья", "Сумма (₽)", "Доля (%)",
            "Риск вкл.", "Риск вкл. %", "Риск искл.", "Риск искл. %"
        ]
        
        # Форматирование
        df_display["Сумма (₽)"] = df_display["Сумма (₽)"].apply(lambda x: f"{x:,.0f}")
        df_display["Доля (%)"] = df_display["Доля (%)"].apply(lambda x: f"{x:.1f}%")
        
        st.dataframe(df_display, use_container_width=True, hide_index=True)
        
        # Детализация по клику
        st.subheader("🔍 Детализация по статье")
        
        selected_article = st.selectbox(
            "Выберите статью для детального анализа",
            options=[f"{a['code']}: {a['name']}" for a in filtered_articles],
            format_func=lambda x: x[:80] + "..." if len(x) > 80 else x
        )
        
        if selected_article:
            article_code = selected_article.split(":")[0]
            article_data = next((a for a in filtered_articles if a["code"] == article_code), None)
            
            if article_data:
                col1, col2 = st.columns(2)
                
                with col1:
                    st.info(f"**Статья {article_data['code']}:** {article_data['name']}")
                    st.metric("Сумма", f"{article_data['amount']:,.0f} ₽")
                    st.metric("Доля в выручке", f"{article_data['percent']:.1f}%")
                
                with col2:
                    st.metric("Риск включения", f"{article_data['inclusion_risk']}% {article_data['inclusion_status']}")
                    st.metric("Риск исключения", f"{article_data['exclusion_risk']}% {article_data['exclusion_status']}")
                
                st.divider()
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.write("**🔴 Риск включения (исключение регулятором):**")
                    st.write(article_data["inclusion_reason"])
                    st.write("**Рекомендации:**")
                    for rec in article_data["inclusion_recommendations"]:
                        st.write(f"• {rec}")
                
                with col2:
                    st.write("**🟢 Риск исключения (если не включить):**")
                    st.write(article_data["exclusion_reason"])
                    st.write("**Последствия:**")
                    for rec in article_data["exclusion_recommendations"]:
                        st.write(f"• {rec}")
    
    else:
        st.info("💡 Загрузите JSON файл для расчёта рисков")
    
    st.divider()
    
    # ─────────────────────────────────────────────────────────────────────
    # 💡 Шаг 4: Рекомендации по снижению рисков
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("4️⃣ Рекомендации по снижению рисков")
    
    if "risk_articles" in st.session_state:
        articles_risks = st.session_state.risk_articles
        
        # ТОП-5 рисков
        high_risk = [a for a in articles_risks if a["inclusion_risk"] >= 70][:5]
        medium_risk = [a for a in articles_risks if 40 <= a["inclusion_risk"] < 70][:5]
        
        if high_risk:
            st.error(f"🔴 Найдено {len(high_risk)} статей с высоким риском:")
            for i, article in enumerate(high_risk, 1):
                with st.expander(f"{i}. {article['name']} (риск {article['inclusion_risk']}%)"):
                    st.write(f"**Причина:** {article['inclusion_reason']}")
                    st.write("**Рекомендации:**")
                    for rec in article["inclusion_recommendations"]:
                        st.write(f"• {rec}")
        else:
            st.success("✅ Нет статей с высоким риском!")
        
        if medium_risk:
            st.warning(f"🟡 Найдено {len(medium_risk)} статей со средним риском:")
            for i, article in enumerate(medium_risk, 1):
                st.write(f"{i}. **{article['name']}** — {article['inclusion_reason']}")
        
        # Общие рекомендации
        st.divider()
        st.subheader("📋 Общие рекомендации")
        
        approval_chance = 100 - st.session_state.risk_app
        
        if approval_chance < 50:
            st.error("⚠️ Критически низкая вероятность одобрения. Рекомендуется:")
            st.write("• Пересмотреть спорные статьи затрат")
            st.write("• Подготовить дополнительные обоснования")
            st.write("• Провести предварительные консультации с регулятором")
        elif approval_chance < 70:
            st.warning("ℹ️ Средняя вероятность одобрения. Рекомендуется:")
            st.write("• Усилить обоснование по статьям с риском ≥40%")
            st.write("• Проверить комплектность документов")
            st.write("• Добавить ссылки на нормативные акты")
        else:
            st.success("✅ Хорошая вероятность одобрения. Рекомендуется:")
            st.write("• Проверить актуальность всех документов")
            st.write("• Подготовить пояснительную записку")
            st.write("• Можно подавать заявку")
    
    else:
        st.info("💡 Загрузите JSON файл для получения рекомендаций")
    
    st.divider()
    
    # ─────────────────────────────────────────────────────────────────────
    # 📥 Шаг 5: Экспорт отчёта
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("5️⃣ Экспорт отчёта о рисках")
    
    if "risk_articles" in st.session_state and fgis_data:
        if st.button("📄 Сформировать отчёт о рисках", type="primary", use_container_width=True):
            org_name = fgis_data.get("header", {}).get("organization_name", "Organization")
            inn = fgis_data.get("header", {}).get("inn", "0000000000")
            
            doc = generate_risk_report(
                st.session_state.risk_articles,
                st.session_state.risk_app,
                st.session_state.risk_status,
                org_name,
                inn
            )
            
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            file_id = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"RiskReport_{inn}_{file_id}.docx"
            
            st.success("✅ Отчёт сформирован!")
            
            st.download_button(
                label="📥 Скачать отчёт (DOCX)",
                data=doc_buffer,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
            
            st.caption(f"📁 Файл: {filename}")
            st.info("💡 Отчёт содержит: общую оценку, матрицу рисков, рекомендации по статьям")
    
    else:
        st.info("💡 Загрузите JSON файл для экспорта отчёта")
    
    # ─────────────────────────────────────────────────────────────────────
    # 💡 Пример
    # ─────────────────────────────────────────────────────────────────────
    with st.expander("💡 Как использовать калькулятор рисков", expanded=False):
        st.write("""
        **Шаги работы:**
        
        1. **Загрузите JSON** из модуля "Экспорт ФГИС"
        2. **Изучите общую оценку** — вероятность одобрения заявки
        3. **Проанализируйте матрицу рисков** — по каждой статье затрат
        4. **Изучите рекомендации** — что сделать для снижения рисков
        5. **Скачайте отчёт** — для внутреннего согласования или приложения к заявке
        
        **Интерпретация рисков:**
        
        - 🔴 **Высокий риск (≥70%)**: Статья很可能 будет исключена регулятором
        - 🟡 **Средний риск (40-69%)**: Требуется дополнительное обоснование
        - 🟢 **Низкий риск (<40%)**: Статья скорее всего будет одобрена
        
        **Риск включения vs Риск исключения:**
        
        - **Риск включения**: Вероятность что регулятор исключит статью из заявки
        - **Риск исключения**: Последствия если вы сами не включите эту статью
        """)

# =============================================================================
# 🚀 Запуск
# =============================================================================

if __name__ == "__main__":
    show_risk_calculator()