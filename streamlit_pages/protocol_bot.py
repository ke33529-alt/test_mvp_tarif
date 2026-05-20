# streamlit_pages/protocol_bot.py
"""
Робот-Протокольщик
─────────────────────────────────────────────────────────────────────────────
• Создаёт официальные протоколы из текста, документов (TXT/DOCX/PDF) и аудио
• Аудио: WAV всегда; M4A/MP3/OGG/AAC/3GP — при наличии ffmpeg в PATH
• Хранение: DOCX + TXT на диске, JSON-индекс
• Вкладка «База протоколов» — карточки с деталями / скачиванием
─────────────────────────────────────────────────────────────────────────────
"""
import io
import json
import os
import subprocess
import sys
import threading
from datetime import datetime
from typing import Dict, List, Optional

import streamlit as st

# Принудительно переключаем stdout на UTF-8 для Windows PowerShell
# (иначе эмодзи в print() вызывают UnicodeEncodeError на cp1251)
if hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

# =============================================================================
# Константы и пути
# =============================================================================
_BASE_DIR       = os.path.join("data", "protocol_bot")
_DB_PATH        = os.path.join(_BASE_DIR, "protocols_db.json")
_PROTOCOLS_DIR  = os.path.join(_BASE_DIR, "protocols")   # DOCX + TXT
_TEMP_DIR       = os.path.join(_BASE_DIR, "temp")
_DB_LOCK        = threading.Lock()

SUPPORTED_AUDIO_FFMPEG = ["m4a", "mp3", "wav", "ogg", "mp4", "3gp", "amr", "aac", "wma", "oga"]
SUPPORTED_AUDIO_NOFFMPEG = ["wav"]
SUPPORTED_DOCS  = ["txt", "docx", "doc"]

DEFAULT_STRUCTURE = """\
1. Дата и время встречи
2. Присутствовали
3. Повестка дня
4. Обсуждаемые вопросы
5. Принятые решения
6. Поручения (кто, что, срок)
7. Следующая встреча"""

DETAIL_CAPTIONS = {
    "краткий":   "Только ключевые факты и решения",
    "средний":   "Факты + важные детали, оптимальный объём",
    "подробный": "Все факты, детали, цитаты, полный контекст",
}
DETAIL_MAX_TOKENS = {"краткий": 1200, "средний": 2500, "подробный": 5000}


# =============================================================================
# Утилиты
# =============================================================================
def _ensure_dirs():
    for d in [_BASE_DIR, _PROTOCOLS_DIR, _TEMP_DIR]:
        os.makedirs(d, exist_ok=True)


def _fmt_size(b: int) -> str:
    if b < 1024:
        return f"{b} Б"
    if b < 1024 ** 2:
        return f"{b / 1024:.1f} КБ"
    return f"{b / 1024 ** 2:.2f} МБ"


# =============================================================================
# Проверка ffmpeg
# =============================================================================

# Прямой путь к ffmpeg.exe — используется если ffmpeg не найден в системном PATH.
# При необходимости скорректируйте под своё расположение.
_FFMPEG_DIRECT = r"C:\ffmpeg-2026-05-18-git-b4d11dffbf-essentials_build\bin\ffmpeg.exe"

_ffmpeg_cache: Optional[bool] = None
_ffmpeg_exe:   Optional[str]  = None   # реальная команда/путь для запуска


def _ffmpeg_available() -> bool:
    """
    Ищет ffmpeg в трёх местах (в порядке приоритета):
    1. Системный PATH («ffmpeg»)
    2. Прямой путь _FFMPEG_DIRECT
    3. Папка bin рядом с exe-файлом скрипта
    Результат кешируется.
    """
    global _ffmpeg_cache, _ffmpeg_exe
    if _ffmpeg_cache is not None:
        return _ffmpeg_cache

    candidates = ["ffmpeg", _FFMPEG_DIRECT]

    # Ещё один вариант: bin\ рядом с текущим файлом
    _local = os.path.join(os.path.dirname(os.path.abspath(__file__)), "bin", "ffmpeg.exe")
    if os.path.exists(_local):
        candidates.append(_local)

    for candidate in candidates:
        try:
            r = subprocess.run(
                [candidate, "-version"],
                capture_output=True,
                timeout=5,
            )
            if r.returncode == 0:
                _ffmpeg_cache = True
                _ffmpeg_exe   = candidate
                return True
        except Exception:
            continue

    _ffmpeg_cache = False
    _ffmpeg_exe   = None
    return False


def _ffmpeg_cmd() -> str:
    """Возвращает команду/путь для запуска ffmpeg."""
    _ffmpeg_available()   # заполняет _ffmpeg_exe
    return _ffmpeg_exe or "ffmpeg"


# =============================================================================
# База данных (JSON-индекс)
# =============================================================================
def _load_db() -> Dict:
    _ensure_dirs()
    if os.path.exists(_DB_PATH):
        try:
            with open(_DB_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {"protocols": [], "stats": {"total": 0, "last_created": None}}


def _save_db(db: Dict):
    _ensure_dirs()
    with _DB_LOCK:
        with open(_DB_PATH, "w", encoding="utf-8") as f:
            json.dump(db, f, ensure_ascii=False, indent=2)


def _db_add(
    protocol_text: str,
    meeting_name: str,
    organization: str,
    meeting_date: str,
    meeting_time: str,
    attendees: List[Dict],
    source_type: str,
    source_filename: str,
    structure: str,
    detail_level: str,
) -> Dict:
    """Сохраняет DOCX + TXT на диск и добавляет запись в JSON-индекс."""
    _ensure_dirs()
    db = _load_db()

    proto_id = f"proto_{len(db['protocols']) + 1:04d}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

    txt_path  = os.path.join(_PROTOCOLS_DIR, f"{proto_id}.txt")
    docx_path = os.path.join(_PROTOCOLS_DIR, f"{proto_id}.docx")

    # TXT
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(protocol_text)

    # DOCX
    _write_docx(
        protocol_text, meeting_name, organization,
        meeting_date, meeting_time, attendees, docx_path,
    )

    record = {
        "id":              proto_id,
        "meeting_name":    meeting_name,
        "organization":    organization,
        "meeting_date":    meeting_date,
        "meeting_time":    meeting_time,
        "attendees":       attendees,
        "source_type":     source_type,
        "source_filename": source_filename,
        "structure":       structure,
        "detail_level":    detail_level,
        "created_at":      datetime.now().isoformat(),
        "txt_path":        txt_path,
        "docx_path":       docx_path,
        "txt_size":        os.path.getsize(txt_path),
        "docx_size":       os.path.getsize(docx_path) if os.path.exists(docx_path) else 0,
        "word_count":      len(protocol_text.split()),
    }

    db["protocols"].append(record)
    db["stats"]["total"] = len(db["protocols"])
    db["stats"]["last_created"] = record["created_at"]
    _save_db(db)

    return record


def _db_delete(proto_id: str):
    db = _load_db()
    rec = next((p for p in db["protocols"] if p["id"] == proto_id), None)
    if rec:
        for key in ("txt_path", "docx_path"):
            path = rec.get(key, "")
            if path and os.path.exists(path):
                try:
                    os.remove(path)
                except Exception:
                    pass
        db["protocols"] = [p for p in db["protocols"] if p["id"] != proto_id]
        db["stats"]["total"] = len(db["protocols"])
        _save_db(db)


# =============================================================================
# DOCX-генератор
# =============================================================================
def _build_docx_document(
    protocol_text: str,
    meeting_name: str,
    organization: str,
    meeting_date: str,
    meeting_time: str,
    attendees: List[Dict],
):
    """Строит и возвращает объект python-docx Document (не сохраняет)."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin   = Cm(3)
    sec.right_margin  = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ── Заголовок ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ПРОТОКОЛ")
    run.bold = True
    run.font.size = Pt(16)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(meeting_name.upper())
    run2.bold = True
    run2.font.size = Pt(14)

    doc.add_paragraph()

    # ── Реквизиты ──────────────────────────────────────────────────────────
    def _bold_line(label: str, value: str):
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(value)

    _bold_line("Организация: ", organization)
    _bold_line("Дата:        ", meeting_date)
    _bold_line("Время:       ", meeting_time)

    # ── Присутствующие ─────────────────────────────────────────────────────
    if attendees:
        p_att = doc.add_paragraph()
        p_att.add_run("Присутствовали:").bold = True
        for att in attendees:
            parts = [att.get("name", ""), att.get("position", ""), att.get("org", "")]
            line = " — ".join(part for part in parts if part.strip())
            if line:
                doc.add_paragraph(f"    • {line}")

    doc.add_paragraph()
    doc.add_paragraph("─" * 60)

    # ── Тело протокола ─────────────────────────────────────────────────────
    for line in protocol_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
        elif any(stripped.startswith(f"{i}.") for i in range(1, 30)):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(13)
        elif stripped.startswith(("-", "•", "*", "–", "—")):
            doc.add_paragraph(stripped, style="List Bullet")
        else:
            doc.add_paragraph(stripped)

    doc.add_paragraph()
    doc.add_paragraph("─" * 60)
    doc.add_paragraph()

    # ── Подписи ────────────────────────────────────────────────────────────
    p_sign = doc.add_paragraph()
    p_sign.add_run("Подписи:").bold = True
    doc.add_paragraph()
    doc.add_paragraph("_________________ / _________________")
    doc.add_paragraph()
    doc.add_paragraph("_________________ / _________________")

    return doc


def _write_docx(
    protocol_text: str,
    meeting_name: str,
    organization: str,
    meeting_date: str,
    meeting_time: str,
    attendees: List[Dict],
    path: str,
):
    doc = _build_docx_document(
        protocol_text, meeting_name, organization,
        meeting_date, meeting_time, attendees,
    )
    doc.save(path)


def _docx_bytes(
    protocol_text: str,
    meeting_name: str,
    organization: str,
    meeting_date: str,
    meeting_time: str,
    attendees: List[Dict],
) -> bytes:
    """Возвращает DOCX как bytes (для download_button)."""
    doc = _build_docx_document(
        protocol_text, meeting_name, organization,
        meeting_date, meeting_time, attendees,
    )
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# =============================================================================
# Транскрибация аудио (faster-whisper)
# =============================================================================
def transcribe_audio(
    file_bytes: bytes,
    filename: str,
    language: str = "ru",
    _status_placeholder=None,
    whisper_model: str = "medium",
) -> Dict:
    """
    Транскрибирует аудио через faster-whisper (CTranslate2).
    Работает на GPU (CUDA) без конфликтов с PyTorch.
    """
    result: Dict = {"status": "error", "text": None, "duration": None, "error": None}

    def _upd(msg: str):
        try:
            import sys
            if hasattr(sys.stdout, "buffer"):
                sys.stdout.buffer.write(f"[WHISPER] {msg}\n".encode("utf-8", errors="replace"))
                sys.stdout.buffer.flush()
        except Exception:
            pass
        if _status_placeholder is not None:
            try:
                _status_placeholder.info(msg)
            except Exception:
                pass

    ext = os.path.splitext(filename.lower())[1]
    _upd(f"Файл: {filename} ({len(file_bytes)/1024/1024:.1f} МБ), формат: {ext.upper()}")

    if ext not in [f".{e}" for e in SUPPORTED_AUDIO_FFMPEG]:
        result["error"] = f"Формат {ext.upper()} не поддерживается."
        return result

    if ext != ".wav" and not _ffmpeg_available():
        result["error"] = (
            f"Для файлов {ext.upper()} требуется ffmpeg.\n"
            "Скачайте: https://ffmpeg.org/download.html\n"
            "Распакуйте и добавьте bin\\ в PATH Windows, затем перезапустите приложение.\n\n"
            "Альтернатива — конвертируйте в WAV: https://cloudconvert.com"
        )
        return result

    try:
        from faster_whisper import WhisperModel  # noqa: F401
    except ImportError:
        result["error"] = "Установите: pip install faster-whisper"
        return result

    # Прописываем путь к ffmpeg в PATH процесса
    _exe = _ffmpeg_cmd()
    if os.path.isabs(_exe) and os.path.exists(_exe):
        _bin_dir = os.path.dirname(_exe)
        if _bin_dir not in os.environ.get("PATH", ""):
            os.environ["PATH"] = _bin_dir + os.pathsep + os.environ.get("PATH", "")
        _upd(f"ffmpeg: {_exe}")

    # Определяем устройство
    try:
        import torch
        _device   = "cuda" if torch.cuda.is_available() else "cpu"
        _devname  = torch.cuda.get_device_name(0) if _device == "cuda" else "CPU"
        _compute  = "int8_float16" if _device == "cuda" else "int8"
        _upd(f"Устройство: {_devname} ({_device.upper()}), compute={_compute}")
    except Exception:
        _device, _devname, _compute = "cpu", "CPU", "int8"
        _upd("torch не найден — используется CPU")

    _ensure_dirs()
    temp_path = os.path.join(_TEMP_DIR, f"audio_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}")
    try:
        with open(temp_path, "wb") as f:
            f.write(file_bytes)

        from faster_whisper import WhisperModel

        _upd(f"Загрузка faster-whisper {whisper_model} на {_devname}...")
        t0 = datetime.now()
        model = WhisperModel(whisper_model, device=_device, compute_type=_compute)
        load_sec = (datetime.now() - t0).seconds
        _upd(f"Модель загружена за {load_sec} сек — транскрибирую...")

        t1 = datetime.now()
        segments_gen, info = model.transcribe(
            temp_path,
            language=language,
            beam_size=5,
            vad_filter=True,         # пропускает тишину — быстрее
            vad_parameters={"min_silence_duration_ms": 500},
        )

        # Собираем текст из генератора сегментов
        full_text = []
        last_end  = 0.0
        for seg in segments_gen:
            full_text.append(seg.text.strip())
            last_end = seg.end

        trans_sec = (datetime.now() - t1).seconds

        result["text"]     = " ".join(full_text).strip()
        result["status"]   = "success"
        result["duration"] = f"{int(last_end // 60)}:{int(last_end % 60):02d}"

        _upd(
            f"Готово! Транскрибация: {trans_sec} сек · "
            f"Аудио: {result['duration']} · "
            f"Символов: {len(result['text'])}"
        )

    except Exception as exc:
        _upd(f"Ошибка: {type(exc).__name__}: {exc}")
        result["error"] = f"Ошибка транскрибации: {type(exc).__name__}: {exc}"
    finally:
        if os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass

    return result

# =============================================================================
# Извлечение текста из документов
# =============================================================================
def extract_doc_text(file_bytes: bytes, filename: str) -> str:
    ext = os.path.splitext(filename.lower())[1]

    if ext == ".txt":
        for enc in ("utf-8", "cp1251", "latin-1"):
            try:
                return file_bytes.decode(enc)
            except Exception:
                pass
        return file_bytes.decode("utf-8", errors="replace")

    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(c.text.strip() for c in row.cells))
            return "\n".join(parts)
        except Exception as exc:
            return f"[Ошибка DOCX: {exc}]"

    if ext == ".doc":
        try:
            import mammoth
            res = mammoth.extract_raw_text(io.BytesIO(file_bytes))
            return res.value
        except Exception as exc:
            return f"[Ошибка DOC: {exc}]"

    if ext == ".pdf":
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text = "\n".join(page.get_text() for page in doc)
            doc.close()
            return text
        except Exception as exc:
            return f"[Ошибка PDF: {exc}]"

    return f"[Формат {ext} не поддерживается]"


# =============================================================================
# Генерация протокола (LM Studio — OpenAI-совместимый API)
# =============================================================================

# Конфиг LM Studio — читается из того же файла что и advisor
_ADVISOR_CONFIG_FILE = os.path.join("config", "advisor_config.json")
_LM_STUDIO_DEFAULTS  = {
    "lm_studio_url": "http://127.0.0.1:1234/v1",
    "default_model": "qwen/qwen3.5-9b",
    "max_tokens":    2048,
    "timeout_seconds": 120,
}

def _load_lm_config() -> Dict:
    if os.path.exists(_ADVISOR_CONFIG_FILE):
        try:
            with open(_ADVISOR_CONFIG_FILE, "r", encoding="utf-8") as f:
                return {**_LM_STUDIO_DEFAULTS, **json.load(f)}
        except Exception:
            pass
    return dict(_LM_STUDIO_DEFAULTS)


def generate_protocol(
    text: str,
    structure: str,
    meeting_name: str,
    organization: str,
    meeting_date: str,
    meeting_time: str,
    attendees: List[Dict],
    detail_level: str = "средний",
    model: Optional[str] = None,
) -> Dict:
    result: Dict = {"status": "error", "protocol": None, "error": None}

    # ── Присутствующие → строка ──────────────────────────────────────────────
    attendees_block = ""
    if attendees:
        lines = []
        for a in attendees:
            parts = [a.get("name", ""), a.get("position", ""), a.get("org", "")]
            line  = " — ".join(p for p in parts if p.strip())
            if line:
                lines.append(f"  • {line}")
        if lines:
            attendees_block = "Присутствовали:\n" + "\n".join(lines)

    if not structure.strip():
        structure = DEFAULT_STRUCTURE

    prompt = f"""Ты — профессиональный секретарь. Составь официальный протокол встречи на русском языке.

РЕКВИЗИТЫ:
Название:     {meeting_name}
Организация:  {organization}
Дата:         {meeting_date}
Время:        {meeting_time}
{attendees_block}

СТРУКТУРА ПРОТОКОЛА (следуй строго, раздел за разделом):
{structure}

УРОВЕНЬ ДЕТАЛИЗАЦИИ: {detail_level}
{DETAIL_CAPTIONS.get(detail_level, "")}

ТЕКСТ / РАСШИФРОВКА ВСТРЕЧИ:
{text[:14000]}

ТРЕБОВАНИЯ:
- Строго следуй указанной структуре
- Выдели ключевые решения и поручения отдельно
- Укажи ответственных и сроки
- Официально-деловой стиль, без лишних слов
- Если информации нет — пиши «Не указано»
- Не добавляй советов и рекомендаций

ПРОТОКОЛ:"""

    try:
        from openai import OpenAI
        cfg     = _load_lm_config()
        _model  = model or cfg["default_model"]
        _client = OpenAI(base_url=cfg["lm_studio_url"], api_key="lm-studio")

        # Qwen3 — отключаем thinking
        is_qwen3   = "qwen3" in _model.lower()
        extra_body = {"enable_thinking": False} if is_qwen3 else {}
        if is_qwen3:
            prompt = "/no_think\n\n" + prompt

        kwargs: Dict = dict(
            model=_model,
            messages=[
                {"role": "system", "content": "Ты профессиональный секретарь. Отвечай только на русском языке."},
                {"role": "user",   "content": prompt},
            ],
            temperature=0.15,
            max_tokens=DETAIL_MAX_TOKENS.get(detail_level, 2500),
            timeout=cfg["timeout_seconds"],
        )
        if extra_body:
            kwargs["extra_body"] = extra_body

        print(f"[PROTOCOL] LM Studio: {_model}", flush=True)
        response = _client.chat.completions.create(**kwargs)
        raw = response.choices[0].message.content or ""

        # Убираем thinking-блоки если есть
        import re as _re
        protocol = _re.sub(r"<think>.*?</think>", "", raw, flags=_re.DOTALL).strip()

        if not protocol:
            result["error"] = "❌ Модель вернула пустой ответ"
        else:
            result["protocol"] = protocol
            result["status"]   = "success"

    except Exception as exc:
        err = str(exc)
        if "Connection" in err or "refused" in err:
            result["error"] = "🔌 Нет подключения к LM Studio. Убедитесь что сервер запущен на 127.0.0.1:1234"
        elif "404" in err:
            result["error"] = f"❌ Модель не найдена в LM Studio: {_model}. Загрузите её в LM Studio."
        else:
            result["error"] = f"❌ {type(exc).__name__}: {exc}"

    return result


# =============================================================================
# Виджет «Присутствующие»
# =============================================================================
def _init_attendees():
    if "proto_attendees" not in st.session_state:
        st.session_state.proto_attendees = [{"name": "", "position": "", "org": ""}]


def _render_attendees_widget() -> List[Dict]:
    """
    Отображает динамический список участников.
    Возвращает список непустых строк.
    """
    _init_attendees()
    rows = st.session_state.proto_attendees
    to_delete: Optional[int] = None

    # Шапка колонок
    hc1, hc2, hc3, _ = st.columns([3, 3, 3, 0.6])
    hc1.caption("ФИО")
    hc2.caption("Должность")
    hc3.caption("Организация")

    for i, row in enumerate(rows):
        c1, c2, c3, c4 = st.columns([3, 3, 3, 0.6])
        with c1:
            rows[i]["name"] = st.text_input(
                f"name_{i}", value=row["name"],
                key=f"att_name_{i}", placeholder="Иванов И.И.",
                label_visibility="collapsed",
            )
        with c2:
            rows[i]["position"] = st.text_input(
                f"pos_{i}", value=row["position"],
                key=f"att_pos_{i}", placeholder="Директор",
                label_visibility="collapsed",
            )
        with c3:
            rows[i]["org"] = st.text_input(
                f"org_{i}", value=row["org"],
                key=f"att_org_{i}", placeholder="ООО «Компания»",
                label_visibility="collapsed",
            )
        with c4:
            if st.button("✕", key=f"att_del_{i}", help="Удалить участника"):
                to_delete = i

    if to_delete is not None and len(rows) > 1:
        st.session_state.proto_attendees.pop(to_delete)
        st.rerun()

    if st.button("＋ Добавить участника", key="att_add_btn"):
        st.session_state.proto_attendees.append({"name": "", "position": "", "org": ""})
        st.rerun()

    return [
        r for r in rows
        if any(r.get(k, "").strip() for k in ("name", "position", "org"))
    ]


# =============================================================================
# Главная функция Streamlit
# =============================================================================
def show_protocol_bot():
    st.header("Робот-протокольщик")

    # ── Инициализация session_state ──────────────────────────────────────────
    for key, default in [
        ("proto_result",    None),   # Dict текущего созданного протокола
        ("proto_saved_id",  None),   # ID сохранённого протокола
        ("proto_transcript", None),  # Текст расшифровки аудио
        ("pb_open_card",    None),   # ID раскрытой карточки в базе
    ]:
        if key not in st.session_state:
            st.session_state[key] = default

    tab_create, tab_search, tab_base = st.tabs(["✏️ Создание протокола", "🔍 Поиск по протоколам", "📚 База протоколов"])

    # =========================================================================
    # ВКЛАДКА 1 — СОЗДАНИЕ ПРОТОКОЛА
    # =========================================================================
    with tab_create:

        # ── Шаг 1: Контекст встречи ──────────────────────────────────────────
        st.subheader("1. Контекст встречи")

        c1, c2 = st.columns(2)
        with c1:
            meeting_name = st.text_input(
                "Название встречи",
                value=st.session_state.get("_proto_name", "Совещание"),
                key="_proto_name",
                placeholder="Совещание",
            )
            organization = st.text_input(
                "Организация-инициатор",
                value=st.session_state.get("_proto_org", ""),
                key="_proto_org",
                placeholder="Ваша организация",
            )
        with c2:
            meeting_date = st.date_input(
                "Дата",
                value=datetime.now().date(),
                key="_proto_date",
            )
            meeting_time = st.time_input(
                "Время",
                value=datetime.now().time().replace(second=0, microsecond=0),
                key="_proto_time",
            )

        st.markdown("**Присутствующие**")
        attendees = _render_attendees_widget()

        st.divider()

        # ── Шаг 2: Источник данных ───────────────────────────────────────────
        st.subheader("2. Источник данных")

        input_method = st.radio(
            "Формат",
            ["✏️ Текст вручную", "📄 Документ", "🎤 Аудио"],
            horizontal=True,
            key="proto_input_method",
        )

        source_text     = ""
        source_filename = ""

        # ── Вариант A: Текст вручную ─────────────────────────────────────────
        if input_method == "✏️ Текст вручную":
            source_text = st.text_area(
                "Введите заметки, расшифровку или тезисы встречи",
                height=280,
                key="proto_manual_text",
                placeholder=(
                    "Обсуждали тариф на 2025 год.\n"
                    "Иванов предложил снизить ставку на 5%.\n"
                    "Решение: согласовать с ФАС до 15 апреля.\n"
                    "Ответственный: Петров."
                ),
            )

        # ── Вариант B: Документ ──────────────────────────────────────────────
        elif input_method == "📄 Документ":
            st.caption("Поддерживаются: TXT, DOCX, DOC, PDF")
            uploaded_doc = st.file_uploader(
                "Загрузите документ",
                type=SUPPORTED_DOCS,
                key="proto_doc_upload",
            )
            if uploaded_doc:
                source_filename = uploaded_doc.name
                raw_bytes = uploaded_doc.read()
                extracted = extract_doc_text(raw_bytes, uploaded_doc.name)
                st.success(f"✅ {uploaded_doc.name} — {len(extracted):,} символов")
                source_text = st.text_area(
                    "Текст (можно отредактировать перед генерацией)",
                    value=extracted,
                    height=240,
                    key="proto_doc_text_edit",
                )

        # ── Вариант C: Аудио ─────────────────────────────────────────────────
        else:
            has_ffmpeg = _ffmpeg_available()

            if has_ffmpeg:
                allowed_ext = SUPPORTED_AUDIO_FFMPEG
                st.caption(
                    "Поддерживаются: **M4A** (iPhone), **MP3**, **WAV**, **OGG**, **AAC**, **3GP**, MP4, WMA"
                )
            else:
                allowed_ext = SUPPORTED_AUDIO_NOFFMPEG
                st.warning(
                    "**ffmpeg не найден** — поддерживается только WAV.  \n"
                    "Для iPhone (M4A) и других форматов:  \n"
                    "1. Скачайте ffmpeg: https://ffmpeg.org/download.html  \n"
                    "2. Распакуйте и добавьте папку `bin\\` в системный PATH  \n"
                    "3. Перезапустите приложение  \n\n"
                    "💡 Или конвертируйте файл в WAV: https://cloudconvert.com"
                )

            uploaded_audio = st.file_uploader(
                "Загрузите аудиозапись",
                type=allowed_ext,
                key="proto_audio_upload",
            )

            if uploaded_audio:
                source_filename = uploaded_audio.name
                st.audio(uploaded_audio)

                # Выбор модели
                _MODELS = {
                    "tiny":     ("⚡ Быстрый",      "tiny",     "~10 сек на 10 мин · низкое качество"),
                    "base":     ("🚀 Быстрый+",     "base",     "~20 сек на 10 мин · приемлемое качество"),
                    "small":    ("⚖️ Средний",      "small",    "~40 сек на 10 мин · хорошее качество"),
                    "medium":   ("✅ Рекомендуемый", "medium",   "~1.5 мин на 10 мин · отличное качество"),
                    "large-v3": ("🏆 Максимальный", "large-v3", "~5+ мин на 10 мин · требует много RAM"),
                }
                _model_labels = {k: f"{v[0]} — {v[2]}" for k, v in _MODELS.items()}
                _selected_key = st.selectbox(
                    "Модель Whisper",
                    options=list(_MODELS.keys()),
                    index=3,   # medium по умолчанию
                    format_func=lambda k: _model_labels[k],
                    key="proto_whisper_model",
                )
                _selected_model = _MODELS[_selected_key][1]

                col_btn, col_hint = st.columns([1, 3])
                with col_btn:
                    do_transcribe = st.button(
                        "🎤 Транскрибировать",
                        type="secondary",
                        key="proto_transcribe_btn",
                        use_container_width=True,
                    )
                with col_hint:
                    import torch as _torch
                    _gpu_ok   = _torch.cuda.is_available()
                    _gpu_name = _torch.cuda.get_device_name(0) if _gpu_ok else "CPU"
                    _icon = "🟢 GPU: " + _gpu_name if _gpu_ok else "🔴 CPU"
                    st.caption(f"Whisper `{_selected_key}` · {_icon} · {_MODELS[_selected_key][2]}")


                if do_transcribe:
                    import time as _time
                    audio_bytes = uploaded_audio.read()
                    _mb = len(audio_bytes) / 1024 / 1024

                    _status = st.empty()
                    _status.info(f"⏳ Подготовка... файл {_mb:.1f} МБ, модель: {_selected_key}")

                    _start = _time.time()
                    tr = transcribe_audio(
                        audio_bytes, uploaded_audio.name,
                        _status_placeholder=_status,
                        whisper_model=_selected_model,
                    )
                    _elapsed = int(_time.time() - _start)

                    _status.empty()
                    if tr["status"] == "success":
                        st.session_state.proto_transcript = tr["text"]
                        _dur = tr.get("duration", "")
                        st.success(
                            f"✅ Готово за **{_elapsed} сек**"
                            + (f" · аудио: **{_dur}**" if _dur else "")
                            + f" · символов: **{len(tr['text']):,}**"
                        )
                        st.rerun()
                    else:
                        st.error(tr["error"])



            # Показываем / редактируем расшифровку
            if st.session_state.proto_transcript:
                source_text = st.text_area(
                    "Расшифровка (можно отредактировать)",
                    value=st.session_state.proto_transcript,
                    height=240,
                    key="proto_transcript_edit",
                )

        st.divider()

        # ── Шаг 3: Структура и настройки ────────────────────────────────────
        st.subheader("3. Структура протокола")

        col_struct, col_settings = st.columns([3, 2])

        with col_struct:
            structure_text = st.text_area(
                "Разделы протокола (AI следует им строго)",
                value=DEFAULT_STRUCTURE,
                height=220,
                key="proto_structure",
                help="Перечислите нужные разделы — по одному в строке",
            )

        with col_settings:
            st.markdown("**Детализация**")
            detail_level = st.radio(
                "detail",
                list(DETAIL_CAPTIONS.keys()),
                index=1,
                key="proto_detail",
                label_visibility="collapsed",
            )
            st.caption(DETAIL_CAPTIONS[detail_level])

            st.markdown("**Макс. длина (токены)**")
            max_tokens = st.slider(
                "tokens",
                500, 5000,
                DETAIL_MAX_TOKENS[detail_level],
                250,
                key="proto_tokens",
                label_visibility="collapsed",
            )

        st.divider()

        # ── Шаг 4: Кнопка генерации ─────────────────────────────────────────
        st.subheader("4. Создать протокол")

        col_gen, col_status = st.columns([1, 3])
        with col_gen:
            gen_btn = st.button(
                "🤖 Создать протокол",
                type="primary",
                use_container_width=True,
                key="proto_gen_btn",
            )
        with col_status:
            if source_text.strip():
                st.caption(f"✅ Источник: {len(source_text):,} символов")
            else:
                st.caption("⬆️ Введите текст или загрузите файл / аудио")

        if gen_btn:
            if not source_text.strip():
                st.warning("⚠️ Введите текст или загрузите файл / аудио")
            else:
                with st.spinner("🔄 AI составляет протокол..."):
                    res = generate_protocol(
                        text=source_text,
                        structure=structure_text,
                        meeting_name=meeting_name,
                        organization=organization or "Ваша организация",
                        meeting_date=str(meeting_date),
                        meeting_time=str(meeting_time)[:5],
                        attendees=attendees,
                        detail_level=detail_level,
                    )

                if res["status"] == "success":
                    st.session_state.proto_result = {
                        "text":            res["protocol"],
                        "meeting_name":    meeting_name,
                        "organization":    organization or "Ваша организация",
                        "meeting_date":    str(meeting_date),
                        "meeting_time":    str(meeting_time)[:5],
                        "attendees":       attendees,
                        "structure":       structure_text,
                        "detail_level":    detail_level,
                        "source_type":     input_method.split()[0],
                        "source_filename": source_filename,
                    }
                    st.session_state.proto_saved_id = None
                    st.success("✅ Протокол готов! Прокрутите вниз.")
                    st.rerun()
                else:
                    st.error(res["error"])

        # ── Шаг 5: Результат ─────────────────────────────────────────────────
        if st.session_state.proto_result:
            proto = st.session_state.proto_result
            st.divider()
            st.subheader("5. Готовый протокол")

            edited = st.text_area(
                "Текст протокола (можно отредактировать перед сохранением)",
                value=proto["text"],
                height=500,
                key="proto_edit_area",
            )
            proto["text"] = edited   # обновляем в session_state

            c_save, c_dl, c_reset = st.columns(3)

            with c_save:
                if st.session_state.proto_saved_id:
                    st.success(f"✅ Сохранён: {st.session_state.proto_saved_id[:20]}…")
                else:
                    if st.button("💾 Сохранить в базу", use_container_width=True):
                        record = _db_add(
                            protocol_text  = edited,
                            meeting_name   = proto["meeting_name"],
                            organization   = proto["organization"],
                            meeting_date   = proto["meeting_date"],
                            meeting_time   = proto["meeting_time"],
                            attendees      = proto["attendees"],
                            source_type    = proto["source_type"],
                            source_filename= proto["source_filename"],
                            structure      = proto["structure"],
                            detail_level   = proto["detail_level"],
                        )
                        st.session_state.proto_saved_id = record["id"]
                        st.rerun()

            with c_dl:
                try:
                    dl_bytes = _docx_bytes(
                        edited,
                        proto["meeting_name"],
                        proto["organization"],
                        proto["meeting_date"],
                        proto["meeting_time"],
                        proto["attendees"],
                    )
                except Exception:
                    dl_bytes = b""

                safe_name = "".join(
                    c for c in proto["meeting_name"][:25]
                    if c.isalnum() or c in (" ", "_", "-")
                ).strip().replace(" ", "_")
                fname = f"Protocol_{safe_name}_{proto['meeting_date'].replace('-','')}.docx"

                st.download_button(
                    "📥 Скачать DOCX",
                    data=dl_bytes,
                    file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="proto_dl_btn",
                    disabled=not bool(dl_bytes),
                )

            with c_reset:
                if st.button("🔄 Создать новый", use_container_width=True):
                    st.session_state.proto_result    = None
                    st.session_state.proto_transcript = None
                    st.session_state.proto_saved_id  = None
                    st.rerun()

            # Предпросмотр (форматированный)
            with st.expander("👁️ Предпросмотр форматированного текста"):
                st.markdown(edited.replace("\n", "  \n"))

    # =========================================================================
    # =========================================================================
    # ВКЛАДКА 2 — ПОИСК ПО ПРОТОКОЛАМ
    # =========================================================================
    with tab_search:
        st.subheader("Поиск по содержимому протоколов")

        db_for_search = _load_db()
        protos_for_search = db_for_search.get("protocols", [])

        if not protos_for_search:
            st.info("База протоколов пуста.")
        else:
            # ── Параметры поиска ─────────────────────────────────────────────
            _sq = st.text_input(
                "Поисковый запрос",
                placeholder="Введите слово, фразу или имя участника...",
                key="psearch_query",
            )

            so1, so2, so3 = st.columns(3)
            with so1:
                _match_type = st.radio(
                    "Тип совпадения",
                    ["Точное", "По словам", "Нечёткое"],
                    key="psearch_match",
                    help="Точное — ищет фразу целиком. По словам — все слова должны присутствовать. Нечёткое — находит похожие слова.",
                )
            with so2:
                _scope = st.radio(
                    "Где искать",
                    ["Текст протокола", "Реквизиты", "Везде"],
                    key="psearch_scope",
                )
            with so3:
                _presence = st.radio(
                    "Наличие",
                    ["Содержит", "Не содержит"],
                    key="psearch_presence",
                    help="Найти протоколы где запрос присутствует или отсутствует",
                )

            _do_search = st.button("Найти", key="psearch_btn", type="primary")

            st.divider()

            def _proto_matches(proto: dict, query: str, match_type: str, scope: str) -> list:
                """Возвращает список совпадений [(контекст, источник)] или [] если не найдено."""
                import re as _re
                q = query.strip()
                if not q:
                    return []

                def _search_in(text: str, src: str) -> list:
                    hits = []
                    if match_type == "Точное":
                        idx = 0
                        while True:
                            pos = text.lower().find(q.lower(), idx)
                            if pos < 0:
                                break
                            start = max(0, pos - 80)
                            end   = min(len(text), pos + len(q) + 80)
                            snippet = ("..." if start > 0 else "") + text[start:end] + ("..." if end < len(text) else "")
                            # Подсветка
                            hi = _re.sub(f"(?i)({_re.escape(q)})", r"**\1**", snippet)
                            hits.append((hi, src))
                            idx = pos + 1
                            if len(hits) >= 3:
                                break
                    elif match_type == "По словам":
                        words = q.lower().split()
                        tl = text.lower()
                        if all(w in tl for w in words):
                            # Находим позицию первого слова
                            pos = tl.find(words[0])
                            start = max(0, pos - 60)
                            end   = min(len(text), pos + 120)
                            snippet = ("..." if start > 0 else "") + text[start:end] + "..."
                            pattern = "|".join(_re.escape(w) for w in words)
                            hi = _re.sub(f"(?i)({pattern})", r"**\1**", snippet)
                            hits.append((hi, src))
                    elif match_type == "Нечёткое":
                        # Простое нечёткое: ищем все слова запроса с учётом начала слова
                        words = q.lower().split()
                        tl = text.lower()
                        found_words = []
                        for w in words:
                            # Ищем слова начинающиеся на первые 4 символа запрашиваемого слова
                            stem = w[:max(4, len(w)-2)]
                            if stem in tl:
                                found_words.append(w)
                        if len(found_words) >= max(1, len(words) // 2):
                            pos = tl.find(found_words[0][:4])
                            start = max(0, pos - 60)
                            end   = min(len(text), pos + 150)
                            snippet = ("..." if start > 0 else "") + text[start:end] + "..."
                            pattern = "|".join(_re.escape(w[:4]) for w in found_words)
                            hi = _re.sub(f"(?i)({pattern})", r"**\1**", snippet)
                            hits.append((hi, src))
                    return hits

                all_hits = []

                # Реквизиты
                if scope in ("Реквизиты", "Везде"):
                    meta_text = " ".join([
                        proto.get("meeting_name", ""),
                        proto.get("organization", ""),
                        " ".join(a.get("name","") for a in proto.get("attendees", [])),
                    ])
                    all_hits += _search_in(meta_text, "реквизиты")

                # Текст протокола
                if scope in ("Текст протокола", "Везде"):
                    txt_path = proto.get("txt_path", "")
                    if txt_path and os.path.exists(txt_path):
                        try:
                            with open(txt_path, "r", encoding="utf-8") as _f:
                                txt = _f.read()
                            all_hits += _search_in(txt, "текст протокола")
                        except Exception:
                            pass

                return all_hits

            if _do_search or st.session_state.get("psearch_last_query"):
                if _do_search:
                    st.session_state.psearch_last_query = _sq

                query_to_use = st.session_state.get("psearch_last_query", _sq)
                want_match   = _presence == "Содержит"

                if query_to_use.strip():
                    _results = []
                    for p in protos_for_search:
                        hits = _proto_matches(p, query_to_use, _match_type, _scope)
                        has_hits = len(hits) > 0
                        if has_hits == want_match:
                            _results.append((p, hits))

                    _icon = "+" if want_match else "−"
                    st.caption(
                        f"{_icon} Найдено протоколов: **{len(_results)}** "
                        f"из {len(protos_for_search)} "
                        f"· запрос: «{query_to_use}» "
                        f"· {_match_type} · {_scope}"
                    )

                    if not _results:
                        st.warning("Ничего не найдено. Попробуйте изменить тип совпадения.")
                    else:
                        for p, hits in _results:
                            att_names = ", ".join(
                                a.get("name","") for a in p.get("attendees",[]) if a.get("name")
                            ) or "—"
                            with st.expander(
                                f"{'📄' if hits else '○'} {p.get('meeting_name','—')} · "
                                f"{p.get('meeting_date','')[:10]} · "
                                f"{p.get('organization','—')} · "
                                f"{len(hits)} совп.",
                                expanded=(len(_results) == 1),
                            ):
                                st.caption(f"Участники: {att_names}")
                                if hits:
                                    for ctx, src in hits[:5]:
                                        st.markdown(
                                            f"<div style='background:#f4f6f9;border-left:3px solid #1B5C74;"
                                            f"padding:6px 10px;margin:4px 0;border-radius:0 4px 4px 0;"
                                            f"font-size:0.85rem;'>{ctx}</div>",
                                            unsafe_allow_html=True,
                                        )
                                        st.caption(f"Источник: {src}")
                                elif not want_match:
                                    st.caption("Запрос не найден в этом протоколе")

                                # Кнопки действий
                                docx_ok = bool(p.get("docx_path") and os.path.exists(p.get("docx_path","")))
                                ca, cb = st.columns(2)
                                with ca:
                                    if docx_ok:
                                        with open(p["docx_path"], "rb") as _df:
                                            _db = _df.read()
                                        _sn = "".join(c for c in p.get("meeting_name","")[:20] if c.isalnum() or c in " _").strip().replace(" ","_")
                                        st.download_button(
                                            "📥 Скачать DOCX",
                                            data=_db,
                                            file_name=f"Protocol_{_sn}_{p.get('meeting_date','').replace('-','')}.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            key=f"psearch_dl_{p['id']}",
                                            use_container_width=True,
                                        )
                                with cb:
                                    if st.button("✏️ Редактировать", key=f"psearch_edit_{p['id']}", use_container_width=True):
                                        st.session_state.pb_open_card  = p["id"]
                                        st.session_state[f"pb_editing_{p['id']}"] = True
                                        st.info("Перейдите на вкладку «База протоколов»")


    # ВКЛАДКА 2 — БАЗА ПРОТОКОЛОВ
    # =========================================================================
    with tab_base:
        db        = _load_db()
        all_protos = list(reversed(db.get("protocols", [])))
        stats      = db.get("stats", {})

        st.subheader("База протоколов")

        # ── Метрики ──────────────────────────────────────────────────────────
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Всего протоколов", stats.get("total", 0))
        m2.metric(
            "Слов (итого)",
            f"{sum(p.get('word_count', 0) for p in all_protos):,}",
        )
        _total_sz = sum(
            (p.get("txt_size") or 0) + (p.get("docx_size") or 0)
            for p in all_protos
        )
        m3.metric("На диске", _fmt_size(_total_sz))
        m4.metric("Последний", (stats.get("last_created") or "—")[:10])

        if not all_protos:
            st.info("📭 База пуста. Создайте первый протокол на вкладке «Создание протокола».")
            return

        st.divider()

        # ── Фильтры ──────────────────────────────────────────────────────────
        ff1, ff2, ff3 = st.columns([3, 2, 2])
        with ff1:
            _search = st.text_input(
                "Поиск",
                placeholder="По названию встречи или организации...",
                key="pb_search",
            )
        with ff2:
            _detail_filter = st.multiselect(
                "Детализация",
                ["краткий", "средний", "подробный"],
                key="pb_detail_filter",
                placeholder="Все",
            )
        with ff3:
            _sort = st.selectbox(
                "Сортировка",
                ["По дате (новые)", "По дате (старые)", "По названию А-Я"],
                key="pb_sort",
            )

        # ── Фильтр по диапазону дат ───────────────────────────────────────────
        from datetime import date as _date, timedelta as _td
        fd1, fd2, fd3 = st.columns([2, 2, 1])
        with fd1:
            # Минимальная и максимальная даты в базе
            _all_dates = []
            for _p in all_protos:
                try:
                    _all_dates.append(_date.fromisoformat(_p.get("meeting_date", "")[:10]))
                except Exception:
                    pass
            _min_date = min(_all_dates) if _all_dates else _date.today() - _td(days=365)
            _max_date = max(_all_dates) if _all_dates else _date.today()
            _date_from = st.date_input(
                "Дата с",
                value=st.session_state.get("pb_date_from", _min_date),
                min_value=_min_date,
                max_value=_max_date,
                key="pb_date_from",
            )
        with fd2:
            _date_to = st.date_input(
                "Дата по",
                value=st.session_state.get("pb_date_to", _max_date),
                min_value=_min_date,
                max_value=_max_date,
                key="pb_date_to",
            )
        with fd3:
            st.markdown("<div style='margin-top:28px'></div>", unsafe_allow_html=True)
            if st.button("Сбросить даты", key="pb_date_reset", use_container_width=True):
                st.session_state.pb_date_from = _min_date
                st.session_state.pb_date_to   = _max_date
                st.rerun()

        # ── Применяем фильтры ────────────────────────────────────────────────
        filtered = list(all_protos)

        # Фильтр по датам встречи
        def _proto_date(p) -> _date:
            try:
                return _date.fromisoformat(p.get("meeting_date","")[:10])
            except Exception:
                return _date(1970, 1, 1)

        filtered = [p for p in filtered if _date_from <= _proto_date(p) <= _date_to]

        if _search.strip():
            q = _search.strip().lower()
            filtered = [
                p for p in filtered
                if q in p.get("meeting_name", "").lower()
                or q in p.get("organization", "").lower()
            ]
        if _detail_filter:
            filtered = [p for p in filtered if p.get("detail_level") in _detail_filter]
        if _sort == "По дате (старые)":
            filtered = list(reversed(filtered))
        elif _sort == "По названию А-Я":
            filtered.sort(key=lambda p: p.get("meeting_name", "").lower())

        _active_hints = []
        if _search.strip():
            _active_hints.append(f"поиск: «{_search.strip()}»")
        if _detail_filter:
            _active_hints.append(f"детализация: {', '.join(_detail_filter)}")
        _date_hint = f"{_date_from.strftime('%d.%m.%Y')} — {_date_to.strftime('%d.%m.%Y')}"
        _active_hints.append(f"период: {_date_hint}")
        hint_str = "  ·  ".join(_active_hints)
        st.caption(f"Показано: {len(filtered)} из {len(all_protos)}  ·  {hint_str}")

        # ── Карточки ─────────────────────────────────────────────────────────
        open_card = st.session_state.get("pb_open_card")

        for proto in filtered:
            pid      = proto["id"]
            is_open  = open_card == pid
            docx_ok  = bool(proto.get("docx_path") and os.path.exists(proto.get("docx_path", "")))
            txt_ok   = bool(proto.get("txt_path")  and os.path.exists(proto.get("txt_path",  "")))

            # Список имён участников для превью
            att_names = ", ".join(
                a.get("name", "") for a in proto.get("attendees", []) if a.get("name")
            ) or "—"

            # Шапка карточки
            detail_badge = {
                "краткий":   "краткий",
                "средний":   "средний",
                "подробный": "подробный",
            }.get(proto.get("detail_level", ""), proto.get("detail_level", "—"))

            st.markdown(
                f"<div style='background:#ffffff;border:1px solid #dce3ec;"
                f"border-radius:6px;padding:10px 14px;margin-bottom:4px;'>"
                f"<div style='font-weight:600;font-size:0.95rem;color:#1B5C74;'>"
                f"{proto.get('meeting_name', 'Без названия')}</div>"
                f"<div style='font-size:0.78rem;color:#5a6a7a;margin-top:4px;'>"
                f"{proto.get('organization', '—')} &nbsp;·&nbsp; "
                f"{proto.get('meeting_date', '')[:10]} &nbsp;·&nbsp; "
                f"{proto.get('meeting_time', '')} &nbsp;·&nbsp; "
                f"{detail_badge} &nbsp;·&nbsp; "
                f"{proto.get('word_count', 0):,} слов"
                f"</div>"
                f"<div style='font-size:0.74rem;color:#7a8a9a;margin-top:3px;'>"
                f"Участники: {att_names[:90]}{'…' if len(att_names) > 90 else ''}"
                f"</div>"
                + (
                    "<div style='font-size:0.78rem;color:#5a6a7a;margin-top:6px;"
                    "border-top:1px solid #eef1f5;padding-top:6px;"
                    "white-space:nowrap;overflow:hidden;text-overflow:ellipsis;'>"
                    + (open(proto["txt_path"], encoding="utf-8").read()[:200].replace("<","&lt;").replace(">","&gt;") + "…"
                       if txt_ok else "<i style='color:#aaa'>текст недоступен</i>")
                    + "</div>"
                    if True else ""
                )
                + f"</div>",
                unsafe_allow_html=True,
            )

            # Кнопки-действия
            ca, cb, cc, cd, ce = st.columns([2, 2, 2, 2, 0.7])

            with ca:
                toggle_label = "Скрыть" if is_open else "Показать детали"
                if st.button(toggle_label, key=f"pb_toggle_{pid}", use_container_width=True):
                    st.session_state.pb_open_card = None if is_open else pid
                    st.rerun()

            with cb:
                if docx_ok:
                    with open(proto["docx_path"], "rb") as _f:
                        _dl_bytes = _f.read()
                    _safe = "".join(
                        c for c in proto.get("meeting_name", "protocol")[:20]
                        if c.isalnum() or c in (" ", "_")
                    ).strip().replace(" ", "_")
                    _fname = f"Protocol_{_safe}_{proto.get('meeting_date','').replace('-','')}.docx"
                    st.download_button(
                        "Скачать DOCX",
                        data=_dl_bytes,
                        file_name=_fname,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"pb_dl_{pid}",
                        use_container_width=True,
                    )
                else:
                    st.button("DOCX недоступен", key=f"pb_dl_{pid}", disabled=True, use_container_width=True)

            with cc:
                # Исходный текст встречи (источник транскрибации/ввода)
                _src_txt_path = proto.get("txt_path", "")
                if txt_ok:
                    with open(proto["txt_path"], "r", encoding="utf-8") as _tf:
                        _src_txt_bytes = _tf.read().encode("utf-8")
                    _src_safe = "".join(
                        c for c in proto.get("meeting_name","")[:20]
                        if c.isalnum() or c in (" ","_")
                    ).strip().replace(" ","_")
                    st.download_button(
                        "Скачать TXT",
                        data=_src_txt_bytes,
                        file_name=f"Protocol_{_src_safe}_{proto.get('meeting_date','').replace('-','')}.txt",
                        mime="text/plain; charset=utf-8",
                        key=f"pb_dl_txt_{pid}",
                        use_container_width=True,
                    )
                else:
                    st.button("TXT недоступен", key=f"pb_dl_txt_{pid}", disabled=True, use_container_width=True)

            with cd:
                edit_key = f"pb_editing_{pid}"
                if edit_key not in st.session_state:
                    st.session_state[edit_key] = False
                if st.button(
                    "Редактировать" if not st.session_state[edit_key] else "Закрыть редактор",
                    key=f"pb_load_{pid}",
                    use_container_width=True,
                ):
                    st.session_state[edit_key] = not st.session_state[edit_key]
                    st.session_state.pb_open_card = pid
                    st.rerun()

            with ce:
                if st.button("✕", key=f"pb_del_{pid}", help="Удалить протокол навсегда"):
                    _db_delete(pid)
                    if st.session_state.get("pb_open_card") == pid:
                        st.session_state.pb_open_card = None
                    st.rerun()

            # ── Детали карточки (раскрытая) ──────────────────────────────────
            if is_open:
                # Участники
                atts = proto.get("attendees") or []
                if atts:
                    st.markdown("**Участники:**")
                    for att in atts:
                        parts = [att.get("name",""), att.get("position",""), att.get("org","")]
                        line = " — ".join(p for p in parts if p.strip())
                        if line:
                            st.markdown(f"&nbsp;&nbsp;&nbsp;• {line}")

                edit_key = f"pb_editing_{pid}"
                is_editing = st.session_state.get(edit_key, False)

                if txt_ok:
                    with open(proto["txt_path"], "r", encoding="utf-8") as _f:
                        txt_content = _f.read()

                    if is_editing:
                        # ── Режим редактирования ──────────────────────────────
                        st.markdown("**Редактирование протокола:**")
                        edited = st.text_area(
                            label="",
                            value=txt_content,
                            height=450,
                            key=f"pb_edit_area_{pid}",
                        )

                        cs1, cs2, cs3 = st.columns(3)
                        with cs1:
                            if st.button("💾 Сохранить", key=f"pb_save_{pid}", use_container_width=True):
                                # Сохраняем TXT
                                with open(proto["txt_path"], "w", encoding="utf-8") as _f:
                                    _f.write(edited)
                                # Пересохраняем DOCX
                                _write_docx(
                                    edited,
                                    proto.get("meeting_name", ""),
                                    proto.get("organization", ""),
                                    proto.get("meeting_date", ""),
                                    proto.get("meeting_time", ""),
                                    proto.get("attendees", []),
                                    proto["docx_path"],
                                )
                                # Обновляем размеры в БД
                                db = _load_db()
                                for p in db["protocols"]:
                                    if p["id"] == pid:
                                        p["word_count"] = len(edited.split())
                                        p["txt_size"]   = os.path.getsize(proto["txt_path"])
                                        p["docx_size"]  = os.path.getsize(proto["docx_path"]) if os.path.exists(proto["docx_path"]) else 0
                                        break
                                _save_db(db)
                                st.session_state[edit_key] = False
                                st.success("✅ Сохранено")
                                st.rerun()

                        with cs2:
                            try:
                                dl_bytes = _docx_bytes(
                                    edited,
                                    proto.get("meeting_name", ""),
                                    proto.get("organization", ""),
                                    proto.get("meeting_date", ""),
                                    proto.get("meeting_time", ""),
                                    proto.get("attendees", []),
                                )
                            except Exception:
                                dl_bytes = b""
                            _safe = "".join(c for c in proto.get("meeting_name","")[:20] if c.isalnum() or c in " _").strip().replace(" ","_")
                            st.download_button(
                                "📥 Скачать DOCX",
                                data=dl_bytes,
                                file_name=f"Protocol_{_safe}_{proto.get('meeting_date','').replace('-','')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"pb_dl_edit_{pid}",
                                use_container_width=True,
                            )

                        with cs3:
                            if st.button("↩ Отмена", key=f"pb_cancel_{pid}", use_container_width=True):
                                st.session_state[edit_key] = False
                                st.rerun()
                    else:
                        # ── Режим просмотра ───────────────────────────────────
                        st.markdown("**Текст протокола:**")
                        st.text_area(
                            label="",
                            value=txt_content,
                            height=360,
                            key=f"pb_preview_txt_{pid}",
                            disabled=True,
                        )
                        meta_parts = []
                        if proto.get("source_filename"):
                            meta_parts.append(f"Источник: {proto['source_filename']}")
                        meta_parts.append(f"Создан: {proto.get('created_at','')[:16].replace('T',' ')}")
                        st.caption("  ·  ".join(meta_parts))
                else:
                    st.warning("⚠️ Файл TXT не найден.")


# =============================================================================
# Точка входа
# =============================================================================
if __name__ == "__main__":
    show_protocol_bot()