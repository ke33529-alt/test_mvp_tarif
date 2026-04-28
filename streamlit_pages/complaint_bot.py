# streamlit_pages/complaint_bot.py
import streamlit as st
import json
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import io
import os
import sys

sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# =============================================================================
# Функции
# =============================================================================

def load_json_data(uploaded_file):
    """Загружает JSON из Экспорта ФГИС или Калькулятора рисков"""
    try:
        uploaded_file.seek(0)
        data = json.load(uploaded_file)
        return data
    except Exception as e:
        st.error(f"Ошибка чтения JSON: {e}")
        return None

def get_complaint_template(addressee_type, sphere):
    """Возвращает шаблон жалобы в зависимости от адресата и сферы"""
    
    templates = {
        "fas": {
            "header": "В Федеральную антимонопольную службу России\n125993, г. Москва, Садовая-Кудринская ул., д. 7, Д-242, ГСП-3",
            "title": "ЖАЛОБА\nна решение регулятора об отказе в установлении тарифа",
            "legal_basis": [
                "Федеральный закон от 26.12.2008 N 294-ФЗ О защите прав юридических лиц",
                "Федеральный закон от 27.07.2010 N 190-ФЗ О теплоснабжении",
                "Приказ ФАС России от 30.12.2020 N 1746-э Об утверждении Методических указаний"
            ],
            "demands": [
                "Признать решение [наименование регулятора] от [дата] N [номер] незаконным",
                "Обязать [наименование регулятора] пересмотреть решение и установить тариф в размере [сумма]",
                "Рассмотреть жалобу в установленный законом срок"
            ]
        },
        "arbitration": {
            "header": "В Арбитражный суд [субъект РФ]\n[адрес суда]",
            "title": "ЗАЯВЛЕНИЕ\nо признании решения регулятора недействительным",
            "legal_basis": [
                "Арбитражный процессуальный кодекс Российской Федерации",
                "Федеральный закон от 27.12.2010 N 416-ФЗ О водоснабжении и водоотведении",
                "Постановление Пленума ВАС РФ от 30.07.2013 N 61"
            ],
            "demands": [
                "Признать решение [наименование регулятора] от [дата] N [номер] недействительным",
                "Обязать [наименование регулятора] принять новое решение об установлении тарифа",
                "Взыскать судебные расходы в размере [сумма] руб."
            ]
        },
        "court_general": {
            "header": "В [наименование суда общей юрисдикции]\n[адрес суда]",
            "title": "АДМИНИСТРАТИВНОЕ ИСКОВОЕ ЗАЯВЛЕНИЕ\nоб оспаривании решения регулятора",
            "legal_basis": [
                "Кодекс административного судопроизводства Российской Федерации",
                "Федеральный закон от 24.06.1998 N 89-ФЗ Об отходах производства и потребления",
                "Постановление Пленума ВС РФ от 27.09.2016 N 36"
            ],
            "demands": [
                "Признать решение [наименование регулятора] от [дата] N [номер] незаконным",
                "Обязать [наименование регулятора] устранить допущенные нарушения",
                "Взыскать компенсацию морального вреда (если применимо)"
            ]
        }
    }
    
    return templates.get(addressee_type, templates["fas"])

def generate_violation_arguments(violations, sphere, articles_data):
    """Генерирует аргументы по нарушениям регулятора"""
    
    arguments = []
    
    violation_templates = {
        "uncalculated_costs": {
            "title": "Неучтенные экономически обоснованные затраты",
            "text": "Регулятор необоснованно исключил из расчета тарифа затраты в размере {amount} руб. по статье {article}. Данное исключение противоречит п. 45 Приказа ФАС N 1746-э.",
            "legal_ref": "Приказ ФАС России N 1746-э, п. 45-50"
        },
        "wrong_methodology": {
            "title": "Неверная трактовка методики расчета",
            "text": "При расчете тарифа регулятор применил методику с нарушениями. Это привело к занижению тарифа на {percent}%.",
            "legal_ref": "Приказ ФАС России N 1746-э, п. 12-18"
        },
        "procedural_violations": {
            "title": "Процессуальные нарушения",
            "text": "Регулятор не соблюдал установленные сроки рассмотрения заявки, не предоставил мотивированное решение.",
            "legal_ref": "ФЗ N 294, ст. 15-17"
        },
        "wrong_comparison": {
            "title": "Некорректное сравнение с аналогами",
            "text": "При сравнении с аналогичными организациями регулятор использовал некорректные данные.",
            "legal_ref": "Приказ ФАС России N 1746-э, п. 89-95"
        },
        "missing_documents": {
            "title": "Отказ из-за отсутствия документов (необоснованный)",
            "text": "Регулятор указал на отсутствие документов, которые были представлены в составе заявки.",
            "legal_ref": "ФЗ N 294, ст. 16"
        },
        "wrong_amortization": {
            "title": "Неверный расчет амортизации",
            "text": "Регулятор неправомерно снизил норму амортизации по группе ОС {group}.",
            "legal_ref": "Постановление Правительства РФ N 1, Приказ ФАС N 1746-э п. 67-75"
        },
        "wrong_numeracy": {
            "title": "Необоснованное снижение численности",
            "text": "Требование регулятора о снижении численности персонала не основано на нормативах.",
            "legal_ref": "Приказ ФАС России N 1746-э, п. 45-50"
        }
    }
    
    for violation in violations:
        if violation in violation_templates:
            template = violation_templates[violation]
            text = template["text"]
            
            if "{amount}" in text:
                amount = sum(a.get("amount", 0) for a in articles_data[:3]) if articles_data else 1000000
                text = text.replace("{amount}", f"{amount:,.0f}")
            if "{article}" in text:
                article_name = articles_data[0].get("name", "затраты") if articles_data else "затраты"
                text = text.replace("{article}", article_name)
            if "{percent}" in text:
                text = text.replace("{percent}", "15-25")
            if "{group}" in text:
                text = text.replace("{group}", "7-10")
            
            arguments.append({
                "title": template["title"],
                "text": text,
                "legal_ref": template["legal_ref"]
            })
    
    return arguments

def generate_complaint_text(data, addressee_type, sphere, violations, additional_info):
    """Генерирует полный текст жалобы"""
    
    template = get_complaint_template(addressee_type, sphere)
    header_data = data.get("header", {})
    application_data = data.get("application", {})
    revenue_data = application_data.get("revenue", {})
    articles_data = revenue_data.get("articles", [])
    period_data = application_data.get("tariff_period", {})
    
    sphere_names = {
        "water_supply": "водоснабжение",
        "water_drainage": "водоотведение",
        "waste_management": "обращение с ТКО",
        "heat_supply": "теплоснабжение"
    }
    sphere_name = sphere_names.get(sphere, "коммунальные услуги")
    
    arguments = generate_violation_arguments(violations, sphere, articles_data)
    total_revenue = revenue_data.get("total", 0)
    
    text = f"""
{template["header"]}

От: {header_data.get("organization_name", "Организация")}
ИНН: {header_data.get("inn", "")} | КПП: {header_data.get("kpp", "")}
Адрес: {header_data.get("address", "")}
Тел.: {header_data.get("phone", "")} | Email: {header_data.get("email", "")}

{'='*80}

{template["title"]}

{'='*80}

1. ОПИСАТЕЛЬНАЯ ЧАСТЬ

{header_data.get("organization_name", "Организация")} (далее - Заявитель) осуществляет деятельность в сфере {sphere_name} на основании лицензии/договора N [номер] от [дата].

[дата] Заявитель подал тарифную заявку на установление тарифа на {sphere_name} на период с {period_data.get('start', '01.01.2025')} по {period_data.get('end', '31.12.2025')}.

[дата решения] [наименование регулятора] (далее - Регулятор) вынес решение N [номер] об отказе в установлении тарифа / о снижении тарифа на [процент]%.

Заявитель считает данное решение незаконным и необоснованным по следующим основаниям:

{'='*80}

2. НАРУШЕНИЯ, ДОПУЩЕННЫЕ РЕГУЛЯТОРОМ

"""
    
    for i, arg in enumerate(arguments, 1):
        text += f"""
{i}. {arg["title"]}

{arg["text"]}

Нормативное обоснование: {arg["legal_ref"]}

"""
    
    text += f"""
{'='*80}

3. НОРМАТИВНОЕ ОБОСНОВАНИЕ

Данная жалоба основана на следующих нормативных правовых актах:

"""
    
    for i, basis in enumerate(template["legal_basis"], 1):
        text += f"{i}. {basis}\n"
    
    text += f"""

{'='*80}

4. ПРОСИТЕЛЬНАЯ ЧАСТЬ

На основании изложенного, руководствуясь вышеуказанными нормативными правовыми актами,

ПРОШУ:

"""
    
    for i, demand in enumerate(template["demands"], 1):
        d = demand
        d = d.replace("[наименование регулятора]", "[наименование регулятора]")
        d = d.replace("[дата]", "[дата решения]")
        d = d.replace("[номер]", "[номер решения]")
        d = d.replace("[сумма]", f"{total_revenue:,.0f}")
        text += f"{i}. {d}\n"
    
    text += f"""

{'='*80}

5. ДОПОЛНИТЕЛЬНАЯ ИНФОРМАЦИЯ

{additional_info if additional_info else 'Дополнительно заявляем, что все доводы подтверждены документами, приложенными к тарифной заявке и настоящей жалобе.'}

{'='*80}

6. ПРИЛОЖЕНИЯ

1. Копия решения регулятора от [дата] N [номер]
2. Копия тарифной заявки со всеми приложениями
3. Документы, подтверждающие экономическую обоснованность затрат
4. Квитанция об уплате госпошлины (для суда)
5. Доверенность представителя (если применимо)
6. Иные документы по усмотрению Заявителя

{'='*80}

[Должность руководителя] _________________ / [ФИО] /

М.П.

[дата подачи]

{'='*80}

Справка для подачи:
- Госпошлина: 3 000 руб. (для арбитражного суда)
- Срок подачи: 3 месяца с даты получения решения регулятора
- Срок рассмотрения: 1 месяц (ФАС), 2 месяца (суд)

{'='*80}
Сформировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}
Система: AI-Тарифщик v1.0 - Робот-жалобщик
{'='*80}
"""
    
    return text

def generate_complaint_docx(text, addressee_type):
    """Создает DOCX документ жалобы"""
    doc = Document()
    
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    lines = text.split('\n')
    
    for line in lines:
        if line.strip() == '':
            doc.add_paragraph()
        elif line.strip().startswith('='):
            p = doc.add_paragraph()
            run = p.add_run(line.strip()[:50])
            run.font.size = Pt(8)
        elif "В Федеральную" in line or "В Арбитражный" in line or "В [наименование" in line:
            p = doc.add_paragraph(line.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.runs[0]
            run.font.size = Pt(12)
        elif "От:" in line:
            p = doc.add_paragraph(line.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.runs[0]
            run.font.size = Pt(12)
        elif line.strip().startswith("ЖАЛОБА") or line.strip().startswith("ЗАЯВЛЕНИЕ") or line.strip().startswith("АДМИНИСТРАТИВНОЕ"):
            p = doc.add_paragraph(line.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(16)
        elif line.strip().startswith("1.") or line.strip().startswith("2.") or line.strip().startswith("3.") or line.strip().startswith("4.") or line.strip().startswith("5.") or line.strip().startswith("6."):
            if "ОПИСАТЕЛЬНАЯ" in line or "НАРУШЕНИЯ" in line or "НОРМАТИВНОЕ" in line or "ПРОСИТЕЛЬНАЯ" in line or "ДОПОЛНИТЕЛЬНАЯ" in line or "ПРИЛОЖЕНИЯ" in line:
                p = doc.add_paragraph(line.strip())
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(14)
            else:
                doc.add_paragraph(line.strip(), style='List Number')
        elif line.strip().startswith("-"):
            doc.add_paragraph(line.strip(), style='List Bullet')
        elif "ПРОШУ:" in line.strip():
            p = doc.add_paragraph(line.strip())
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(14)
        elif "[Должность]" in line or "М.П." in line:
            p = doc.add_paragraph(line.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            doc.add_paragraph(line.strip())
    
    return doc

def validate_complaint(fgis_data, addressee_type, violations):
    """Проверяет данные для жалобы"""
    errors = []
    
    if not fgis_data:
        errors.append("Загрузите JSON из Экспорта ФГИС или Калькулятора рисков")
    
    if not addressee_type:
        errors.append("Выберите адресата жалобы")
    
    if not violations or len(violations) == 0:
        errors.append("Выберите хотя бы одно нарушение регулятора")
    
    return errors

# =============================================================================
# Интерфейс Streamlit
# =============================================================================

def show_complaint_bot():
    """Страница Робота-жалобщика"""
    
    st.header("Робот-жалобщик")
    st.info("Загрузите JSON -> Выберите нарушения -> Сформируйте жалобу в ФАС/суд")
    
    # Шаг 1: Загрузка JSON
    st.subheader("1. Загрузка данных")
    
    uploaded_json = st.file_uploader(
        "Загрузите JSON файл (Экспорт ФГИС или Калькулятор рисков)",
        type=['json'],
        help="Файл должен содержать данные тарифной заявки"
    )
    
    fgis_data = None
    
    if uploaded_json:
        fgis_data = load_json_data(uploaded_json)
        
        if fgis_data:
            st.success(f"JSON загружен: {uploaded_json.name}")
            
            header = fgis_data.get("header", {})
            st.info(f"**Организация:** {header.get('organization_name', 'Не указано')}")
            st.caption(f"ИНН: {header.get('inn', '')} | Сфера: {fgis_data.get('application', {}).get('activity_sphere', 'Не указано')}")
    
    st.divider()
    
    # Шаг 2: Выбор адресата жалобы
    st.subheader("2. Выбор адресата жалобы")
    
    addressee_type = st.radio(
        "Куда подаем жалобу?",
        ["fas", "arbitration", "court_general"],
        format_func=lambda x: {
            "fas": "ФАС России (досудебный порядок)",
            "arbitration": "Арбитражный суд",
            "court_general": "Суд общей юрисдикции"
        }.get(x, x),
        horizontal=True
    )
    
    if addressee_type == "fas":
        st.info("**ФАС России:**\n- Срок подачи: 1 месяц с даты решения\n- Срок рассмотрения: 1 месяц\n- Госпошлина: не требуется\n- Досудебный порядок обязателен")
    elif addressee_type == "arbitration":
        st.info("**Арбитражный суд:**\n- Срок подачи: 3 месяца с даты решения\n- Срок рассмотрения: 2 месяца\n- Госпошлина: 3 000 руб.\n- Требуется представитель")
    elif addressee_type == "court_general":
        st.info("**Суд общей юрисдикции:**\n- Срок подачи: 3 месяца с даты решения\n- Срок рассмотрения: 2 месяца\n- Госпошлина: 300 руб. (физлица)\n- Для защиты прав потребителей")
    
    st.divider()
    
    # Шаг 3: Выбор нарушений регулятора
    st.subheader("3. Нарушения регулятора")
    st.info("Выберите все нарушения, которые допущены регулятором в вашем случае")
    
    violations = st.multiselect(
        "Отметьте нарушения:",
        [
            "uncalculated_costs",
            "wrong_methodology",
            "procedural_violations",
            "wrong_comparison",
            "missing_documents",
            "wrong_amortization",
            "wrong_numeracy"
        ],
        format_func=lambda x: {
            "uncalculated_costs": "Неучтенные экономически обоснованные затраты",
            "wrong_methodology": "Неверная трактовка методики расчета",
            "procedural_violations": "Процессуальные нарушения (сроки, порядок)",
            "wrong_comparison": "Некорректное сравнение с аналогами",
            "missing_documents": "Отказ из-за отсутствия документов (необоснованный)",
            "wrong_amortization": "Неверный расчет амортизации",
            "wrong_numeracy": "Необоснованное снижение численности"
        }.get(x, x)
    )
    
    if violations:
        st.success(f"Выбрано нарушений: {len(violations)}")
    
    st.divider()
    
    # Шаг 4: Дополнительная информация
    st.subheader("4. Дополнительная информация")
    
    additional_info = st.text_area(
        "Дополнительные обстоятельства дела",
        placeholder="Укажите особые обстоятельства, которые важно отразить в жалобе",
        height=150
    )
    
    with st.expander("Рекомендации по заполнению", expanded=False):
        st.write("""
        **Что указать в дополнительной информации:**
        
        1. **Критические последствия:** "Отказ в тарифе приведет к остановке теплоснабжения 5000 потребителей"
        
        2. **История спора:** "Это уже третье решение регулятора об отказе за 2 года"
        
        3. **Прецеденты:** "Аналогичное дело NА40-12345/2024 решено в пользу РСО"
        
        4. **Социальная значимость:** "Организация является единственным поставщиком услуги в районе"
        """)
    
    st.divider()
    
    # Шаг 5: Генерация жалобы
    st.subheader("5. Генерация жалобы")
    
    if st.button("Сформировать жалобу", type="primary", use_container_width=True):
        errors = validate_complaint(fgis_data, addressee_type, violations)
        
        if errors:
            st.error("Обнаружены ошибки:")
            for e in errors:
                st.write(e)
            st.info("Исправьте ошибки перед генерацией жалобы")
        else:
            sphere = fgis_data.get("application", {}).get("activity_sphere", "heat_supply")
            
            complaint_text = generate_complaint_text(
                fgis_data,
                addressee_type,
                sphere,
                violations,
                additional_info
            )
            
            st.success("Жалоба сформирована!")
            
            st.subheader("Предпросмотр")
            st.text_area(
                "Текст жалобы",
                value=complaint_text,
                height=500,
                key="preview_complaint"
            )
            
            doc = generate_complaint_docx(complaint_text, addressee_type)
            
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            org_name = fgis_data.get("header", {}).get("organization_name", "Organization")
            inn = fgis_data.get("header", {}).get("inn", "0000000000")
            file_id = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            addressee_name = {
                "fas": "FAS",
                "arbitration": "ArbitrationCourt",
                "court_general": "GeneralCourt"
            }.get(addressee_type, "Complaint")
            
            filename_docx = f"Complaint_{addressee_name}_{inn}_{file_id}.docx"
            filename_txt = f"Complaint_{addressee_name}_{inn}_{file_id}.txt"
            
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="Скачать DOCX",
                    data=doc_buffer,
                    file_name=filename_docx,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            with col2:
                st.download_button(
                    label="Скачать TXT",
                    data=complaint_text,
                    file_name=filename_txt,
                    mime="text/plain",
                    use_container_width=True,
                )
            
            st.caption(f"Файлы: {filename_docx}, {filename_txt}")
            
            st.divider()
            st.subheader("Контрольный список перед подачей")
            
            checklist = {
                "fas": [
                    "Копия решения регулятора",
                    "Копия тарифной заявки",
                    "Документы по нарушениям",
                    "Доверенность представителя",
                    "Опись вложения"
                ],
                "arbitration": [
                    "Квитанция госпошлины (3 000 руб.)",
                    "Копия решения регулятора",
                    "Копия тарифной заявки",
                    "Документы по нарушениям",
                    "Доверенность представителя",
                    "Опись вложения"
                ],
                "court_general": [
                    "Квитанция госпошлины (300 руб.)",
                    "Копия решения регулятора",
                    "Документы по нарушениям",
                    "Доверенность представителя",
                    "Опись вложения"
                ]
            }
            
            for item in checklist.get(addressee_type, []):
                st.write(item)
            
            st.info("DOCX файл можно отредактировать в Word перед подачей")
    
    with st.expander("Пример заполнения", expanded=False):
        st.write("**Адресат:** ФАС России (досудебный порядок)")
        st.write("**Нарушения:**")
        st.write("- Неучтенные экономически обоснованные затраты")
        st.write("- Неверный расчет амортизации")
        st.write("- Процессуальные нарушения")

# =============================================================================
# Запуск
# =============================================================================

if __name__ == "__main__":
    show_complaint_bot()