# streamlit_pages/peer_comparison.py
import streamlit as st
import pandas as pd
import json
import os
from datetime import datetime
import io
from docx import Document as DocxDocument
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =============================================================================
# Функции
# =============================================================================

def get_regional_benchmarks():
    """Демо-данные по Тамбовской области (MVP)"""
    
    return {
        "region": "Тамбовская область",
        "sphere": "heat",
        "period": "2024-2025",
        "articles": {
            "Амортизация": {"avg": 180, "min": 120, "max": 250, "unit": "тыс. ₽/Гкал"},
            "Расходы на ремонт ОС": {"avg": 45, "min": 20, "max": 80, "unit": "тыс. ₽/Гкал"},
            "Заработная плата": {"avg": 120, "min": 80, "max": 180, "unit": "тыс. ₽/Гкал"},
            "Электроэнергия на собственные нужды": {"avg": 25, "min": 15, "max": 40, "unit": "тыс. ₽/Гкал"},
            "Топливо": {"avg": 350, "min": 250, "max": 500, "unit": "тыс. ₽/Гкал"},
            "Потери в сетях": {"avg": 15, "min": 8, "max": 25, "unit": "%"},
            "Численность персонала": {"avg": 12, "min": 8, "max": 20, "unit": "чел./1000 абонентов"},
            "Хозяйственные расходы": {"avg": 30, "min": 15, "max": 50, "unit": "тыс. ₽/Гкал"},
            "Прочие расходы": {"avg": 20, "min": 10, "max": 35, "unit": "тыс. ₽/Гкал"},
            "Налог на прибыль": {"avg": 25, "min": 15, "max": 40, "unit": "тыс. ₽/Гкал"}
        },
        "organizations": [
            {
                "name": "ООО «ТамбовТеплоСеть»",
                "inn": "6829012345",
                "tariff": 1850,
                "sphere": "heat",
                "release": 150000,
                "decision": "Решение №ТБ-2024-089 от 15.11.2024",
                "status": "approved"
            },
            {
                "name": "ООО «ТамбовЭнерго»",
                "inn": "6829054321",
                "tariff": 1920,
                "sphere": "heat",
                "release": 180000,
                "decision": "Решение №ТБ-2024-076 от 20.10.2024",
                "status": "approved"
            },
            {
                "name": "ООО «МичуринскТепло»",
                "inn": "6828011122",
                "tariff": 1780,
                "sphere": "heat",
                "release": 120000,
                "decision": "Решение №ТБ-2024-063 от 05.09.2024",
                "status": "approved"
            },
            {
                "name": "ООО «КотовскТеплосбыт»",
                "inn": "6827033344",
                "tariff": 2050,
                "sphere": "heat",
                "release": 95000,
                "decision": "Решение №ТБ-2024-051 от 12.08.2024",
                "status": "approved_with_conditions"
            },
            {
                "name": "ООО «РассказовоТепло»",
                "inn": "6826055566",
                "tariff": 1690,
                "sphere": "heat",
                "release": 85000,
                "decision": "Решение №ТБ-2024-042 от 18.07.2024",
                "status": "approved"
            }
        ]
    }

def compare_article(user_value, benchmark):
    """Сравнивает значение со средним по региону"""
    
    if user_value > benchmark * 1.1:  # >10% выше
        return "more", "🔴", "Выше среднего"
    elif user_value < benchmark * 0.9:  # >10% ниже
        return "less", "🟢", "Ниже среднего"
    else:
        return "equal", "🟡", "В пределах нормы"

def calculate_deviation_percent(user_value, benchmark):
    """Считает процент отклонения"""
    if benchmark == 0:
        return 0
    return round(((user_value - benchmark) / benchmark) * 100, 2)

def get_comparison_status_icon(status):
    """Возвращает иконку по статусу организации"""
    icons = {
        "approved": "✅",
        "approved_with_conditions": "⚠️",
        "rejected": "❌",
        "pending": "⏳"
    }
    return icons.get(status, "⚪")

def load_analyzer_data():
    """Загружает данные из Анализатора заявок (демо)"""
    
    return {
        "organization": "ООО «ТеплоСеть»",
        "inn": "6829000000",
        "sphere": "heat",
        "tariff": 1900,
        "release": 160000,
        "articles": {
            "Амортизация": 200,
            "Расходы на ремонт ОС": 50,
            "Заработная плата": 130,
            "Электроэнергия на собственные нужды": 28,
            "Топливо": 380,
            "Потери в сетях": 18,
            "Численность персонала": 14,
            "Хозяйственные расходы": 35,
            "Прочие расходы": 22,
            "Налог на прибыль": 28
        }
    }

def load_forecaster_data():
    """Загружает данные из Прогнозиста тарифов (демо)"""
    
    return {
        "organization": "ООО «ТеплоСеть»",
        "inn": "6829000000",
        "sphere": "heat",
        "tariff_base": 1850,
        "tariff_forecast": 1950,
        "release": 160000,
        "articles": {
            "Амортизация": 195,
            "Расходы на ремонт ОС": 48,
            "Заработная плата": 125,
            "Электроэнергия на собственные нужды": 26,
            "Топливо": 365,
            "Потери в сетях": 16,
            "Численность персонала": 13,
            "Хозяйственные расходы": 32,
            "Прочие расходы": 21,
            "Налог на прибыль": 26
        }
    }

def generate_comparison_report(comparison_data, organization_name):
    """Генерирует отчёт в DOCX"""
    
    doc = DocxDocument()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    # Заголовок
    p = doc.add_paragraph("СРАВНЕНИЕ С АНАЛОГАМИ В РЕГИОНЕ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph(f"Организация: {organization_name}")
    doc.add_paragraph(f"Регион: Тамбовская область")
    doc.add_paragraph(f"Дата отчёта: {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph(f"Период тарифа: {comparison_data.get('period', '2024-2025')}")
    
    doc.add_paragraph()
    doc.add_paragraph("-" * 80)
    doc.add_paragraph()
    
    # Итоговые цифры
    p = doc.add_paragraph("ИТОГОВЫЕ ПОКАЗАТЕЛИ:")
    run = p.runs[0]
    run.bold = True
    
    doc.add_paragraph(f"Ваш тариф: {comparison_data['user_tariff']} ₽/Гкал")
    doc.add_paragraph(f"Средний по региону: {comparison_data['avg_tariff']} ₽/Гкал")
    doc.add_paragraph(f"Отклонение: {comparison_data['tariff_deviation']:+.2f}%")
    doc.add_paragraph(f"Позиция: {comparison_data['tariff_status']}")
    
    doc.add_paragraph()
    doc.add_paragraph("-" * 80)
    doc.add_paragraph()
    
    # Таблица сравнения
    p = doc.add_paragraph("СРАВНЕНИЕ ПО СТАТЬЯМ ЗАТРАТ:")
    run = p.runs[0]
    run.bold = True
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Заголовки
    header_cells = table.rows[0].cells
    headers = ["Статья", "Ваше значение", "Среднее", "Отклонение", "Статус"]
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
    
    # Данные
    for article_data in comparison_data['articles']:
        row_cells = table.add_row().cells
        row_cells[0].text = article_data['name']
        row_cells[1].text = f"{article_data['user_value']:.2f}"
        row_cells[2].text = f"{article_data['benchmark']:.2f}"
        row_cells[3].text = f"{article_data['deviation']:+.2f}%"
        row_cells[4].text = article_data['status_text']
    
    doc.add_paragraph()
    doc.add_paragraph("-" * 80)
    doc.add_paragraph()
    
    # Подобные организации
    p = doc.add_paragraph("ПОДОБНЫЕ ОРГАНИЗАЦИИ В РЕГИОНЕ:")
    run = p.runs[0]
    run.bold = True
    
    for org in comparison_data['organizations'][:5]:
        doc.add_paragraph()
        # ИСПРАВЛЕНО: используем функцию для получения иконки
        status_icon = get_comparison_status_icon(org.get('status', 'pending'))
        p = doc.add_paragraph(f"{status_icon} {org['name']} (ИНН: {org['inn']})")
        run = p.runs[0]
        run.bold = True
        
        doc.add_paragraph(f"Тариф: {org['tariff']} ₽/Гкал")
        doc.add_paragraph(f"Полезный отпуск: {org['release']:,.0f} Гкал")
        doc.add_paragraph(f"Решение: {org['decision']}")
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# =============================================================================
# Интерфейс Streamlit
# =============================================================================

def show_peer_comparison():
    """Страница Сравнения с аналогами"""
    
    st.header("🌐 Сравнение с аналогами в регионе")
    st.info("📌 Сравнение ваших показателей с аналогичными РСО в Тамбовской области")
    
    # Инициализация session_state
    if "benchmark_data" not in st.session_state:
        st.session_state.benchmark_data = get_regional_benchmarks()
    if "user_data" not in st.session_state:
        st.session_state.user_data = None
    if "comparison_result" not in st.session_state:
        st.session_state.comparison_result = None
    
    # ─────────────────────────────────────────────────────────────────────
    # Шаг 1: Источник данных
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("1. Источник данных о тарифе")
    
    data_source = st.radio(
        "Откуда загрузить данные?",
        ["📁 Из Анализатора заявок", "🔮 Из Прогнозиста тарифов", "✍️ Ручной ввод"],
        horizontal=True,
        key="data_source_radio"
    )
    
    user_data = None
    
    if data_source == "📁 Из Анализатора заявок":
        if st.button("📁 Загрузить из Анализатора", key="load_analyzer_btn"):
            st.session_state.user_data = load_analyzer_data()
            st.success(f"✅ Загружены данные: {st.session_state.user_data['organization']}")
            st.rerun()
    
    elif data_source == "🔮 Из Прогнозиста тарифов":
        if st.button("🔮 Загрузить из Прогнозиста", key="load_forecaster_btn"):
            st.session_state.user_data = load_forecaster_data()
            st.success(f"✅ Загружены данные: {st.session_state.user_data['organization']}")
            st.rerun()
    
    else:  # Ручной ввод
        col1, col2 = st.columns(2)
        with col1:
            org_name = st.text_input("Организация", value="ООО «ТеплоСеть»", key="manual_org")
            tariff = st.number_input("Тариф (₽/Гкал)", value=1900, key="manual_tariff")
        with col2:
            inn = st.text_input("ИНН", value="6829000000", key="manual_inn")
            release = st.number_input("Полезный отпуск (Гкал)", value=160000, key="manual_release")
        
        if st.button("💾 Сохранить данные", key="save_manual_btn"):
            st.session_state.user_data = {
                "organization": org_name,
                "inn": inn,
                "sphere": "heat",
                "tariff": tariff,
                "release": release,
                "articles": {
                    "Амортизация": 200,
                    "Расходы на ремонт ОС": 50,
                    "Заработная плата": 130,
                    "Электроэнергия на собственные нужды": 28,
                    "Топливо": 380,
                    "Потери в сетях": 18,
                    "Численность персонала": 14,
                    "Хозяйственные расходы": 35,
                    "Прочие расходы": 22,
                    "Налог на прибыль": 28
                }
            }
            st.success("✅ Данные сохранены")
            st.rerun()
    
    # ─────────────────────────────────────────────────────────────────────
    # Шаг 2: Сравнение
    # ─────────────────────────────────────────────────────────────────────
    if st.session_state.user_data:
        user_data = st.session_state.user_data
        
        st.divider()
        st.subheader("2. Информация о тарифе")
        
        col1, col2, col3 = st.columns(3)
        col1.metric("Организация", user_data['organization'][:20])
        col2.metric("Тариф", f"{user_data.get('tariff', user_data.get('tariff_forecast', 0))} ₽/Гкал")
        col3.metric("Полезный отпуск", f"{user_data['release']:,.0f} Гкал")
        
        st.divider()
        st.subheader("3. Сравнение с регионом")
        
        benchmark = st.session_state.benchmark_data
        
        if st.button("🔍 Выполнить сравнение", use_container_width=True, type="primary", key="compare_btn"):
            # Сравнение по статьям
            articles_comparison = []
            
            user_articles = user_data.get('articles', {})
            
            for article_name, bench_data in benchmark['articles'].items():
                user_value = user_articles.get(article_name, bench_data['avg'])
                
                status, icon, status_text = compare_article(user_value, bench_data['avg'])
                deviation = calculate_deviation_percent(user_value, bench_data['avg'])
                
                articles_comparison.append({
                    "name": article_name,
                    "user_value": user_value,
                    "benchmark": bench_data['avg'],
                    "deviation": deviation,
                    "status": status,
                    "status_icon": icon,
                    "status_text": status_text,
                    "unit": bench_data['unit']
                })
            
            # Сортировка по отклонению
            articles_comparison.sort(key=lambda x: abs(x['deviation']), reverse=True)
            
            # Расчёт общего отклонения тарифа
            user_tariff = user_data.get('tariff', user_data.get('tariff_forecast', 0))
            avg_tariff = sum(org['tariff'] for org in benchmark['organizations']) / len(benchmark['organizations'])
            tariff_deviation = calculate_deviation_percent(user_tariff, avg_tariff)
            
            if tariff_deviation > 10:
                tariff_status = "🔴 Выше среднего"
            elif tariff_deviation < -10:
                tariff_status = "🟢 Ниже среднего"
            else:
                tariff_status = "🟡 В пределах нормы"
            
            # Сохранение результата
            st.session_state.comparison_result = {
                "articles": articles_comparison,
                "user_tariff": user_tariff,
                "avg_tariff": round(avg_tariff, 2),
                "tariff_deviation": tariff_deviation,
                "tariff_status": tariff_status,
                "organizations": benchmark['organizations'],
                "region": benchmark['region'],
                "period": benchmark['period']
            }
            
            st.success("✅ Сравнение выполнено!")
            st.rerun()
    
    # ─────────────────────────────────────────────────────────────────────
    # Шаг 3: Результаты
    # ─────────────────────────────────────────────────────────────────────
    if st.session_state.comparison_result:
        result = st.session_state.comparison_result
        
        st.divider()
        st.subheader("4. Результаты сравнения")
        
        # Итоговый блок
        st.markdown(f"""
        <div style="background: linear-gradient(90deg, #3498db, #2c3e50); 
                    padding: 2rem; border-radius: 10px; text-align: center; margin: 1rem 0;">
            <h2 style="color: white; margin: 0;">📊 Позиция относительно рынка</h2>
            <h1 style="color: white; margin: 0.5rem 0; font-size: 3rem;">
                {result['tariff_status']}
            </h1>
            <p style="color: #ecf0f1; margin: 0;">
                Ваш тариф: {result['user_tariff']} ₽/Гкал | 
                Средний: {result['avg_tariff']} ₽/Гкал | 
                Отклонение: {result['tariff_deviation']:+.2f}%
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        st.divider()
        
        # Таблица сравнения по статьям
        st.subheader("📋 Сравнение по статьям затрат")
        
        articles_df = pd.DataFrame([
            {
                "Статья": art['name'],
                "Ваше значение": f"{art['user_value']:.2f} {art['unit']}",
                "Среднее по региону": f"{art['benchmark']:.2f} {art['unit']}",
                "Отклонение": f"{art['deviation']:+.2f}%",
                "Статус": f"{art['status_icon']} {art['status_text']}"
            }
            for art in result['articles']
        ])
        
        # Цветовое форматирование
        def color_status(val):
            if "🔴" in val:
                return "background-color: #ffebee; color: #c62828"
            elif "🟢" in val:
                return "background-color: #e8f5e9; color: #2e7d32"
            else:
                return "background-color: #fff8e1; color: #f57f17"
        
        st.dataframe(
            articles_df.style.applymap(color_status, subset=["Статус"]),
            use_container_width=True,
            hide_index=True
        )
        
        st.divider()
        
        # Подобные организации
        st.subheader("🏢 Подобные организации в регионе")
        
        for i, org in enumerate(result['organizations'][:5], 1):
            with st.expander(f"{get_comparison_status_icon(org['status'])} {org['name']} (Тариф: {org['tariff']} ₽/Гкал)", expanded=(i<=3)):
                col1, col2 = st.columns(2)
                
                with col1:
                    st.write(f"**ИНН:** {org['inn']}")
                    st.write(f"**Полезный отпуск:** {org['release']:,.0f} Гкал")
                    st.write(f"**Сфера:** {'🔥 Теплоснабжение' if org['sphere'] == 'heat' else org['sphere']}")
                
                with col2:
                    st.write(f"**Статус:** {org['status']}")
                    st.write(f"**Решение:** {org['decision']}")
                    
                    if st.button("📄 Открыть решение", key=f"decision_{org['inn']}"):
                        st.info("🚧 Функция в разработке (требуется база документов)")
        
        st.divider()
        
        # ─────────────────────────────────────────────────────────────────────
        # Шаг 4: Экспорт
        # ─────────────────────────────────────────────────────────────────────
        st.subheader("5. Экспорт")
        
        col1, col2 = st.columns(2)
        
        with col1:
            docx_output = generate_comparison_report(result, user_data['organization'])
            filename = f"Peer_Comparison_{user_data['organization'][:20]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            st.download_button(
                label="📥 Скачать DOCX",
                data=docx_output,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        
        with col2:
            if st.button("📧 Отправить регулятору", use_container_width=True):
                st.info("🚧 Функция в разработке")
        
        # Рекомендации
        st.divider()
        st.subheader("💡 Рекомендации для обоснования")
        
        high_articles = [a for a in result['articles'] if a['status'] == 'more']
        
        if high_articles:
            st.warning(f"⚠️ {len(high_articles)} статей выше среднего — требуют обоснования:")
            for art in high_articles[:3]:
                st.write(f"• **{art['name']}**: {art['deviation']:+.2f}% (аргумент: износ сетей, климат, рельеф)")
        else:
            st.success("✅ Все статьи в пределах нормы или ниже среднего")
    
    else:
        st.info("👈 Загрузите данные и нажмите «Выполнить сравнение»")
    
    # ─────────────────────────────────────────────────────────────────────
    # Справка
    # ─────────────────────────────────────────────────────────────────────
    with st.expander("💡 Как использовать", expanded=False):
        st.write("""
**Назначение:**

Сравнение с аналогами показывает положение вашей организации относительно других РСО в регионе по:
- Уровню тарифа
- Структуре затрат по статьям
- Ключевым показателям (потери, численность, рентабельность)

**Источники данных:**

- **Анализатор заявок**: фактические данные из вашей заявки
- **Прогнозист тарифов**: прогнозные значения тарифа и затрат
- **Ручной ввод**: для сценарного анализа

**Как работает сравнение:**

1. Каждая статья затрат сравнивается со средним по региону
2. Отклонение >10% = выше/ниже среднего
3. Отклонение ≤10% = в пределах нормы
4. Статьи сортируются по величине отклонения

**Подобные организации:**

- Отбираются по сфере деятельности (тепло/вода/ТКО)
- Показываются тарифные решения регулятора
- Можно использовать как аргумент: «Организация X имеет тариф выше нашего»

**Регион:**

- В MVP: Тамбовская область
- В будущем: выбор любого региона РФ

**Экспорт:**

- DOCX отчёт для включения в пояснительную записку
- Автоматическая генерация аргументов для регулятора
        """)

# =============================================================================
# Запуск
# =============================================================================

if __name__ == "__main__":
    show_peer_comparison()