# streamlit_pages/law_tracker.py
import streamlit as st
import json
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os

# =============================================================================
# Функции
# =============================================================================

def get_law_changes_data():
    """Возвращает демонстрационные данные об изменениях законов"""
    
    changes = [
        {
            "id": 1,
            "date": "2025-01-15",
            "title": "Изменения в Приказе ФАС N 1746-э",
            "source": "ФАС России",
            "category": "fas",
            "sphere": ["heat_supply", "water_supply", "waste_management"],
            "impact": "high",
            "description": "Уточнены правила включения затрат на ремонт в тариф. Требуется дополнительная документация.",
            "affected_articles": ["Расходы на ремонт ОС", "Амортизация"],
            "link": "https://fas.gov.ru/documents/123456",
            "status": "new"
        },
        {
            "id": 2,
            "date": "2025-01-10",
            "title": "Новые нормативы накопления ТКО",
            "source": "Минприроды",
            "category": "npa",
            "sphere": ["waste_management"],
            "impact": "high",
            "description": "Утверждены новые нормативы накопления ТКО для различных категорий потребителей.",
            "affected_articles": ["Транспортирование ТКО", "Размещение на полигоне"],
            "link": "https://minprirody.gov.ru/doc/789",
            "status": "new"
        },
        {
            "id": 3,
            "date": "2025-01-08",
            "title": "Разъяснение по учёту электроэнергии",
            "source": "ФАС России",
            "category": "fas",
            "sphere": ["water_supply", "heat_supply"],
            "impact": "medium",
            "description": "Разъяснён порядок учёта затрат на электроэнергию для собственных нужд.",
            "affected_articles": ["Электроэнергия", "Расходы на перекачку"],
            "link": "https://fas.gov.ru/documents/654321",
            "status": "reviewed"
        },
        {
            "id": 4,
            "date": "2025-01-05",
            "title": "Индексация тарифов на 2025 год",
            "source": "Правительство РФ",
            "category": "npa",
            "sphere": ["heat_supply", "water_supply", "water_drainage", "waste_management"],
            "impact": "high",
            "description": "Установлены предельные индексы изменения тарифов на 2025 год по регионам.",
            "affected_articles": ["Валовая выручка", "Тарифные ставки"],
            "link": "https://government.ru/doc/456",
            "status": "new"
        },
        {
            "id": 5,
            "date": "2025-01-03",
            "title": "Методичка по расчёту потерь",
            "source": "Минэнерго",
            "category": "methodics",
            "sphere": ["heat_supply", "water_supply"],
            "impact": "medium",
            "description": "Обновлена методика расчёта нормативов потерь в сетях.",
            "affected_articles": ["Потери в сетях", "Техническое обслуживание"],
            "link": "https://minenergo.gov.ru/doc/321",
            "status": "reviewed"
        },
        {
            "id": 6,
            "date": "2024-12-28",
            "title": "Судебная практика по амортизации",
            "source": "Арбитражный суд",
            "category": "court",
            "sphere": ["heat_supply", "water_supply", "waste_management"],
            "impact": "low",
            "description": "Суд подтвердил правомерность включения затрат на модернизацию ОС в тариф.",
            "affected_articles": ["Амортизация", "Инвестиционная программа"],
            "link": "https://kad.arbitr.ru/Card/123456",
            "status": "reviewed"
        },
        {
            "id": 7,
            "date": "2024-12-25",
            "title": "Изменения в классификации ОС",
            "source": "Правительство РФ",
            "category": "npa",
            "sphere": ["heat_supply", "water_supply", "water_drainage", "waste_management"],
            "impact": "medium",
            "description": "Обновлена классификация основных средств для целей амортизации.",
            "affected_articles": ["Амортизация", "Основные средства"],
            "link": "https://government.ru/doc/789",
            "status": "reviewed"
        },
        {
            "id": 8,
            "date": "2024-12-20",
            "title": "Требования к отчётности РСО",
            "source": "ФАС России",
            "category": "fas",
            "sphere": ["heat_supply", "water_supply", "water_drainage", "waste_management"],
            "impact": "low",
            "description": "Введены новые формы отчётности для ресурсоснабжающих организаций.",
            "affected_articles": ["Административные расходы", "Отчётность"],
            "link": "https://fas.gov.ru/documents/987654",
            "status": "reviewed"
        }
    ]
    
    return changes

def filter_changes(changes, sphere_filter, category_filter, date_from, date_to):
    """Фильтрует изменения по параметрам"""
    
    filtered = changes.copy()
    
    if sphere_filter and sphere_filter != "all":
        filtered = [c for c in filtered if sphere_filter in c.get("sphere", [])]
    
    if category_filter and category_filter != "all":
        filtered = [c for c in filtered if c.get("category") == category_filter]
    
    if date_from:
        filtered = [c for c in filtered if c.get("date", "") >= date_from]
    
    if date_to:
        filtered = [c for c in filtered if c.get("date", "") <= date_to]
    
    return sorted(filtered, key=lambda x: x.get("date", ""), reverse=True)

def get_impact_label(impact):
    """Возвращает метку влияния"""
    labels = {
        "high": "🔴 Высокое",
        "medium": "🟡 Среднее",
        "low": "🟢 Низкое"
    }
    return labels.get(impact, "⚪ Не определено")

def get_category_label(category):
    """Возвращает метку категории"""
    labels = {
        "npa": "📜 НПА",
        "fas": "⚖️ ФАС",
        "court": "🏛️ Суд",
        "methodics": "📋 Методичка"
    }
    return labels.get(category, "📄 Документ")

def generate_report(changes, sphere_filter, category_filter):
    """Генерирует отчёт об изменениях"""
    
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    p = doc.add_paragraph("ОТЧЁТ ОБ ИЗМЕНЕНИЯХ В ЗАКОНОДАТЕЛЬСТВЕ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph(f"Сфера: {sphere_filter if sphere_filter != 'all' else 'Все сферы'}")
    doc.add_paragraph(f"Категория: {category_filter if category_filter != 'all' else 'Все категории'}")
    
    doc.add_paragraph()
    p = doc.add_paragraph(f"Всего изменений: {len(changes)}")
    run = p.runs[0]
    run.bold = True
    
    doc.add_paragraph()
    
    for change in changes[:20]:
        p = doc.add_paragraph(f"{change['date']} - {change['title']}")
        run = p.runs[0]
        run.bold = True
        
        doc.add_paragraph(f"Источник: {change['source']}")
        doc.add_paragraph(f"Влияние: {get_impact_label(change['impact'])}")
        doc.add_paragraph(f"Описание: {change['description']}")
        doc.add_paragraph(f"Затрагиваемые статьи: {', '.join(change['affected_articles'])}")
        doc.add_paragraph("-" * 80)
    
    return doc

# =============================================================================
# Интерфейс Streamlit
# =============================================================================

def show_law_tracker():
    """Страница Трекера изменений законов"""
    
    st.header("🔄 Трекер изменений законов")
    st.info("📌 Мониторинг изменений в законодательстве по тарифному регулированию")
    
    # Инициализация session_state
    if "law_changes" not in st.session_state:
        st.session_state.law_changes = get_law_changes_data()
    if "viewed_changes" not in st.session_state:
        st.session_state.viewed_changes = []
    
    # ─────────────────────────────────────────────────────────────────────
    # Шаг 1: Фильтры
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("1. Фильтры")
    
    col1, col2 = st.columns(2)
    
    with col1:
        sphere_filter = st.selectbox(
            "Сфера деятельности",
            ["all", "heat_supply", "water_supply", "water_drainage", "waste_management"],
            format_func=lambda x: {
                "all": "Все сферы",
                "heat_supply": "🔥 Теплоснабжение",
                "water_supply": "💧 Водоснабжение",
                "water_drainage": "🚰 Водоотведение",
                "waste_management": "🗑️ ТКО"
            }.get(x, x)
        )
    
    with col2:
        category_filter = st.selectbox(
            "Категория документа",
            ["all", "npa", "fas", "court", "methodics"],
            format_func=lambda x: {
                "all": "Все категории",
                "npa": "📜 НПА (законы)",
                "fas": "⚖️ Документы ФАС",
                "court": "🏛️ Судебная практика",
                "methodics": "📋 Методички"
            }.get(x, x)
        )
    
    col1, col2 = st.columns(2)
    with col1:
        date_from = st.date_input("С даты", value=datetime.now() - timedelta(days=30))
    with col2:
        date_to = st.date_input("По дату", value=datetime.now())
    
    # ─────────────────────────────────────────────────────────────────────
    # Шаг 2: Применение фильтров
    # ─────────────────────────────────────────────────────────────────────
    st.divider()
    st.subheader("2. Изменения в законодательстве")
    
    filtered_changes = filter_changes(
        st.session_state.law_changes,
        sphere_filter,
        category_filter,
        date_from.isoformat(),
        date_to.isoformat()
    )
    
    st.caption(f"Найдено изменений: {len(filtered_changes)}")
    
    # Статистика
    col1, col2, col3, col4 = st.columns(4)
    high_impact = len([c for c in filtered_changes if c.get("impact") == "high"])
    medium_impact = len([c for c in filtered_changes if c.get("impact") == "medium"])
    low_impact = len([c for c in filtered_changes if c.get("impact") == "low"])
    new_changes = len([c for c in filtered_changes if c.get("status") == "new"])
    
    col1.metric("🔴 Высокое влияние", high_impact)
    col2.metric("🟡 Среднее влияние", medium_impact)
    col3.metric("🟢 Низкое влияние", low_impact)
    col4.metric("🆕 Новые", new_changes)
    
    st.divider()
    
    # ─────────────────────────────────────────────────────────────────────
    # Шаг 3: Список изменений
    # ─────────────────────────────────────────────────────────────────────
    if filtered_changes:
        for change in filtered_changes:
            impact_label = get_impact_label(change["impact"])
            category_label = get_category_label(change["category"])
            
            is_new = change["id"] not in st.session_state.viewed_changes
            
            with st.expander(
                f"{'🆕' if is_new else ''} {change['date']} - {change['title']} [{impact_label}]",
                expanded=is_new
            ):
                col1, col2 = st.columns([3, 1])
                
                with col1:
                    st.write(f"**Источник:** {change['source']}")
                    st.write(f"**Категория:** {category_label}")
                    st.write(f"**Сферы:** {', '.join([{'heat_supply': '🔥 Теплоснабжение', 'water_supply': '💧 Водоснабжение', 'water_drainage': '🚰 Водоотведение', 'waste_management': '🗑️ ТКО'}.get(s, s) for s in change.get('sphere', [])])}")
                
                with col2:
                    if is_new:
                        if st.button("✅ Отметить как просмотренное", key=f"view_{change['id']}"):
                            st.session_state.viewed_changes.append(change["id"])
                            st.rerun()
                    else:
                        st.success("✅ Просмотрено")
                
                st.write(f"**Описание:** {change['description']}")
                
                st.write("**Затрагиваемые статьи затрат:**")
                for article in change.get("affected_articles", []):
                    st.write(f"- {article}")
                
                if change.get("link"):
                    st.write(f"**Ссылка:** [{change['link']}]({change['link']})")
                
                st.divider()
                
                # Интеграция с другими модулями
                st.write("**Действия:**")
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    if st.button("📝 В пояснительную записку", key=f"note_{change['id']}"):
                        st.info("💡 Это изменение будет учтено в Генераторе пояснительной записки")
                
                with col2:
                    if st.button("📊 Пересчитать риски", key=f"risk_{change['id']}"):
                        st.info("💡 Это изменение будет учтено в Калькуляторе рисков")
                
                with col3:
                    if st.button("⚖️ Позиция ФАС", key=f"fas_{change['id']}"):
                        st.info("💡 Откройте Позицию ФАС для анализа")
    else:
        st.info("📭 По выбранным фильтрам изменений не найдено")
    
    # ─────────────────────────────────────────────────────────────────────
    # Шаг 4: Экспорт отчёта
    # ─────────────────────────────────────────────────────────────────────
    st.divider()
    st.subheader("3. Экспорт отчёта")
    
    if st.button("📄 Сформировать отчёт об изменениях", type="primary", use_container_width=True):
        doc = generate_report(filtered_changes, sphere_filter, category_filter)
        
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        file_id = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"LawTracker_Report_{file_id}.docx"
        
        st.success("✅ Отчёт сформирован!")
        
        st.download_button(
            label="📥 Скачать отчёт (DOCX)",
            data=doc_buffer,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        
        st.caption(f"📁 Файл: {filename}")
    
    # ─────────────────────────────────────────────────────────────────────
    # Шаг 5: Настройки уведомлений
    # ─────────────────────────────────────────────────────────────────────
    st.divider()
    st.subheader("4. Настройки уведомлений")
    
    st.info("📌 Уведомления о новых изменениях приходят только в приложении")
    
    col1, col2 = st.columns(2)
    
    with col1:
        notify_daily = st.checkbox("Ежедневные уведомления", value=True)
        notify_weekly = st.checkbox("Еженедельный дайджест", value=False)
    
    with col2:
        notify_high_only = st.checkbox("Только высокое влияние", value=False)
        notify_all = st.checkbox("Все изменения", value=True)
    
    if st.button("💾 Сохранить настройки"):
        st.success("✅ Настройки сохранены!")
        st.session_state.tracker_settings = {
            "daily": notify_daily,
            "weekly": notify_weekly,
            "high_only": notify_high_only,
            "all": notify_all
        }
    
    # ─────────────────────────────────────────────────────────────────────
    # Справка
    # ─────────────────────────────────────────────────────────────────────
    with st.expander("💡 Как использовать Трекер", expanded=False):
        st.write("""
        **Назначение:**
        
        Трекер изменений законов помогает отслеживать нововведения в тарифном регулировании:
        - Новые НПА и законы
        - Разъяснения ФАС
        - Судебная практика
        - Методические указания
        
        **Как работать:**
        
        1. Выберите сферу деятельности (ТКО, тепло, вода)
        2. Выберите категорию документов (НПА, ФАС, суд)
        3. Укажите период для поиска
        4. Изучите изменения и отметьте просмотренные
        5. Скачайте отчёт для внутреннего документооборота
        
        **Интеграция с другими модулями:**
        
        - 📝 Пояснительная записка: учёт новых требований в обосновании
        - 📊 Калькулятор рисков: пересчёт рисков с учётом изменений
        - ⚖️ Позиция ФАС: анализ трактовки новых норм
        
        **Частота обновлений:**
        
        - Проверка источников: ежедневно
        - Уведомления: по настроенному расписанию
        - База изменений: обновляется автоматически
        """)

# =============================================================================
# Запуск
# =============================================================================

if __name__ == "__main__":
    show_law_tracker()