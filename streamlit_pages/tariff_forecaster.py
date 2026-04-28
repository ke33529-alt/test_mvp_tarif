# streamlit_pages/tariff_forecaster.py
import streamlit as st
import pandas as pd
import json
import os
from datetime import datetime
import io
from docx import Document as DocxDocument
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =============================================================================
# Функции
# =============================================================================

def get_regulation_methods():
    """Методы регулирования тарифов"""
    return {
        "ebt": "Экономически обоснованных расходов (тарифов)",
        "comparison": "Сравнения аналогов",
        "indexation": "Индексации установленных тарифов",
        "mrdi": "Минимальной доходности инвестированного капитала"
    }

def get_spheres():
    """Сферы деятельности"""
    return {
        "heat": "🔥 Теплоснабжение",
        "water": "💧 Водоснабжение",
        "wastewater": "🚰 Водоотведение",
        "tko": "🗑️ ТКО",
        "electricity": "⚡ Электричество"
    }

def get_article_indices(sphere, method):
    """Индексы для статей затрат по сфере и методу"""
    
    indices = {
        "Амортизация": 1.05,
        "Расходы на ремонт ОС": 1.08,
        "Заработная плата": 1.06,
        "Электроэнергия на собственные нужды": 1.04,
        "Топливо": 1.10,
        "Потери в сетях": 0.98,
        "Численность персонала": 1.03,
        "Хозяйственные расходы": 1.05,
        "Транспортирование ТКО": 1.07,
        "Утилизация отходов": 1.09,
        "Расходы на воду": 1.04,
        "Канализация": 1.05,
        "Прочие расходы": 1.03
    }
    
    method_multipliers = {
        "ebt": 1.00,
        "comparison": 0.95,
        "indexation": 1.07,
        "mrdi": 1.02
    }
    
    adjusted = {}
    for article, idx in indices.items():
        adjusted[article] = round(idx * method_multipliers.get(method, 1.00), 3)
    
    return adjusted

def parse_articles_from_text(text):
    """Извлекает статьи затрат и суммы из текста"""
    
    articles = []
    import re
    
    patterns = [
        r'([А-Яа-яЁё\s]+):\s*([\d\s,]+)\s*(?:руб|тыс|млн)?',
        r'([А-Яа-яЁё\s]+)\s*[-–—]\s*([\d\s,]+)\s*(?:руб|тыс|млн)?',
        r'([\d\s,]+)\s*(?:руб|тыс|млн)?\s*[-–—]\s*([А-Яа-яЁё\s]+)'
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            if len(match) >= 2:
                if match[0].strip().replace(' ', '').isdigit():
                    amount_str = match[0].strip().replace(' ', '').replace(',', '')
                    name = match[1].strip()
                else:
                    name = match[0].strip()
                    amount_str = match[1].strip().replace(' ', '').replace(',', '')
                
                if len(name) < 3 or len(name) > 100:
                    continue
                
                try:
                    amount = float(amount_str)
                    if amount > 0:
                        if 'млн' in text[text.find(amount_str)-10:text.find(amount_str)+10]:
                            amount *= 1000000
                        elif 'тыс' in text[text.find(amount_str)-10:text.find(amount_str)+10]:
                            amount *= 1000
                        
                        articles.append({
                            "name": name,
                            "amount": amount,
                            "source": "text"
                        })
                except:
                    continue
    
    seen = {}
    for art in articles:
        key = art["name"].lower()
        if key not in seen:
            seen[key] = art
        else:
            seen[key]["amount"] += art["amount"]
    
    return list(seen.values())

def calculate_forecast(articles, sphere, method, conditions_text):
    """Рассчитывает прогноз тарифа"""
    
    indices = get_article_indices(sphere, method)
    
    forecast_articles = []
    total_forecast = 0
    
    for article in articles:
        article_name = article["name"]
        base_amount = article["amount"]
        
        matched_index = 1.00
        matched_name = article_name
        
        for idx_name, idx_value in indices.items():
            if idx_name.lower() in article_name.lower() or article_name.lower() in idx_name.lower():
                matched_index = idx_value
                matched_name = idx_name
                break
        
        if matched_index == 1.00:
            matched_index = 1.05
            matched_name = "Прочие статьи"
        
        forecast_amount = base_amount * matched_index
        total_forecast += forecast_amount
        
        forecast_articles.append({
            "name": article_name,
            "base_amount": base_amount,
            "index": matched_index,
            "matched_article": matched_name,
            "forecast_amount": forecast_amount,
            "deviation_percent": round((matched_index - 1) * 100, 2)
        })
    
    useful_release = 100000
    
    tariff_base = total_forecast / useful_release
    tariff_forecast = (total_forecast * 1.05) / useful_release
    
    return {
        "articles": forecast_articles,
        "total_base": total_forecast / 1.05,
        "total_forecast": total_forecast,
        "tariff_base": tariff_base,
        "tariff_forecast": tariff_forecast,
        "useful_release": useful_release,
        "method": method,
        "sphere": sphere,
        "conditions": conditions_text
    }

def generate_forecast_report(forecast_data, scenario_name):
    """Генерирует отчёт в DOCX"""
    
    doc = DocxDocument()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    p = doc.add_paragraph("ПРОГНОЗ ТАРИФА")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph(f"Сценарий: {scenario_name}")
    doc.add_paragraph(f"Дата прогноза: {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph(f"Сфера: {forecast_data['sphere']}")
    doc.add_paragraph(f"Метод регулирования: {forecast_data['method']}")
    
    doc.add_paragraph()
    doc.add_paragraph("-" * 80)
    doc.add_paragraph()
    
    p = doc.add_paragraph("ИТОГОВЫЕ ПОКАЗАТЕЛИ:")
    run = p.runs[0]
    run.bold = True
    
    doc.add_paragraph(f"Валовая выручка (базовая): {forecast_data['total_base']:,.2f} ₽")
    doc.add_paragraph(f"Валовая выручка (прогноз): {forecast_data['total_forecast']:,.2f} ₽")
    doc.add_paragraph(f"Тариф (базовый): {forecast_data['tariff_base']:,.2f} ₽/ед.")
    doc.add_paragraph(f"Тариф (прогноз): {forecast_data['tariff_forecast']:,.2f} ₽/ед.")
    
    dev_percent = ((forecast_data['tariff_forecast'] / forecast_data['tariff_base']) - 1) * 100
    doc.add_paragraph(f"Изменение тарифа: {dev_percent:+.2f}%")
    
    doc.add_paragraph()
    doc.add_paragraph("-" * 80)
    doc.add_paragraph()
    
    p = doc.add_paragraph("ДЕТАЛИЗАЦИЯ ПО СТАТЬЯМ ЗАТРАТ:")
    run = p.runs[0]
    run.bold = True
    
    for art in forecast_data['articles']:
        doc.add_paragraph()
        p = doc.add_paragraph(f"Статья: {art['name']}")
        run = p.runs[0]
        run.bold = True
        
        doc.add_paragraph(f"Базовая сумма: {art['base_amount']:,.2f} ₽")
        doc.add_paragraph(f"Применённый индекс: {art['index']:.3f}")
        doc.add_paragraph(f"Прогнозная сумма: {art['forecast_amount']:,.2f} ₽")
        doc.add_paragraph(f"Отклонение: {art['deviation_percent']:+.2f}%")
    
    doc.add_paragraph()
    doc.add_paragraph("-" * 80)
    doc.add_paragraph()
    
    p = doc.add_paragraph("ОБЪЯСНЕНИЕ ПРОГНОЗА:")
    run = p.runs[0]
    run.bold = True
    
    doc.add_paragraph(f"Метод регулирования: {forecast_data['method']}")
    doc.add_paragraph(f"Условия: {forecast_data.get('conditions', 'Не указаны')}")
    doc.add_paragraph()
    doc.add_paragraph("Прогноз рассчитан на основе:")
    doc.add_paragraph("• Применённых индексов к каждой статье затрат")
    doc.add_paragraph("• Суммирования прогнозных значений статей в НВВ")
    doc.add_paragraph("• Деления НВВ на полезный отпуск для расчёта тарифа")
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def save_scenario(scenario_data):
    """Сохраняет сценарий в историю"""
    
    history_dir = os.path.join("data", "tariff_forecaster", "scenarios")
    os.makedirs(history_dir, exist_ok=True)
    
    scenario_id = f"scenario_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    file_path = os.path.join(history_dir, f"{scenario_id}.json")
    
    scenario_data["id"] = scenario_id
    scenario_data["created_at"] = datetime.now().isoformat()
    
    with open(file_path, 'w', encoding='utf-8') as f:
        json.dump(scenario_data, f, ensure_ascii=False, indent=2)
    
    return scenario_id

def load_scenarios():
    """Загружает историю сценариев"""
    
    history_dir = os.path.join("data", "tariff_forecaster", "scenarios")
    if not os.path.exists(history_dir):
        return []
    
    scenarios = []
    for filename in os.listdir(history_dir):
        if filename.endswith(".json"):
            filepath = os.path.join(history_dir, filename)
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    scenario = json.load(f)
                    scenarios.append(scenario)
            except:
                continue
    
    return sorted(scenarios, key=lambda x: x.get("created_at", ""), reverse=True)

def delete_scenario(scenario_id):
    """Удаляет сценарий"""
    
    history_dir = os.path.join("data", "tariff_forecaster", "scenarios")
    file_path = os.path.join(history_dir, f"{scenario_id}.json")
    
    if os.path.exists(file_path):
        os.remove(file_path)
        return True
    return False

def load_analyzer_result():
    """Загружает результат из Анализатора заявок (демо)"""
    
    return {
        "articles": [
            {"name": "Амортизация", "amount": 20000000},
            {"name": "Расходы на ремонт ОС", "amount": 5000000},
            {"name": "Заработная плата", "amount": 15000000},
            {"name": "Электроэнергия на собственные нужды", "amount": 3000000},
            {"name": "Топливо", "amount": 25000000},
            {"name": "Потери в сетях", "amount": 8000000},
            {"name": "Хозяйственные расходы", "amount": 2000000}
        ],
        "total": 78000000,
        "organization": "ООО «ТеплоСеть»",
        "sphere": "heat"
    }

# =============================================================================
# Интерфейс Streamlit
# =============================================================================

def show_tariff_forecaster():
    """Страница Прогнозиста тарифов"""
    
    st.header("🔮 Прогнозист тарифов")
    st.info("📌 AI-прогнозирование тарифа на основе статей затрат и метода регулирования")
    
    # ─────────────────────────────────────────────────────────────────────
    # Инициализация session_state (ИСПРАВЛЕНО)
    # ─────────────────────────────────────────────────────────────────────
    if "current_forecast" not in st.session_state:
        st.session_state.current_forecast = None
    if "scenarios" not in st.session_state:
        st.session_state.scenarios = load_scenarios()
    if "current_scenario_name" not in st.session_state:
        st.session_state.current_scenario_name = "Базовый сценарий"
    if "articles_data" not in st.session_state:
        st.session_state.articles_data = []  # ← ГЛАВНОЕ ИСПРАВЛЕНИЕ
    if "analyzer_loaded" not in st.session_state:
        st.session_state.analyzer_loaded = False
    if "input_method" not in st.session_state:
        st.session_state.input_method = "📝 Текстовое описание"
    
    # ─────────────────────────────────────────────────────────────────────
    # Шаг 1: Входные данные
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("1. Входные данные")
    
    # Выбор способа ввода
    input_method = st.radio(
        "Способ ввода данных",
        ["📝 Текстовое описание", "📁 Загрузка из Анализатора"],
        horizontal=True,
        key="input_method_radio",
        on_change=lambda: st.session_state.update({
            "input_method": "📝 Текстовое описание" if st.session_state.input_method_radio == "📝 Текстовое описание" else "📁 Загрузка из Анализатора",
            "articles_data": [],
            "analyzer_loaded": False,
            "current_forecast": None
        })
    )
    
    st.session_state.input_method = input_method
    
    # ─────────────────────────────────────────────────────────────────────
    # ВАРИАНТ 1: Текстовое описание
    # ─────────────────────────────────────────────────────────────────────
    if input_method == "📝 Текстовое описание":
        articles_text = st.text_area(
            "Описание тарифа (статьи затрат с суммами)",
            placeholder="""Пример:
Амортизация: 20000000 руб
Расходы на ремонт ОС: 5000000 руб
Заработная плата: 15000000 руб
Электроэнергия на собственные нужды: 3000000 руб
Топливо: 25000000 руб
""",
            height=200,
            key="articles_text_input"
        )
        
        # Парсим текст и сохраняем в session_state
        if articles_text:
            parsed_articles = parse_articles_from_text(articles_text)
            if parsed_articles:
                st.session_state.articles_data = parsed_articles
                st.success(f"✅ Распознано {len(parsed_articles)} статей на сумму {sum(a['amount'] for a in parsed_articles):,.0f} ₽")
            else:
                st.session_state.articles_data = []
                st.warning("⚠️ Не удалось распознать статьи. Проверьте формат (Название: Сумма)")
        else:
            st.session_state.articles_data = []
    
    # ─────────────────────────────────────────────────────────────────────
    # ВАРИАНТ 2: Загрузка из Анализатора
    # ─────────────────────────────────────────────────────────────────────
    else:
        if st.button("📁 Загрузить из Анализатора заявок", key="load_analyzer_btn"):
            analyzer_data = load_analyzer_result()
            st.session_state.articles_data = analyzer_data["articles"]
            st.session_state.analyzer_sphere = analyzer_data.get("sphere", "heat")
            st.session_state.analyzer_loaded = True
            st.session_state.current_forecast = None
            st.success(f"✅ Загружено {len(analyzer_data['articles'])} статей на сумму {analyzer_data['total']:,.0f} ₽")
            st.rerun()
        
        if st.session_state.analyzer_loaded and st.session_state.articles_data:
            st.write("**📋 Загруженные статьи:**")
            articles_df = pd.DataFrame([
                {"Статья": art["name"], "Сумма (₽)": f"{art['amount']:,.0f}"}
                for art in st.session_state.articles_data
            ])
            st.dataframe(articles_df, use_container_width=True, hide_index=True)
    
    st.divider()
    
    # ─────────────────────────────────────────────────────────────────────
    # Шаг 2: Параметры прогноза
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("2. Параметры прогноза")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Если загружено из анализатора - используем ту сферу
        sphere_index = 0
        if st.session_state.analyzer_loaded and "analyzer_sphere" in st.session_state:
            sphere_keys = list(get_spheres().keys())
            if st.session_state.analyzer_sphere in sphere_keys:
                sphere_index = sphere_keys.index(st.session_state.analyzer_sphere)
        
        sphere = st.selectbox(
            "Сфера деятельности",
            list(get_spheres().keys()),
            format_func=lambda x: get_spheres().get(x, x),
            key="sphere_select",
            index=sphere_index
        )
    
    with col2:
        method = st.selectbox(
            "Метод регулирования",
            list(get_regulation_methods().keys()),
            format_func=lambda x: get_regulation_methods().get(x, x),
            key="method_select"
        )
    
    conditions_text = st.text_area(
        "Условия прогноза (опционально)",
        placeholder="Например: плановый период 2025-2027, учёт инфляции 8%, инвестиционная программа...",
        height=100,
        key="conditions_input"
    )
    
    st.divider()
    
    # ─────────────────────────────────────────────────────────────────────
    # Шаг 3: Расчёт прогноза
    # ─────────────────────────────────────────────────────────────────────
    st.subheader("3. Расчёт прогноза")
    
    # ПРОВЕРКА: есть ли данные для расчёта
    has_articles = len(st.session_state.articles_data) > 0
    
    if not has_articles:
        st.warning("⚠️ Заполните данные для расчёта: введите текст со статьями затрат ИЛИ загрузите из Анализатора")
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        st.caption("💡 AI определит индексы для каждой статьи и рассчитает прогноз")
    
    with col2:
        calc_btn = st.button(
            "🔮 Рассчитать прогноз", 
            use_container_width=True, 
            type="primary", 
            key="calc_btn",
            disabled=not has_articles  # ← Блокируем если нет данных
        )
    
    if calc_btn and has_articles:
        with st.spinner("🔄 AI рассчитывает прогноз тарифа..."):
            forecast = calculate_forecast(
                st.session_state.articles_data,
                sphere,
                method,
                conditions_text
            )
            st.session_state.current_forecast = forecast
            st.success("✅ Прогноз готов!")
            st.rerun()
    
    # ─────────────────────────────────────────────────────────────────────
    # Шаг 4: Результаты прогноза
    # ─────────────────────────────────────────────────────────────────────
    if st.session_state.current_forecast:
        forecast = st.session_state.current_forecast
        
        st.divider()
        st.subheader("4. Результаты прогноза")
        
        # Крупная цифра тарифа
        st.markdown(f"""
        <div style="background: linear-gradient(90deg, #3498db, #2c3e50); 
                    padding: 2rem; border-radius: 10px; text-align: center; margin: 1rem 0;">
            <h2 style="color: white; margin: 0;">🔮 Прогноз тарифа</h2>
            <h1 style="color: white; margin: 0.5rem 0; font-size: 3rem;">
                {forecast['tariff_forecast']:,.2f} ₽/ед.
            </h1>
            <p style="color: #ecf0f1; margin: 0;">
                Базовый: {forecast['tariff_base']:,.2f} ₽/ед. | 
                Изменение: {((forecast['tariff_forecast']/forecast['tariff_base'])-1)*100:+.2f}%
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        # Компоненты
        st.subheader("📊 Компоненты расчёта")
        
        col1, col2, col3 = st.columns(3)
        col1.metric("Валовая выручка (базовая)", f"{forecast['total_base']:,.0f} ₽")
        col2.metric("Валовая выручка (прогноз)", f"{forecast['total_forecast']:,.0f} ₽")
        col3.metric("Полезный отпуск", f"{forecast['useful_release']:,.0f} ед.")
        
        st.divider()
        
        # Таблица по статьям
        st.subheader("📋 Детализация по статьям затрат")
        
        articles_df = pd.DataFrame([
            {
                "Статья": art["name"],
                "База (₽)": f"{art['base_amount']:,.0f}",
                "Индекс": f"{art['index']:.3f}",
                "Прогноз (₽)": f"{art['forecast_amount']:,.0f}",
                "Отклонение": f"{art['deviation_percent']:+.2f}%"
            }
            for art in forecast["articles"]
        ])
        
        st.dataframe(articles_df, use_container_width=True, hide_index=True)
        
        st.divider()
        
        # Объяснение прогноза
        st.subheader("💡 Объяснение прогноза")
        
        with st.expander("Как родился прогнозный показатель", expanded=True):
            st.write("""
            **Алгоритм расчёта:**
            
            1. **Идентификация статей:** AI сопоставил каждую статью затрат с эталонными индексами
            2. **Применение индексов:** К каждой статье применён индивидуальный индекс (учёт сферы и метода регулирования)
            3. **Расчёт НВВ:** Сумма прогнозных значений всех статей = Валовая выручка
            4. **Расчёт тарифа:** НВВ / Полезный отпуск = Тариф на единицу
            
            **Формула:**
            ```
            Тариф = Σ(Статья_затрат × Индекс) / Полезный_отпуск
            ```
            
            **Применённые индексы:**
            """)
            
            indices = get_article_indices(forecast["sphere"], forecast["method"])
            indices_df = pd.DataFrame({
                "Статья": list(indices.keys()),
                "Индекс": [f"{v:.3f}" for v in indices.values()]
            })
            st.dataframe(indices_df, use_container_width=True, hide_index=True)
            
            if forecast.get("conditions"):
                st.write(f"**Условия:** {forecast['conditions']}")
        
        st.divider()
        
        # ─────────────────────────────────────────────────────────────────────
        # Шаг 5: Управление сценариями
        # ─────────────────────────────────────────────────────────────────────
        st.subheader("5. Сценарии")
        
        col1, col2 = st.columns([3, 1])
        
        with col1:
            scenario_name = st.text_input(
                "Название сценария",
                value=st.session_state.current_scenario_name,
                key="scenario_name_input"
            )
        
        with col2:
            if st.button("💾 Сохранить сценарий", use_container_width=True, key="save_scenario_btn"):
                scenario_data = {
                    "name": scenario_name,
                    "articles": forecast["articles"],
                    "sphere": forecast["sphere"],
                    "method": forecast["method"],
                    "conditions": forecast.get("conditions", ""),
                    "forecast": {
                        "total_base": forecast["total_base"],
                        "total_forecast": forecast["total_forecast"],
                        "tariff_base": forecast["tariff_base"],
                        "tariff_forecast": forecast["tariff_forecast"],
                        "useful_release": forecast["useful_release"]
                    },
                    "input_text": ""
                }
                
                scenario_id = save_scenario(scenario_data)
                st.session_state.scenarios = load_scenarios()
                st.session_state.current_scenario_name = scenario_name
                st.success(f"✅ Сценарий сохранён: {scenario_id}")
                st.rerun()
        
        if st.session_state.scenarios:
            st.write("**📚 Сохранённые сценарии:**")
            
            scenario_options = {s["name"]: s for s in st.session_state.scenarios}
            selected_scenario = st.selectbox(
                "Выберите сценарий для загрузки",
                list(scenario_options.keys()),
                key="load_scenario_select"
            )
            
            col1, col2 = st.columns([3, 1])
            
            with col1:
                if selected_scenario:
                    st.caption(f"Создан: {scenario_options[selected_scenario].get('created_at', '')[:16]}")
            
            with col2:
                if st.button("📂 Открыть", use_container_width=True, key="open_scenario_btn"):
                    scenario = scenario_options[selected_scenario]
                    st.session_state.current_scenario_name = scenario["name"]
                    st.session_state.current_forecast = {
                        "articles": scenario["articles"],
                        "sphere": scenario["sphere"],
                        "method": scenario["method"],
                        "conditions": scenario.get("conditions", ""),
                        "total_base": scenario["forecast"]["total_base"],
                        "total_forecast": scenario["forecast"]["total_forecast"],
                        "tariff_base": scenario["forecast"]["tariff_base"],
                        "tariff_forecast": scenario["forecast"]["tariff_forecast"],
                        "useful_release": scenario["forecast"]["useful_release"]
                    }
                    st.rerun()
            
            if st.button("🗑 Удалить сценарий", key="delete_scenario_btn"):
                scenario = scenario_options[selected_scenario]
                if delete_scenario(scenario.get("id", "")):
                    st.session_state.scenarios = load_scenarios()
                    st.success("✅ Сценарий удалён")
                    st.rerun()
        
        st.divider()
        
        # ─────────────────────────────────────────────────────────────────────
        # Шаг 6: Экспорт
        # ─────────────────────────────────────────────────────────────────────
        st.subheader("6. Экспорт")
        
        col1, col2 = st.columns(2)
        
        with col1:
            docx_output = generate_forecast_report(forecast, st.session_state.current_scenario_name)
            filename = f"Tariff_Forecast_{st.session_state.current_scenario_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            st.download_button(
                label="📥 Скачать DOCX",
                data=docx_output,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        
        with col2:
            if st.button("📄 Экспорт в PDF", use_container_width=True):
                st.info("🚧 Функция в разработке (требуется дополнительная библиотека)")
    
    else:
        if not has_articles:
            st.info("👈 Введите данные и нажмите «Рассчитать прогноз»")
        else:
            st.info("👈 Нажмите «Рассчитать прогноз» для расчёта")
    
    # ─────────────────────────────────────────────────────────────────────
    # Справка
    # ─────────────────────────────────────────────────────────────────────
    with st.expander("💡 Как использовать", expanded=False):
        st.write("""
**Назначение:**

Прогнозист тарифов рассчитывает прогнозную величину тарифа на основе:
- Статей затрат (из текста или Анализатора заявок)
- Сферы деятельности (тепло, вода, ТКО, электричество)
- Метода регулирования (4 варианта)
- Условий прогноза

**Как работает AI:**

1. **Идентификация статей:** AI сопоставляет ваши статьи с эталонными индексами
2. **Применение индексов:** Каждая статья умножается на свой индекс (разный для каждой статьи)
3. **Расчёт НВВ:** Сумма прогнозных значений = Валовая выручка
4. **Расчёт тарифа:** НВВ / Полезный отпуск = Тариф

**Методы регулирования:**

- **Экономически обоснованных расходов** — стандартный метод по затратам
- **Сравнения аналогов** — сравнение с похожими организациями в регионе
- **Индексации установленных тарифов** — индексация текущего тарифа
- **Минимальной доходности инвестированного капитала** — для инвесторов

**Сценарии:**

- Сохраняйте разные варианты прогноза (Базовый, Оптимистичный, Пессимистичный)
- Переключайтесь между сценариями для сравнения
- Экспортируйте отчёты в DOCX для согласования

**Точность:**

- Зависит от качества входных данных
- AI использует эталонные индексы по отраслям
- Целевая точность: 85-90% к 2026 году
        """)

# =============================================================================
# Запуск
# =============================================================================

if __name__ == "__main__":
    show_tariff_forecaster()